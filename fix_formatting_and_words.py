"""
Fix paragraph formatting and update Hebrew words in YY Word documents.
1. Fix first paragraph after taglines (justify and indent)
2. Replace Hebrew words with proper modifiers
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time
import re

# Character definitions
ALEPH = chr(0x02BE)  # ʾ
AYIN = chr(0x02BF)   # ʿ
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

def normalize_quotes(s):
    """Normalize quote variations to straight apostrophe."""
    return s.replace(LEFT_CURLY, "'").replace(RIGHT_CURLY, "'").replace(chr(0x2032), "'")

# Replacement patterns
REPLACEMENTS = [
    # Fix previous: Zarowaʾ → Zarowaʿ (change Aleph to Ayin)
    (re.compile(rf'Zarowa{ALEPH}\b', re.IGNORECASE),
     lambda m: f'Zarowa{AYIN}' if m.group()[0].isupper() else f'zarowa{AYIN}'),

    # Miqra' → Miqraʿ
    (re.compile(r"Miqra['\u2018\u2019\u2032]\b", re.IGNORECASE),
     lambda m: f'Miqra{AYIN}' if m.group()[0].isupper() else f'miqra{AYIN}'),

    # Mow'ed → Mowʿed
    (re.compile(r"Mow['\u2018\u2019\u2032]ed\b"),
     lambda m: f'Mow{AYIN}ed'),

    # mow'ed → mowʿed
    (re.compile(r"mow['\u2018\u2019\u2032]ed\b"),
     lambda m: f'mow{AYIN}ed'),

    # 'eduwth → ʾeduwth
    (re.compile(r"['\u2018\u2019\u2032]eduwth\b", re.IGNORECASE),
     lambda m: f'{ALEPH}eduwth'),

    # 'eth → ʾeth
    (re.compile(r"['\u2018\u2019\u2032]eth\b", re.IGNORECASE),
     lambda m: f'{ALEPH}eth'),

    # 'aza'zel → ʾazaʾzel (Aleph at start and middle)
    (re.compile(r"['\u2018\u2019\u2032]aza['\u2018\u2019\u2032]zel\b", re.IGNORECASE),
     lambda m: f'{ALEPH}aza{ALEPH}zel'),

    # pa'am → paʿam
    (re.compile(r"pa['\u2018\u2019\u2032]am\b", re.IGNORECASE),
     lambda m: f'pa{AYIN}am'),

    # Zarowa' → Zarowaʿ
    (re.compile(r"Zarowa['\u2018\u2019\u2032]\b", re.IGNORECASE),
     lambda m: f'Zarowa{AYIN}' if m.group()[0].isupper() else f'zarowa{AYIN}'),

    # Yahowsha' → Yahowshaʿ (ensure Ayin)
    (re.compile(r"Yahowsha['\u2018\u2019\u2032]\b", re.IGNORECASE),
     lambda m: f'Yahowsha{AYIN}' if m.group()[0].isupper() else f'yahowsha{AYIN}'),

    # 'echad → ʾechad
    (re.compile(r"['\u2018\u2019\u2032]echad\b", re.IGNORECASE),
     lambda m: f'{ALEPH}echad'),

    # ta'en → taʿen
    (re.compile(r"ta['\u2018\u2019\u2032]en\b", re.IGNORECASE),
     lambda m: f'ta{AYIN}en'),
]

def apply_replacements(text):
    """Apply all replacement patterns."""
    original = text
    for pattern, replacement_func in REPLACEMENTS:
        text = pattern.sub(replacement_func, text)
    return text if text != original else None

def create_formatted_runs(run, new_text):
    """Create runs with proper font formatting for modifiers."""
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name or "Times New Roman"
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    para = run._element.getparent()
    run_idx = list(para).index(run._element)
    run.text = ""

    def create_run(text, use_yada_font=False):
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else font_name
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_modifier=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_modifier)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    for char in new_text:
        if char in (ALEPH, AYIN):
            flush_buffer(is_modifier=False)
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)

    flush_buffer(is_modifier=False)

def ensure_modifier_fonts(run):
    """Ensure existing modifiers use Yada Towrah font."""
    if not run.text:
        return False
    if ALEPH not in run.text and AYIN not in run.text:
        return False
    create_formatted_runs(run, run.text)
    return True

def is_tagline_paragraph(para):
    """
    Detect if paragraph is a tagline.
    Taglines are typically centered, bold, or have specific formatting.
    """
    if not para.text.strip():
        return False

    # Check if centered
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True

    # Check if all text is bold
    if para.runs:
        all_bold = all(run.bold for run in para.runs if run.text.strip())
        if all_bold and len(para.text.strip()) < 100:  # Short bold text
            return True

    return False

def fix_paragraph_formatting(doc, stats):
    """Fix first paragraph after taglines - make justified and indented."""
    paragraphs = doc.paragraphs

    for i in range(len(paragraphs) - 1):
        para = paragraphs[i]
        next_para = paragraphs[i + 1]

        # If current is tagline and next is content paragraph
        if is_tagline_paragraph(para):
            # Check if next paragraph is left-aligned or not indented
            if (next_para.text.strip() and
                (next_para.alignment == WD_ALIGN_PARAGRAPH.LEFT or
                 next_para.alignment is None)):

                # Fix alignment - set to justified
                next_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Fix indentation - set first line indent to 0.5 inches (36 points)
                if next_para.paragraph_format.first_line_indent is None or \
                   next_para.paragraph_format.first_line_indent.pt < 10:
                    next_para.paragraph_format.first_line_indent = Pt(36)
                    stats['para_fixes'] += 1

def process_paragraph(para, stats):
    """Process paragraph for word replacements."""
    for run in list(para.runs):
        if not run.text:
            continue

        new_text = apply_replacements(run.text)
        if new_text:
            create_formatted_runs(run, new_text)
            stats['replacements'] += 1
        else:
            if ensure_modifier_fonts(run):
                stats['font_fixes'] += 1

def process_document(docx_path, verbose=False):
    """Process a document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'replacements': 0, 'font_fixes': 0, 'para_fixes': 0}

    stats = {
        'replacements': 0,
        'font_fixes': 0,
        'para_fixes': 0,
        'verbose': verbose
    }

    # Fix paragraph formatting first
    fix_paragraph_formatting(doc, stats)

    # Process paragraphs for word replacements
    for para in doc.paragraphs:
        process_paragraph(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    repl = stats['replacements']
    fonts = stats['font_fixes']
    paras = stats['para_fixes']

    if repl > 0 or fonts > 0 or paras > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Words: {repl}, Fonts: {fonts}, Paras: {paras} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'replacements': 0, 'font_fixes': 0, 'para_fixes': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'replacements': repl, 'font_fixes': fonts, 'para_fixes': paras}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing formatting and updating Hebrew words...")
print()

docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-') and f.endswith('.docx')
    and '_updated' not in f and not f.startswith('~$')
]

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_repl = 0
total_fonts = 0
total_paras = 0

for path in docx_files:
    result = process_document(path)
    if result['replacements'] > 0 or result['font_fixes'] > 0 or result['para_fixes'] > 0:
        total_updated += 1
        total_repl += result['replacements']
        total_fonts += result['font_fixes']
        total_paras += result['para_fixes']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Word replacements: {total_repl}")
print(f"Font corrections: {total_fonts}")
print(f"Paragraph fixes: {total_paras}")
print(f"\nUpdated files: {docs_dir}/*_updated.docx")
