#!/usr/bin/env python3
"""
Fix translations with missing page numbers.

Uses a VBScript that iterates ALL paragraphs (not text-based Find), getting
page numbers directly via Range.Information(1). Matches COM paragraphs to
python-docx paragraphs by text content.
"""

import sys
import os
import re
import logging
import tempfile
import subprocess
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from parse_word_translations import (
    extract_translations_from_doc,
    setup_logging,
    DEFAULT_DIRECTORY,
)
from docx import Document

try:
    import psycopg2
except ImportError:
    print("ERROR: psycopg2 not installed")
    sys.exit(1)

try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: python-dotenv not installed")
    sys.exit(1)


def get_all_paragraph_pages(doc_path: str, timeout: int = 300) -> list:
    """
    Iterate ALL paragraphs in a document via COM and return [(page, text), ...].
    Uses Range.Information(1) for footer page numbers.
    """
    out = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8')
    out.close()

    vbs_content = f'''
Dim word, fso, outputFile
Set fso = CreateObject("Scripting.FileSystemObject")
Set word = CreateObject("Word.Application")
word.Visible = False
word.DisplayAlerts = 0

Set outputFile = fso.CreateTextFile("{out.name}", True, True)

WScript.StdErr.WriteLine "Opening document..."
Dim doc
Set doc = word.Documents.Open("{doc_path}", , True)

Dim paraCount
paraCount = doc.Paragraphs.Count
WScript.StdErr.WriteLine "Opened, paragraphs: " & paraCount

Dim i, pageNum, paraText
For i = 1 To paraCount
    On Error Resume Next
    pageNum = doc.Paragraphs(i).Range.Information(1)
    paraText = Left(doc.Paragraphs(i).Range.Text, 150)
    paraText = Replace(paraText, vbCr, "")
    paraText = Replace(paraText, vbLf, "")
    paraText = Replace(paraText, vbTab, " ")
    If Err.Number = 0 Then
        outputFile.WriteLine pageNum & "|" & paraText
    End If
    On Error GoTo 0
Next

doc.Close False
outputFile.Close
word.Quit
WScript.StdErr.WriteLine "Done - " & paraCount & " paragraphs"
'''

    vbs_file = tempfile.NamedTemporaryFile(mode='w', suffix='.vbs', delete=False, encoding='utf-8')
    vbs_file.write(vbs_content)
    vbs_file.close()

    try:
        result = subprocess.run(
            ["cscript", "//NoLogo", vbs_file.name],
            capture_output=True, text=True, timeout=timeout
        )

        if result.stderr:
            for line in result.stderr.strip().split('\n'):
                if line.strip():
                    logging.debug(f"  VBS: {line.strip()}")

        # Parse output: page|text
        paragraphs = []
        if os.path.exists(out.name):
            with open(out.name, 'r', encoding='utf-16', errors='replace') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('\ufeff'):
                        line = line.lstrip('\ufeff').strip()
                    if not line:
                        continue
                    pipe = line.find('|')
                    if pipe > 0:
                        try:
                            page = int(line[:pipe])
                            text = line[pipe+1:]
                            paragraphs.append((page, text))
                        except ValueError:
                            pass

        return paragraphs

    except subprocess.TimeoutExpired:
        logging.warning(f"  COM timed out after {timeout}s")
        try:
            subprocess.run(["powershell", "-Command",
                "Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue"],
                capture_output=True, timeout=10)
        except:
            pass
        return []
    except Exception as e:
        logging.warning(f"  COM failed: {e}")
        return []
    finally:
        for fp in [vbs_file.name, out.name]:
            try:
                os.unlink(fp)
            except:
                pass


def norm(text: str) -> str:
    """Normalize text for comparison."""
    return re.sub(r'\s+', ' ', text).strip()


def main():
    setup_logging(verbose=False)
    load_dotenv()

    logging.info("Fix Missing Page Numbers (all-paragraph iteration approach)")
    logging.info("=" * 60)

    conn = psycopg2.connect(
        host=os.getenv('POSTGRES_HOST', 'localhost'),
        port=os.getenv('POSTGRES_PORT', '5432'),
        user=os.getenv('POSTGRES_USER', 'postgres'),
        password=os.getenv('POSTGRES_PASSWORD'),
        database=os.getenv('POSTGRES_DB', 'yada'),
    )
    cur = conn.cursor()

    # Get books with missing pages
    cur.execute("""
        SELECT translation_book, COUNT(*) as missing
        FROM translation WHERE translation_page IS NULL
        GROUP BY translation_book ORDER BY missing DESC
    """)
    missing_books = cur.fetchall()
    total_missing = sum(row[1] for row in missing_books)
    logging.info(f"{total_missing} missing pages across {len(missing_books)} books\n")

    doc_dir = Path(DEFAULT_DIRECTORY)
    total_fixed = 0

    for book_name, missing_count in missing_books:
        doc_path = doc_dir / f"{book_name}.docx"
        if not doc_path.exists():
            logging.warning(f"  File not found: {doc_path}")
            continue

        file_size_mb = doc_path.stat().st_size / (1024 * 1024)
        logging.info(f"{book_name} ({file_size_mb:.1f}MB, {missing_count} missing)")

        # Step 1: Extract translations to get para_idx for missing ones
        translations = extract_translations_from_doc(doc_path)

        # Step 2: Get missing translation IDs from DB
        cur.execute("""
            SELECT translation_id, translation_text_word
            FROM translation
            WHERE translation_book = %s AND translation_page IS NULL
        """, (book_name,))
        missing_db = {row[0]: row[1] for row in cur.fetchall()}

        # Step 3: Map missing translations to their paragraph indices
        # Build text_word -> para_idx mapping
        missing_para_indices = {}  # translation_id -> para_idx
        for t in translations:
            text_word = t.get('text_word', '')
            para_idx = t.get('_para_idx')
            for tid, db_text in missing_db.items():
                if tid not in missing_para_indices and db_text == text_word:
                    missing_para_indices[tid] = para_idx
                    break

        if not missing_para_indices:
            logging.warning(f"  Could not match any missing translations to paragraphs")
            continue

        # Step 4: Get python-docx paragraph texts for the needed indices
        doc = Document(str(doc_path))
        pdx_texts = {}  # para_idx -> normalized text
        for tid, pidx in missing_para_indices.items():
            if pidx < len(doc.paragraphs):
                pdx_texts[pidx] = norm(doc.paragraphs[pidx].text)

        # Step 5: Get ALL paragraph pages from COM
        timeout = int(90 + file_size_mb * 50)

        # Try original, then clean copy
        com_paragraphs = []
        for attempt_label in ["original", "clean copy"]:
            if attempt_label == "original":
                attempt_path = str(doc_path)
            else:
                logging.info(f"  Trying clean copy...")
                try:
                    clean_dir = tempfile.mkdtemp()
                    clean_path = os.path.join(clean_dir, "clean.docx")
                    Document(str(doc_path)).save(clean_path)
                    attempt_path = clean_path
                except Exception as e:
                    logging.warning(f"  Clean copy failed: {e}")
                    break

            subprocess.run(["powershell", "-Command",
                "Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue"],
                capture_output=True, timeout=10)

            logging.info(f"  COM all-paragraphs ({attempt_label}, timeout {timeout}s)...")
            com_paragraphs = get_all_paragraph_pages(attempt_path, timeout=timeout)

            if attempt_label == "clean copy":
                try:
                    os.unlink(attempt_path)
                    os.rmdir(os.path.dirname(attempt_path))
                except:
                    pass

            if com_paragraphs:
                logging.info(f"  Got {len(com_paragraphs)} COM paragraphs")
                break
            else:
                logging.warning(f"  COM failed ({attempt_label})")

        if not com_paragraphs:
            logging.warning(f"  Could not get COM paragraphs for {book_name}")
            continue

        # Step 6: Build COM text -> page lookup
        # COM paragraphs are in order, normalize text for matching
        com_text_pages = []  # [(normalized_text, page), ...]
        for page, text in com_paragraphs:
            com_text_pages.append((norm(text), page))

        # Step 7: Match python-docx paragraphs to COM paragraphs
        fixed = 0
        for tid, pidx in missing_para_indices.items():
            pdx_text = pdx_texts.get(pidx, '')
            if not pdx_text:
                continue

            # Find best match in COM paragraphs
            best_page = None
            best_match_len = 0

            for com_text, com_page in com_text_pages:
                if not com_text:
                    continue

                # Check prefix match at various lengths
                check_len = min(80, len(pdx_text), len(com_text))
                if check_len >= 20 and pdx_text[:check_len] == com_text[:check_len]:
                    if check_len > best_match_len:
                        best_match_len = check_len
                        best_page = com_page

            if best_page is not None:
                cur.execute("""
                    UPDATE translation SET translation_page = %s
                    WHERE translation_id = %s AND translation_page IS NULL
                """, (best_page, tid))
                if cur.rowcount > 0:
                    fixed += 1

        conn.commit()
        total_fixed += fixed
        logging.info(f"  Fixed {fixed}/{missing_count}")

    # Final check
    cur.execute("SELECT COUNT(*) FROM translation WHERE translation_page IS NULL")
    still_missing = cur.fetchone()[0]

    logging.info(f"\n{'=' * 60}")
    logging.info(f"SUMMARY")
    logging.info(f"Total fixed: {total_fixed}/{total_missing}")
    logging.info(f"Still missing: {still_missing}")

    cur.close()
    conn.close()
    logging.info("Done.")


if __name__ == "__main__":
    main()
