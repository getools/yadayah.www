"""
Replace Hebrew transliterations with proper ALEPH (ʾ) and AYIN (ʿ) modifiers.
Simple approach: direct text replacement in runs.
Handles all apostrophe variations: ' ' ' `
Processes body text, tables, TOC, headers, and footers.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

ALEPH = chr(0x02BE)  # ʾ (right half ring)
AYIN = chr(0x02BF)   # ʿ (left half ring)

# All possible apostrophe variations
APOSTROPHES = ["'", "'", "'", "`", "´", "ʼ"]

# Define replacements: (old_text_with_apostrophe, new_text_with_modifier)
# We'll generate versions for all apostrophe variants
BASE_REPLACEMENTS = [
    # ALEPH replacements (ʾ)
    ("Yisrae'l", "Yisraʾel"),
    ("'Arek", "ʾArek"),
    ("'arek", "ʾarek"),
    ("'Akad", "ʾAkad"),
    ("'akad", "ʾakad"),
    ("'Ashuwr", "ʾAshuwr"),
    ("'ashuwr", "ʾashuwr"),
    ("'elowah", "ʾelowah"),
    ("'elowahym", "ʾelowahym"),
    ("'Adam", "ʾAdam"),
    ("ʾadam", "ʾAdam"),  # Fix case
    ("Miqra'", "Miqraʾ"),
    ("'Abraham's", "ʾAbraham's"),
    ("'Abshalowm's", "ʾAbshalowm's"),
    ("'Abraham", "ʾAbraham"),
    ("'Ayil", "ʾAyil"),
    ("'Abshalowm", "ʾAbshalowm"),
    ("'Amal", "ʾAmal"),
    ("'amal", "ʾamal"),
    ("'atem", "ʾatem"),
    ("'any", "ʾany"),
    ("'Edowm", "ʾEdowm"),
    ("'Amown", "ʾAmown"),
    ("'atah", "ʾatah"),
    ("'el", "ʾel"),
    ("'ElYah", "ʾElYah"),
    ("'anachnuw", "ʾanachnuw"),
    ("'Elowym", "ʾElowym"),
    ("'ozen", "ʾozen"),
    ("'Abyb", "ʾAbyb"),
    ("'ishah", "ʾishah"),
    ("'chasydy", "ʾchasydy"),

    # AYIN replacements (ʿ)
    ("Shin'ar", "Shinʿar"),
    ("shin'ar", "shinʿar"),
    ("'arabah", "ʿarabah"),
    ("'owd", "ʿowd"),
    ("'Eden", "ʿEden"),
    ("'aleichem", "ʿaleichem"),
    ("Ya'aqob", "Yaʿaqob"),
    ("Mow'ab", "Mowʿab"),
    ("Mow'edym", "Mowʿedym"),
    ("Yahowsha'", "Yahowshaʿ"),
    ("Naby'", "Nabyʿ"),
    ("Ma'achach", "Maʿachach"),
    ("Yow'ab", "Yowʿab"),
    ("Yisra'el", "Yisraʿel"),
    ("Ga'al", "Gaʿal"),
    ("yasha'", "yashaʿ"),
    ("Yisra'elites", "Yisraʿelites"),
    ("Mow'abites", "Mowʿabites"),
    ("Yada'", "Yadaʿ"),
    ("Zarowa'", "Zarowaʿ"),
    ("Yow'el", "Yowʿel"),
    ("Yasha'", "Yashaʿ"),
    ("Yahowshuwa'", "Yahowshuwaʿ"),
    ("mari'yth", "mariʿyth"),
    ("mal'akym", "malʿakym"),
    ("yesha'y", "yeshaʿy"),
    ("arba'ym", "arbaʿym"),
    ("sha'b", "shaʿb"),
    ("Howsha'", "Howshaʿ"),
    ("mits'ar", "mitsʿar"),
    ("Sha'uwl", "Shaʿuwl"),
]

def generate_all_variants():
    """Generate all replacement pairs for each apostrophe variant."""
    all_replacements = []

    for old_text, new_text in BASE_REPLACEMENTS:
        # For each apostrophe in old_text, create versions with all variants
        for apos in APOSTROPHES:
            variant_old = old_text.replace("'", apos)
            if variant_old != new_text:  # Avoid replacing if already correct
                all_replacements.append((variant_old, new_text))

    return all_replacements

ALL_REPLACEMENTS = generate_all_variants()

def process_run(run):
    """Replace text in a single run. Returns number of replacements."""
    if not run.text:
        return 0

    original_text = run.text
    new_text = original_text
    changes = 0

    for old_pattern, new_pattern in ALL_REPLACEMENTS:
        if old_pattern in new_text:
            new_text = new_text.replace(old_pattern, new_pattern)
            changes += new_text.count(new_pattern) - original_text.count(new_pattern)

    if new_text != original_text:
        run.text = new_text
        return 1 if changes > 0 else 0

    return 0

def ensure_modifier_font(para):
    """Ensure all ALEPH and AYIN modifiers have Yada Towrah font."""
    fixes = 0

    for element in para._element.iter():
        if element.tag == qn('w:t'):
            if element.text and (ALEPH in element.text or AYIN in element.text):
                # Check if this text element contains ONLY a modifier
                if element.text in [ALEPH, AYIN]:
                    run_elem = element.getparent()
                    if run_elem.tag == qn('w:r'):
                        rPr = run_elem.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            run_elem.insert(0, rPr)

                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.append(rFonts)

                        current_font = rFonts.get(qn('w:ascii'))
                        if current_font != 'Yada Towrah':
                            rFonts.set(qn('w:ascii'), 'Yada Towrah')
                            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
                            fixes += 1

    return fixes

def process_paragraph(para):
    """Process all runs in a paragraph."""
    fixes = 0

    for run in para.runs:
        fixes += process_run(run)

    # Ensure modifiers have correct font
    fixes += ensure_modifier_font(para)

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables (including TOC)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Replacing Hebrew transliterations with proper modifiers...")
print("(Processing body, tables, TOC, headers, footers)")
print()

# Get all YY files (exclude s07)
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in sorted(docx_files):
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total replacements: {total_fixes}")
