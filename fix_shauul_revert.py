"""Fix Shaʿuwlʿs back to Shaʿuwl's and verify TOC is clean."""
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')
HALF_RING_AYIN = '\u02BF'   # ʿ
QUOTE_CHARS = set("'\u2018\u2019`\u00B4\u02BC")

def fix_shauul_in_para(para):
    """Find Shaʿuwlʿs and revert the second ʿ back to '."""
    if not para.runs:
        return 0
    text = ''
    cmap = []
    for ri, run in enumerate(para.runs):
        for ci, ch in enumerate(run.text):
            text += ch
            cmap.append((ri, ci))

    # Match Sha + ʿ + uwl + ʿ + s (case-insensitive on Sha)
    pattern = re.compile(r'(?<![a-zA-Z])[Ss]ha\u02BFuwl\u02BFs(?![a-zA-Z])')
    count = 0
    for m in pattern.finditer(text):
        # The second ʿ is at position m.start() + 7 (S-h-a-ʿ-u-w-l-ʿ-s)
        pos = m.start() + 7
        if pos < len(cmap):
            ri, ci = cmap[pos]
            t = list(para.runs[ri].text)
            t[ci] = '\u2019'  # right single quote (matching Word style)
            para.runs[ri].text = ''.join(t)
            count += 1
    return count


def main():
    files = sorted(DOCS_DIR.glob('YY*.docx'))
    files = [f for f in files if '_backup' not in f.name and '_updated' not in f.name
             and '_halfring' not in f.name]

    total = 0
    for i, fp in enumerate(files):
        doc = Document(str(fp))
        count = 0

        for para in doc.paragraphs:
            count += fix_shauul_in_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count += fix_shauul_in_para(para)
        for section in doc.sections:
            try:
                if not section.header.is_linked_to_previous:
                    for para in section.header.paragraphs:
                        count += fix_shauul_in_para(para)
            except Exception:
                pass

        if count > 0:
            doc.save(str(fp))
            print(f'  {fp.name}: reverted {count} instances')
            total += count

    print(f'\nReverted {total} Shaʿuwlʿs -> Shaʿuwl\u2019s instances')

    # Now verify TOC entries
    print('\n--- Remaining quotes in TOC entries ---')
    remaining = 0
    for fp in files:
        doc = Document(str(fp))
        for i, para in enumerate(doc.paragraphs[:80]):
            has_field = bool(para._element.findall('.//' + qn('w:instrText')))
            style_name = para.style.name if para.style else ''
            is_toc = has_field or 'toc' in style_name.lower()
            if not is_toc:
                continue
            text = para.text
            quotes = [ch for ch in text if ch in QUOTE_CHARS]
            if quotes:
                remaining += 1
                print(f'  {fp.name} para[{i}]: "{text[:80]}"')

    if remaining == 0:
        print('  None! All TOC Hebrew transliterations have been updated.')
    else:
        print(f'\n  {remaining} TOC entries still have quotes (likely English possessives/contractions)')


if __name__ == '__main__':
    main()
