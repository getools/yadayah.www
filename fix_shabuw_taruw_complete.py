"""
Complete solution: Replace Shabuw'ah and Taruw'ah with proper modifiers
AND ensure all modifier characters use Yada Towrah font.

Phase 1: Replace text in ALL XML elements (including TOC hyperlinks)
Phase 2: Apply Yada Towrah font to all ʿ characters
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

AYIN = chr(0x02BF)  # ʿ

def phase1_replace_text(doc):
    """Replace apostrophes with proper modifiers in ALL text elements."""
    replacements = 0

    for element in doc.element.iter():
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            if element.text:
                original = element.text

                # Replace Shabuw'ah → Shabuwʿah
                text = re.sub(r"Shabuw['\u2018\u2019\u2032]ah",
                             lambda m: f'Shabuw{AYIN}ah' if m.group()[0].isupper() else f'shabuw{AYIN}ah',
                             original, flags=re.IGNORECASE)

                # Replace Taruw'ah → Taruwʿah
                text = re.sub(r"Taruw['\u2018\u2019\u2032]ah",
                             lambda m: f'Taruw{AYIN}ah' if m.group()[0].isupper() else f'taruw{AYIN}ah',
                             text, flags=re.IGNORECASE)

                if text != original:
                    element.text = text
                    replacements += 1

    return replacements

def phase2_apply_fonts(doc):
    """Apply Yada Towrah font to all ʿ modifier characters in runs."""
    font_fixes = 0

    def process_run(run):
        nonlocal font_fixes
        if not run.text or AYIN not in run.text:
            return

        # Split run into segments with modifiers using Yada Towrah font
        is_bold = run.bold
        is_italic = run.italic
        font_name = run.font.name or "Times New Roman"
        font_size = run.font.size
        font_color = None
        if run.font.color and run.font.color.rgb:
            font_color = run.font.color.rgb

        para = run._element.getparent()
        run_idx = list(para).index(run._element)
        text = run.text
        run.text = ""

        def create_run(content, use_yada_font=False):
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            if is_bold:
                b = OxmlElement('w:b')
                rPr.append(b)

            if is_italic:
                i_elem = OxmlElement('w:i')
                rPr.append(i_elem)

            rFonts = OxmlElement('w:rFonts')
            target_font = "Yada Towrah" if use_yada_font else font_name
            rFonts.set(qn('w:ascii'), target_font)
            rFonts.set(qn('w:hAnsi'), target_font)
            rPr.append(rFonts)

            if font_size:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(font_size.pt * 2)))
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
                rPr.append(szCs)

            if font_color:
                color = OxmlElement('w:color')
                color.set(qn('w:val'), str(font_color))
                rPr.append(color)

            new_run.append(rPr)

            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = content
            new_run.append(t)

            return new_run

        current_run_idx = run_idx
        buffer = []

        def flush_buffer(is_modifier=False):
            nonlocal current_run_idx
            if not buffer:
                return
            content = ''.join(buffer)
            new_run = create_run(content, use_yada_font=is_modifier)
            para.insert(current_run_idx + 1, new_run)
            current_run_idx += 1
            buffer.clear()

        for char in text:
            if char == AYIN:
                flush_buffer(is_modifier=False)
                buffer.append(char)
                flush_buffer(is_modifier=True)
            else:
                buffer.append(char)

        flush_buffer(is_modifier=False)
        font_fixes += 1

    # Process body paragraphs
    for para in doc.paragraphs:
        for run in list(para.runs):
            process_run(run)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        process_run(run)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                for run in list(para.runs):
                    process_run(run)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in list(para.runs):
                                process_run(run)

        if section.footer:
            for para in section.footer.paragraphs:
                for run in list(para.runs):
                    process_run(run)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in list(para.runs):
                                process_run(run)

    return font_fixes

def process_document(docx_path):
    """Process a single document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'replacements': 0, 'font_fixes': 0}

    # Phase 1: Replace text everywhere (including TOC)
    replacements = phase1_replace_text(doc)

    # Phase 2: Apply fonts to modifier characters
    font_fixes = phase2_apply_fonts(doc)

    elapsed = time.time() - start_time

    if replacements > 0 or font_fixes > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Text: {replacements}, Fonts: {font_fixes} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'replacements': 0, 'font_fixes': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'replacements': replacements, 'font_fixes': font_fixes}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Updating Shabuw'ah and Taruw'ah with proper modifiers...")
print("(Body, Tables, TOC, Headers, Footers)")
print()

# Get all YY files (excluding intermediate files)
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_repl = 0
total_fonts = 0

for path in docx_files:
    result = process_document(path)
    if result['replacements'] > 0 or result['font_fixes'] > 0:
        total_updated += 1
        total_repl += result['replacements']
        total_fonts += result['font_fixes']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Text replacements: {total_repl}")
print(f"Font corrections: {total_fonts}")
print(f"\nUpdated files: {docs_dir}\\*_updated.docx")
