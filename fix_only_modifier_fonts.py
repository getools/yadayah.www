"""
Apply Yada Towrah font ONLY to ʿ (Ayin) modifier characters.
Leave all other text with its original font.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

AYIN = chr(0x02BF)  # ʿ

def fix_modifier_font(t_elem, run_elem):
    """Set Yada Towrah font ONLY if this text element is exactly the ʿ character."""
    if t_elem.text != AYIN:
        return False

    # Get or create run properties
    rPr = run_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run_elem.insert(0, rPr)

    # Get or create rFonts element
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    # Check current font
    current_font = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')

    # Update to Yada Towrah if different
    if current_font != "Yada Towrah":
        rFonts.set(qn('w:ascii'), "Yada Towrah")
        rFonts.set(qn('w:hAnsi'), "Yada Towrah")
        return True

    return False

def process_paragraph(para):
    """Process all text elements in a paragraph."""
    fixes = 0

    for element in list(para._element.iter()):
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            # Get parent run
            run_elem = element.getparent()
            if run_elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                if fix_modifier_font(element, run_elem):
                    fixes += 1

    return fixes

def process_document(docx_path):
    """Process a document to fix modifier fonts."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} modifiers ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Setting Yada Towrah font ONLY on ʿ (Ayin) modifiers...")
print("(All other text keeps its original font)")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total ʿ modifiers with Yada Towrah font: {total_fixes}")
