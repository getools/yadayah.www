import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Checking specific paragraphs mentioned by user...")
print("=" * 80)

# Check YY-s02v05 chapter 1
print("\n1. YY-s02v05 - Chapter 1 paragraph:")
doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')
target = 'The fourth Miqra'
for i, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f"   Found at paragraph {i}")
        print(f"   Text: {para.text[:80]}...")
        print(f"   Alignment: {para.alignment}")
        indent = para.paragraph_format.first_line_indent
        print(f"   First-line indent: {indent.pt if indent else 'None'}pt")
        print(f"   Style: {para.style.name if para.style else 'None'}")

        # Check previous paragraphs to find the tagline
        print(f"\n   Checking previous paragraphs:")
        for j in range(max(0, i-5), i):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"   Para {j}: {p.text[:60]}...")
        break

# Check YY-s02v06 chapter 4
print("\n" + "=" * 80)
print("\n2. YY-s02v06 - Chapter 4 paragraph:")
import glob
s02v06_files = glob.glob(r'C:\users\joe\work\dev\yada\docs\YY-s02v06*.docx')
doc = Document(s02v06_files[0])
target = 'More than anything, Yahowah wants His people'
for i, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f"   Found at paragraph {i}")
        print(f"   Text: {para.text[:80]}...")
        print(f"   Alignment: {para.alignment}")
        indent = para.paragraph_format.first_line_indent
        print(f"   First-line indent: {indent.pt if indent else 'None'}pt")
        print(f"   Style: {para.style.name if para.style else 'None'}")

        # Check previous paragraphs
        print(f"\n   Checking previous paragraphs:")
        for j in range(max(0, i-5), i):
            p = doc.paragraphs[j]
            if p.text.strip():
                print(f"   Para {j}: {p.text[:60]}...")
        break

# Check About the Author
print("\n" + "=" * 80)
print("\n3. About the Author paragraph:")
doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')
target = 'Twenty-five years ago, Craig Winn was an entrepreneur'
for i, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f"   Found at paragraph {i}")
        print(f"   Text: {para.text[:80]}...")
        print(f"   Alignment: {para.alignment}")
        indent = para.paragraph_format.first_line_indent
        print(f"   First-line indent: {indent.pt if indent else 'None'}pt")
        print(f"   Style: {para.style.name if para.style else 'None'}")
        break

print("\n" + "=" * 80)
