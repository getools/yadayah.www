"""
Simple script to format all Normal style paragraphs from About the Author onwards.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time

def process_document(docx_path):
    """Process a document."""
    filename = os.path.basename(docx_path)
    print(f"\\nProcessing: {filename}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    # Find About the Author section
    start_idx = 0
    for i, para in enumerate(doc.paragraphs):
        if 'AUTHOR' in para.text.upper():
            start_idx = i + 1
            break

    print(f"  Starting at paragraph {start_idx}")

    formatted_count = 0

    # Process all paragraphs from About the Author onwards
    for i in range(start_idx, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # Skip empty
        if not text:
            continue

        # Skip edition info
        if text.startswith('Ver. 202'):
            continue

        # Skip centered
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue

        # Skip short lines (headings)
        if len(text) < 20 and not text[0].islower():
            continue

        # Skip taglines
        if text.endswith('...') or text.endswith('…'):
            continue

        # Skip if not Normal style
        if para.style and para.style.name != 'Normal':
            continue

        # Check current formatting
        is_justified = para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        indent = para.paragraph_format.first_line_indent
        is_indented = indent and indent.pt >= 35

        if is_justified and is_indented:
            continue

        # Format it
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Pt(36)
        formatted_count += 1

    elapsed = time.time() - start_time

    if formatted_count > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_fmt{ext}"
        doc.save(output_path)
        print(f"  ✓ Formatted {formatted_count} paragraphs ({elapsed:.1f}s)")
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return formatted_count

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Formatting all content paragraphs...")

import glob
docx_files = glob.glob(os.path.join(docs_dir, 'YY-*.docx'))
docx_files = [f for f in docx_files if '_fmt' not in f and not os.path.basename(f).startswith('~$')]

print(f"Found {len(docx_files)} documents\\n" + "=" * 80)

total_formatted = 0
for path in docx_files:
    count = process_document(path)
    total_formatted += count

print("\\n" + "=" * 80)
print(f"Total paragraphs formatted: {total_formatted}")
