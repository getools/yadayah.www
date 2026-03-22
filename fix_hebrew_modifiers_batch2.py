"""
Comprehensive Hebrew modifier replacement - Batch 2.
Replaces all apostrophe variants with proper ALEPH (ʾ U+02BE) and AYIN (ʿ U+02BF).
Handles same-run and cross-run patterns. Applies Yada Towrah font to all modifiers.
Patterns sorted by length (longest first) to avoid substring conflicts.
"""

import sys, os, glob, time, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ALEPH = chr(0x02BE)  # ʾ
AYIN = chr(0x02BF)   # ʿ

APOSTROPHE_VARIANTS = [
    chr(0x0027),  # Regular apostrophe '
    chr(0x2019),  # Right single quote
    chr(0x2018),  # Left single quote
    chr(0x0060),  # Backtick
    chr(0x00B4),  # Acute accent
    chr(0x02BC),  # Modifier letter apostrophe
]

# =====================================================================
# ALL replacement patterns: (base_pattern_with_ascii_apostrophe, replacement)
# The ' in base patterns will be expanded to all apostrophe variants.
# Sorted by length at generation time (longest first).
# =====================================================================
ALL_BASE_REPLACEMENTS = [
    # === OLD PATTERNS (from previous batches) ===
    # ALEPH leading
    ("'Abraham's", "\u02BEAbraham\u2019s"),
    ("'Abshalowm's", "\u02BEAbshalowm\u2019s"),
    ("'Abshalowm", "\u02BEAbshalowm"),
    ("'Abraham", "\u02BEAbraham"),
    ("'anachnuw", "\u02BEanachnuw"),
    ("'elowahym", "\u02BEelowahym"),
    ("'Ashuwr", "\u02BEAshuwr"),
    ("'ashuwr", "\u02BEashuwr"),
    ("'chasydy", "\u02BEchasydy"),
    ("'Elowym", "\u02BEElowym"),
    ("'elowah", "\u02BEelowah"),
    ("'Edowm", "\u02BEEdowm"),
    ("'Amnown", "\u02BEAmnown"),
    ("'Amown", "\u02BEAmown"),
    ("'Abram", "\u02BEAbram"),
    ("'ElYah", "\u02BEElYah"),
    ("'ishah", "\u02BEishah"),
    ("'Abyb", "\u02BEAbyb"),
    ("'Adam", "\u02BEAdam"),
    ("'Arek", "\u02BEArek"),
    ("'arek", "\u02BEarek"),
    ("'Akad", "\u02BEAkad"),
    ("'akad", "\u02BEakad"),
    ("'Amal", "\u02BEAmal"),
    ("'amal", "\u02BEamal"),
    ("'Ayil", "\u02BEAyil"),
    ("'atah", "\u02BEatah"),
    ("'atem", "\u02BEatem"),
    ("'ozen", "\u02BEozen"),
    ("'any", "\u02BEany"),
    ("'el", "\u02BEel"),
    # ALEPH mid/trailing (old)
    ("Yisrae'l", "Yisra\u02BEel"),
    ("Miqra'", "Miqra\u02BE"),
    # AYIN leading (old)
    ("'aleichem", "\u02BFaleichem"),
    ("'arabah", "\u02BFarabah"),
    ("'Eden", "\u02BFEden"),
    ("'owd", "\u02BFowd"),
    # AYIN mid/trailing (old)
    ("Yahowshuwa'", "Yahowshuwa\u02BF"),
    ("Yisra'elites", "Yisra\u02BFelites"),
    ("Mow'abites", "Mow\u02BFabites"),
    ("Yahowsha'", "Yahowsha\u02BF"),
    ("Ma'achach", "Ma\u02BFachach"),
    ("Mow'edym", "Mow\u02BFedym"),
    ("Yisra'el", "Yisra\u02BFel"),
    ("Ya'aqob", "Ya\u02BFaqob"),
    ("Sha'uwl", "Sha\u02BFuwl"),
    ("Howsha'", "Howsha\u02BF"),
    ("Shin'ar", "Shin\u02BFar"),
    ("shin'ar", "shin\u02BFar"),
    ("Zarowa'", "Zarowa\u02BF"),
    ("Mow'ab", "Mow\u02BFab"),
    ("Yow'ab", "Yow\u02BFab"),
    ("Yow'el", "Yow\u02BFel"),
    ("Yasha'", "Yasha\u02BF"),
    ("yasha'", "yasha\u02BF"),
    ("Ga'al", "Ga\u02BFal"),
    ("Yada'", "Yada\u02BF"),
    ("Naby'", "Naby\u02BF"),
    ("mari'yth", "mari\u02BFyth"),
    ("mal'akym", "mal\u02BFakym"),
    ("yesha'y", "yesha\u02BFy"),
    ("arba'ym", "arba\u02BFym"),
    ("mits'ar", "mits\u02BFar"),
    ("sha'b", "sha\u02BFb"),

    # === NEW PATTERNS (current batch) ===
    # ALEPH leading (new)
    ("'achazath", "\u02BEachazath"),
    ("'Uzyahuw", "\u02BEUzyahuw"),
    ("'uzyahuw", "\u02BEuzyahuw"),
    ("'Asherah", "\u02BEAsherah"),
    ("'Ephraym", "\u02BEEphraym"),
    ("'Amowts", "\u02BEAmowts"),
    ("'amowts", "\u02BEamowts"),
    ("'Amorah", "\u02BEAmorah"),
    ("'amorah", "\u02BEamorah"),
    ("'Uwryah", "\u02BEUwryah"),
    ("'uwryah", "\u02BEuwryah"),
    ("'adamah", "\u02BEadamah"),
    ("'Adonai", "\u02BEAdonai"),
    ("'adonai", "\u02BEadonai"),
    ("'enowsh", "\u02BEenowsh"),
    ("'Adony", "\u02BEAdony"),
    ("'adony", "\u02BEadony"),
    ("'Ashur", "\u02BEAshur"),
    ("'ashur", "\u02BEashur"),
    ("'Achaz", "\u02BEAchaz"),
    ("'achaz", "\u02BEachaz"),
    ("'anaph", "\u02BEanaph"),
    ("'orach", "\u02BEorach"),
    ("'Arar", "\u02BEArar"),
    ("'iysh", "\u02BEiysh"),
    ("'iwer", "\u02BEiwer"),
    ("'eth", "\u02BEeth"),
    ("'El", "\u02BEEl"),
    # AYIN leading (new)
    ("'ashaq", "\u02BFashaq"),
    ("'osheq", "\u02BFosheq"),
    ("'owlam", "\u02BFowlam"),
    ("'Ammi", "\u02BFAmmi"),
    ("'anah", "\u02BFanah"),
    ("'aqab", "\u02BFaqab"),
    ("'asah", "\u02BFasah"),

    # ALEPH mid/trailing (new)
    ("t-'anah", "t-\u02BEanah"),
    ("L'Chayim", "L\u02BEChayim"),
    ("ta'abah", "ta\u02BEabah"),
    ("t-naba'", "t-naba\u02BE"),
    ("Mal'aky", "Mal\u02BEaky"),
    ("showa'", "showa\u02BE"),
    ("Mal'ak", "Mal\u02BEak"),
    ("pito'm", "pito\u02BEm"),
    ("luwle'", "luwle\u02BE"),
    ("tsama'", "tsama\u02BE"),
    ("me'od", "me\u02BEod"),
    ("naba'", "naba\u02BE"),
    ("naby'", "naby\u02BE"),
    ("qara'", "qara\u02BE"),
    ("ro'eh", "ro\u02BEeh"),
    ("mala'", "mala\u02BE"),

    # AYIN mid/trailing (new)
    ("sha'shuwa'ym", "sha\u02BFshuwa\u02BFym"),
    ("tse'etsa'ym", "tse\u02BFetsa\u02BEym"),
    ("n-'anah-ythy", "n-\u02BFanah-ythy"),
    ("Yahowyada'", "Yahowyada\u02BF"),
    ("Yirma'yah", "Yirma\u02BFyah"),
    ("Yasha'yah", "Yasha\u02BFyah"),
    ("Phar'oah", "Phar\u02BFoah"),
    ("ma'alal", "ma\u02BFalal"),
    ("yow'ets", "yow\u02BFets"),
    ("sha'ar", "sha\u02BFar"),
    ("shama'", "shama\u02BF"),
    ("ma'owz", "ma\u02BFowz"),
    ("ya'ats", "ya\u02BFats"),
    ("ra'a'", "ra\u02BFra\u02BF"),
    ("raga'", "raga\u02BF"),
    ("Sha'a", "Sha\u02BFa"),
    ("sha'a", "sha\u02BFa"),
    ("rea'", "rea\u02BF"),
    ("ra'", "ra\u02BF"),

    # Special: pattern already contains modifier
    ("Lo\u02BE-'Ammi", "Lo\u02BE-\u02BFAmmi"),
    ("Lo'-'Ammi", "Lo\u02BE-\u02BFAmmi"),
]


def generate_all_variants():
    """Generate all apostrophe variant replacements, sorted longest first."""
    all_patterns = []
    seen = set()

    for old_pattern, new_pattern in ALL_BASE_REPLACEMENTS:
        for apos in APOSTROPHE_VARIANTS:
            variant = old_pattern.replace("'", apos)
            if variant != new_pattern and variant not in seen:
                all_patterns.append((variant, new_pattern))
                seen.add(variant)

    # Sort by search pattern length (longest first) to avoid substring conflicts
    all_patterns.sort(key=lambda x: len(x[0]), reverse=True)
    return all_patterns


ALL_PATTERNS = generate_all_variants()


def replace_in_text(text):
    """Apply all pattern replacements to text."""
    for old, new in ALL_PATTERNS:
        if old in text:
            text = text.replace(old, new)
    return text


def split_run_on_modifiers(run_elem, text):
    """Split run at every ALEPH/AYIN, applying Yada Towrah font to modifiers."""
    if ALEPH not in text and AYIN not in text:
        return False

    para = run_elem.getparent()
    if para is None:
        return False

    run_idx = list(para).index(run_elem)
    rPr = run_elem.find(qn('w:rPr'))

    # Build segments: alternating text and modifier
    segments = []
    buf = []
    for ch in text:
        if ch in (ALEPH, AYIN):
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            segments.append(('modifier', ch))
        else:
            buf.append(ch)
    if buf:
        segments.append(('text', ''.join(buf)))

    if len(segments) <= 1:
        return False

    def clone_rPr(exclude_font=False):
        new_rPr = OxmlElement('w:rPr')
        if rPr is not None:
            for child in rPr:
                if exclude_font and child.tag == qn('w:rFonts'):
                    continue
                new_rPr.append(copy.deepcopy(child))
        return new_rPr

    def make_run(content, is_modifier=False):
        r = OxmlElement('w:r')
        p = clone_rPr(exclude_font=is_modifier)
        if is_modifier:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Yada Towrah')
            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
            p.append(rFonts)
        r.append(p)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        r.append(t)
        return r

    para.remove(run_elem)
    idx = run_idx
    for seg_type, seg_text in segments:
        new_run = make_run(seg_text, is_modifier=(seg_type == 'modifier'))
        para.insert(idx, new_run)
        idx += 1

    return True


def ensure_modifier_font(run):
    """Ensure a single-modifier run has Yada Towrah font."""
    text = run.text or ""
    if text not in (ALEPH, AYIN):
        return False

    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            if rFonts.get(qn('w:ascii'), '') == 'Yada Towrah':
                return False  # Already correct

    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Yada Towrah')
    rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
    return True


def process_paragraph(para):
    """Process paragraph: replace patterns, handle cross-run, split modifiers."""
    fixes = 0

    # === PHASE 1: Same-run replacements ===
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        new_text = replace_in_text(text)
        if new_text != text:
            old_m = text.count(ALEPH) + text.count(AYIN)
            new_m = new_text.count(ALEPH) + new_text.count(AYIN)
            fixes += max(0, new_m - old_m)
            run.text = new_text

    # === PHASE 2: Cross-run replacements ===
    for _ in range(100):  # Safety limit
        full_text = para.text
        if not full_text:
            break

        found = False
        for search, replacement in ALL_PATTERNS:
            pos = full_text.find(search)
            if pos == -1:
                continue

            # Map character position to run index
            char_count = 0
            start_run = end_run = None
            start_off = end_off = 0

            for ri, run in enumerate(para.runs):
                rt = run.text or ""
                run_start = char_count
                run_end = char_count + len(rt)

                if start_run is None and pos >= run_start and pos < run_end:
                    start_run = ri
                    start_off = pos - run_start

                if start_run is not None and pos + len(search) <= run_end:
                    end_run = ri
                    end_off = pos + len(search) - run_start
                    break

                char_count = run_end

            if start_run is None or end_run is None:
                continue

            if start_run == end_run:
                # Same-run match missed in Phase 1 (shouldn't happen, but handle it)
                run = para.runs[start_run]
                old_text = run.text
                run.text = old_text[:start_off] + replacement + old_text[end_off:]
                old_m = old_text.count(ALEPH) + old_text.count(AYIN)
                new_m = run.text.count(ALEPH) + run.text.count(AYIN)
                fixes += max(0, new_m - old_m)
                found = True
                break
            else:
                # Cross-run: merge spanning runs
                combined = ""
                for ri in range(start_run, end_run + 1):
                    combined += para.runs[ri].text or ""

                new_combined = combined.replace(search, replacement, 1)
                para.runs[start_run].text = new_combined

                # Remove merged runs (reverse order to keep indices valid)
                p_elem = para._element
                for ri in range(end_run, start_run, -1):
                    p_elem.remove(para.runs[ri]._element)

                old_m = combined.count(ALEPH) + combined.count(AYIN)
                new_m = new_combined.count(ALEPH) + new_combined.count(AYIN)
                fixes += max(0, new_m - old_m)
                found = True
                break

        if not found:
            break

    # === PHASE 3: Split modifiers and ensure Yada Towrah font ===
    run_idx = 0
    while run_idx < len(para.runs):
        run = para.runs[run_idx]
        text = run.text or ""

        if ALEPH in text or AYIN in text:
            if text in (ALEPH, AYIN):
                # Single modifier character - just ensure font
                ensure_modifier_font(run)
                run_idx += 1
            else:
                # Mixed text + modifiers - split
                if split_run_on_modifiers(run._element, text):
                    continue  # Re-check from same index (new runs created)
                run_idx += 1
        else:
            run_idx += 1

    return fixes


def process_document(docx_path):
    """Process a single document."""
    doc_name = os.path.basename(docx_path)
    print(f"Processing: {doc_name}")
    start = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR opening: {e}")
        return 0

    fixes = 0

    # Body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Tables (including TOC)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Headers and footers
    for section in doc.sections:
        for part in [section.header, section.footer]:
            if part:
                for para in part.paragraphs:
                    fixes += process_paragraph(para)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                fixes += process_paragraph(para)

    elapsed = time.time() - start

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  No changes ({elapsed:.1f}s)")

    return fixes


# === Main ===
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print(f"Hebrew modifier replacement - Batch 2")
print(f"Patterns: {len(ALL_BASE_REPLACEMENTS)} base x {len(APOSTROPHE_VARIANTS)} variants = {len(ALL_PATTERNS)} total")
print(f"Processing body, tables, TOC, headers, footers")
print("=" * 80)

all_files = glob.glob(os.path.join(docs_dir, 'YY*.docx'))
docx_files = [f for f in all_files
              if not any(x in os.path.basename(f).lower()
                         for x in ['_updated', '_final', '_fmt', '_formatted',
                                   '_fixed', '_test', '_reverted', '_tagline',
                                   '_debug', '_backup'])]

print(f"Found {len(docx_files)} documents\n")

total = 0
for path in sorted(docx_files):
    total += process_document(path)

print("\n" + "=" * 80)
print(f"COMPLETE: {total} total fixes across {len(docx_files)} documents")
