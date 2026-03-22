"""Targeted fix for the 2 remaining ba'al instances in s04v04."""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ALEPH = chr(0x02BE)
AYIN = chr(0x02BF)
RIGHT_SINGLE = chr(0x2019)

doc_path = r'C:\users\joe\work\dev\yada\docs\YY-s04v04-Coming Home-Basar-Herald.docx'
doc = Document(doc_path)

fixes = 0

for pi in [1884, 2050]:
    para = doc.paragraphs[pi]
    p_elem = para._element

    # Get all direct w:r children
    runs = list(p_elem.findall(qn('w:r')))
    texts = []
    for r in runs:
        te = r.find(qn('w:t'))
        texts.append((te.text or '') if te is not None else '')

    full = ''.join(texts)
    search = f'ba{RIGHT_SINGLE}al'
    pos = full.find(search)

    print(f"Para {pi}: full text length={len(full)}")
    print(f"  Searching for {repr(search)}")
    print(f"  Found at position: {pos}")

    if pos == -1:
        print(f"  NOT FOUND - checking other variants...")
        for apos in [chr(0x0027), chr(0x2018), chr(0x0060)]:
            s = f'ba{apos}al'
            p2 = full.find(s)
            if p2 != -1:
                print(f"  Found with {repr(apos)} at position {p2}")
        continue

    # Map position to runs
    char_count = 0
    sr = er = None
    for ri, t in enumerate(texts):
        rs = char_count
        re = char_count + len(t)
        if sr is None and pos >= rs and pos < re:
            sr = ri
        if sr is not None and pos + len(search) <= re:
            er = ri
            break
        char_count = re

    print(f"  Start run: {sr}, End run: {er}")
    if sr is not None:
        print(f"  Run {sr} text: {repr(texts[sr])}")
    if er is not None and er != sr:
        print(f"  Run {er} text: {repr(texts[er])}")

    if sr is not None and er is not None and sr != er:
        # Merge runs sr to er
        combined = ''.join(texts[sr:er+1])
        replacement = f'ba{AYIN}al'
        new_combined = combined.replace(search, replacement, 1)
        print(f"  Combined: {repr(combined)}")
        print(f"  New: {repr(new_combined)}")

        # Update first run
        te = runs[sr].find(qn('w:t'))
        if te is None:
            te = OxmlElement('w:t')
            te.set(qn('xml:space'), 'preserve')
            runs[sr].append(te)
        te.text = new_combined
        te.set(qn('xml:space'), 'preserve')

        # Remove merged runs
        for ri in range(er, sr, -1):
            p_elem.remove(runs[ri])

        fixes += 1
        print(f"  FIXED!")

    elif sr is not None and er is not None and sr == er:
        # Same run
        old_t = texts[sr]
        new_t = old_t.replace(search, f'ba{AYIN}al', 1)
        te = runs[sr].find(qn('w:t'))
        te.text = new_t
        fixes += 1
        print(f"  FIXED (same run)!")

# Now split modifiers and apply font
if fixes > 0:
    for pi in [1884, 2050]:
        para = doc.paragraphs[pi]
        p_elem = para._element

        all_r = list(p_elem.findall(qn('w:r')))
        i = 0
        while i < len(all_r):
            r = all_r[i]
            te = r.find(qn('w:t'))
            text = (te.text or '') if te is not None else ''

            if text and (ALEPH in text or AYIN in text) and text not in (ALEPH, AYIN):
                # Split
                parent = r.getparent()
                idx = list(parent).index(r)
                rPr = r.find(qn('w:rPr'))

                segments = []
                buf = []
                for ch in text:
                    if ch in (ALEPH, AYIN):
                        if buf:
                            segments.append(('text', ''.join(buf)))
                            buf = []
                        segments.append(('modifier', ch))
                    else:
                        buf.append(ch)
                if buf:
                    segments.append(('text', ''.join(buf)))

                if len(segments) > 1:
                    def clone_rPr(exclude_font=False):
                        new_rPr = OxmlElement('w:rPr')
                        if rPr is not None:
                            for child in rPr:
                                if exclude_font and child.tag == qn('w:rFonts'):
                                    continue
                                new_rPr.append(copy.deepcopy(child))
                        return new_rPr

                    parent.remove(r)
                    cur = idx
                    for st, stxt in segments:
                        nr = OxmlElement('w:r')
                        p = clone_rPr(exclude_font=(st == 'modifier'))
                        if st == 'modifier':
                            rf = OxmlElement('w:rFonts')
                            rf.set(qn('w:ascii'), 'Yada Towrah')
                            rf.set(qn('w:hAnsi'), 'Yada Towrah')
                            p.append(rf)
                        nr.append(p)
                        t = OxmlElement('w:t')
                        t.set(qn('xml:space'), 'preserve')
                        t.text = stxt
                        nr.append(t)
                        parent.insert(cur, nr)
                        cur += 1

                    all_r = list(p_elem.findall(qn('w:r')))
                    continue
            elif text in (ALEPH, AYIN):
                # Ensure font
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), 'Yada Towrah')
                rFonts.set(qn('w:hAnsi'), 'Yada Towrah')

            i += 1

    doc.save(doc_path)
    print(f"\nSaved! {fixes} fixes applied.")
else:
    print("\nNo fixes needed.")
