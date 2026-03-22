"""
Update more specific Hebrew words in all YY Word documents:
- Ya'aqob → Yaʿaqob
- Yisra'el → Yisraʿel
- Taruw'ah → Taruwʿah
- Sha'uwl → Shaʿuwl
- Shaʿuwlʿs → Shaʿuwl's
- Howsha' → Howshaʾ
- Howsha's → Howshaʾ's
- She'owl → Sheʿowl
- Lo' → Loʾ, lo' → loʾ
- Na'aph → Naʾaph
- Sha'aruwr → Shaʿaruwr
- Yahowsha' → Yahowshaʾ
- 'Azab → ʾAzab
- 'aharown → ʾAharown
- 'Aharown → ʾAharown
- Ensure all ʾ and ʿ use Yada Towrah font
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

# Character definitions
ALEPH = chr(0x02BE)  # ʾ (modifier letter right half ring)
AYIN = chr(0x02BF)   # ʿ (modifier letter left half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

def normalize_quotes(s):
    """Normalize all quote variations to straight apostrophe."""
    return s.replace(LEFT_CURLY, "'").replace(RIGHT_CURLY, "'").replace(chr(0x2032), "'")

# Replacement patterns (search -> replacement)
# Order matters - do more specific patterns first
REPLACEMENTS = [
    # Fix: Shaʿuwlʿs → Shaʿuwl's (modifier shouldn't be on possessive 's)
    (re.compile(rf'Sha{AYIN}uwl{AYIN}s\b'),
     lambda m: f'Sha{AYIN}uwl\'s'),

    # Fix: Howshaʾ's pattern (if exists with modifier before 's)
    (re.compile(rf'Howsha{ALEPH}{AYIN}s\b'),
     lambda m: f'Howsha{ALEPH}\'s'),

    # Ya'aqob → Yaʿaqob
    (re.compile(r"Ya['\u2018\u2019\u2032]aqob\b", re.IGNORECASE),
     lambda m: f'Ya{AYIN}aqob' if m.group()[0].isupper() else f'ya{AYIN}aqob'),

    # Yisra'el → Yisraʿel
    (re.compile(r"Yisra['\u2018\u2019\u2032]el\b", re.IGNORECASE),
     lambda m: f'Yisra{AYIN}el' if m.group()[0].isupper() else f'yisra{AYIN}el'),

    # Taruw'ah → Taruwʿah
    (re.compile(r"Taruw['\u2018\u2019\u2032]ah\b", re.IGNORECASE),
     lambda m: f'Taruw{AYIN}ah' if m.group()[0].isupper() else f'taruw{AYIN}ah'),

    # Sha'uwl → Shaʿuwl
    (re.compile(r"Sha['\u2018\u2019\u2032]uwl\b", re.IGNORECASE),
     lambda m: f'Sha{AYIN}uwl' if m.group()[0].isupper() else f'sha{AYIN}uwl'),

    # Howsha' → Howshaʾ (no 's at end)
    (re.compile(r"Howsha['\u2018\u2019\u2032](?!s)\b", re.IGNORECASE),
     lambda m: f'Howsha{ALEPH}' if m.group()[0].isupper() else f'howsha{ALEPH}'),

    # Howsha's → Howshaʾ's (with 's)
    (re.compile(r"Howsha['\u2018\u2019\u2032]s\b", re.IGNORECASE),
     lambda m: f'Howsha{ALEPH}\'s' if m.group()[0].isupper() else f'howsha{ALEPH}\'s'),

    # She'owl → Sheʿowl
    (re.compile(r"She['\u2018\u2019\u2032]owl\b", re.IGNORECASE),
     lambda m: f'She{AYIN}owl' if m.group()[0].isupper() else f'she{AYIN}owl'),

    # Lo' → Loʾ (uppercase L)
    (re.compile(r"Lo['\u2018\u2019\u2032]\b"),
     lambda m: f'Lo{ALEPH}'),

    # lo' → loʾ (lowercase l)
    (re.compile(r"lo['\u2018\u2019\u2032]\b"),
     lambda m: f'lo{ALEPH}'),

    # Na'aph → Naʾaph
    (re.compile(r"Na['\u2018\u2019\u2032]aph\b", re.IGNORECASE),
     lambda m: f'Na{ALEPH}aph' if m.group()[0].isupper() else f'na{ALEPH}aph'),

    # Sha'aruwr → Shaʿaruwr
    (re.compile(r"Sha['\u2018\u2019\u2032]aruwr\b", re.IGNORECASE),
     lambda m: f'Sha{AYIN}aruwr' if m.group()[0].isupper() else f'sha{AYIN}aruwr'),

    # Yahowsha' → Yahowshaʾ
    (re.compile(r"Yahowsha['\u2018\u2019\u2032]\b", re.IGNORECASE),
     lambda m: f'Yahowsha{ALEPH}' if m.group()[0].isupper() else f'yahowsha{ALEPH}'),

    # 'Azab → ʾAzab
    (re.compile(r"['\u2018\u2019\u2032]Azab\b"),
     lambda m: f'{ALEPH}Azab'),

    # 'aharown → ʾAharown (capitalize A)
    (re.compile(r"['\u2018\u2019\u2032]aharown\b", re.IGNORECASE),
     lambda m: f'{ALEPH}Aharown'),

    # 'Aharown → ʾAharown (already capitalized)
    (re.compile(r"['\u2018\u2019\u2032]Aharown\b"),
     lambda m: f'{ALEPH}Aharown'),
]

def apply_replacements(text):
    """Apply all replacement patterns to text."""
    original = text
    for pattern, replacement_func in REPLACEMENTS:
        text = pattern.sub(replacement_func, text)
    return text if text != original else None

def create_formatted_runs(run, new_text):
    """
    Create new runs with proper font formatting.
    All ʾ and ʿ characters use "Yada Towrah" font.
    """
    # Save original formatting
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name or "Times New Roman"
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    # Get parent paragraph and run position
    para = run._element.getparent()
    run_idx = list(para).index(run._element)

    # Clear original run
    run.text = ""

    def create_run(text, use_yada_font=False):
        """Create a run with specified formatting."""
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        # Italic
        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Font
        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else font_name
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        # Font size
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        # Color
        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    # Split text at modifier characters
    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_modifier=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_modifier)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    for char in new_text:
        if char in (ALEPH, AYIN):
            # Flush regular text
            flush_buffer(is_modifier=False)
            # Add modifier and flush
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)

    # Flush remaining
    flush_buffer(is_modifier=False)

def ensure_modifier_fonts(run):
    """
    Ensure any existing ʾ or ʿ characters in a run use Yada Towrah font.
    Returns True if run was modified.
    """
    if not run.text:
        return False

    # Check if text contains modifiers
    if ALEPH not in run.text and AYIN not in run.text:
        return False

    # If run contains modifiers, need to split it
    create_formatted_runs(run, run.text)
    return True

def process_paragraph(para, stats):
    """Process a paragraph: apply replacements and ensure fonts."""
    runs_to_process = list(para.runs)

    for run in runs_to_process:
        if not run.text:
            continue

        # Try to apply replacements
        new_text = apply_replacements(run.text)
        if new_text:
            create_formatted_runs(run, new_text)
            stats['replacements'] += 1
            if stats['verbose']:
                print(f"    {run.text} -> {new_text}")
        else:
            # Even if no replacement, ensure modifier fonts are correct
            if ensure_modifier_fonts(run):
                stats['font_fixes'] += 1

def process_document(docx_path, verbose=False):
    """Process a single Word document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: Could not open document: {e}")
        return {'replacements': 0, 'font_fixes': 0}

    stats = {
        'replacements': 0,
        'font_fixes': 0,
        'verbose': verbose
    }

    # Process paragraphs
    for para in doc.paragraphs:
        process_paragraph(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    repl_count = stats['replacements']
    font_count = stats['font_fixes']

    if repl_count > 0 or font_count > 0:
        # Save with _updated suffix
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Replacements: {repl_count}, Font fixes: {font_count}, saved ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: Could not save document: {e}")
            return {'replacements': 0, 'font_fixes': 0}
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return {'replacements': repl_count, 'font_fixes': font_count}

# Main execution
docs_dir = r'C:\users\joe\work\dev\yada\docs'
if not os.path.exists(docs_dir):
    print(f"ERROR: Documents directory not found: {docs_dir}")
    sys.exit(1)

print("Updating more specific Hebrew words in YY documents...")
print("Replacements:")
print("  Ya'aqob → Yaʿaqob")
print("  Yisra'el → Yisraʿel")
print("  Taruw'ah → Taruwʿah")
print("  Sha'uwl → Shaʿuwl")
print("  Shaʿuwlʿs → Shaʿuwl's")
print("  Howsha' → Howshaʾ")
print("  Howsha's → Howshaʾ's")
print("  She'owl → Sheʿowl")
print("  Lo' → Loʾ, lo' → loʾ")
print("  Na'aph → Naʾaph")
print("  Sha'aruwr → Shaʿaruwr")
print("  Yahowsha' → Yahowshaʾ")
print("  'Azab → ʾAzab")
print("  'aharown → ʾAharown")
print("  'Aharown → ʾAharown")
print("  + Ensuring all ʾ and ʿ use Yada Towrah font")
print()

# Find all YY documents
docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-')
    and f.endswith('.docx')
    and not '_updated' in f
    and not f.startswith('~$')
]

print(f"Found {len(docx_files)} YY documents to process\n")
print("=" * 80)

total_files_updated = 0
total_replacements = 0
total_font_fixes = 0

for docx_path in docx_files:
    result = process_document(docx_path, verbose=False)
    if result['replacements'] > 0 or result['font_fixes'] > 0:
        total_files_updated += 1
        total_replacements += result['replacements']
        total_font_fixes += result['font_fixes']

print("\n" + "=" * 80)
print("UPDATE COMPLETE")
print("=" * 80)
print(f"  Files processed: {len(docx_files)}")
print(f"  Files updated: {total_files_updated}")
print(f"  Word replacements: {total_replacements}")
print(f"  Font corrections: {total_font_fixes}")
print(f"\nUpdated files saved with '_updated.docx' suffix in: {docs_dir}")
