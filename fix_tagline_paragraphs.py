"""
Fix first paragraph after taglines to be justified and indented.

Taglines are:
- Italicized text
- Left-aligned
- End with ... (ellipsis)
- Only in main content area (pages with footer numbers, not TOC/Resources)
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from lxml import etree
import time

def get_content_start_idx(doc):
    """
    Get the paragraph index where main content starts.
    YY docs have: front matter (lowerRoman) → TOC (decimal start=1) → content (decimal start=1)
    Content starts at LAST section with pgNumType fmt=decimal start=1.
    """
    # Parse the document XML to find section properties
    body = doc._element.body
    sections = []

    # Find all section properties in the document
    for i, para in enumerate(doc.paragraphs):
        para_element = para._element
        # Check if this paragraph has a section break
        pPr = para_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            sectPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sectPr is not None:
                pgNumType = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
                if pgNumType is not None:
                    fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
                    start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
                    sections.append({
                        'para_idx': i,
                        'fmt': fmt,
                        'start': start
                    })

    # Also check document-level sectPr (last section)
    sectPr = body.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
    if sectPr is not None:
        pgNumType = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
        if pgNumType is not None:
            fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
            start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
            sections.append({
                'para_idx': len(doc.paragraphs) - 1,
                'fmt': fmt,
                'start': start
            })

    # Find LAST section with decimal start=1 (this is where content begins)
    content_start = 0
    for section in reversed(sections):
        if section['fmt'] == 'decimal' and section['start'] == '1':
            content_start = section['para_idx']
            break

    return content_start

def is_tagline_paragraph(para):
    """
    Detect if paragraph is a tagline:
    - Italicized text (runs italic OR style has italic)
    - Left-aligned
    - Ends with ... (ellipsis)
    """
    if not para.text.strip():
        return False

    # Check if ends with ... (ellipsis)
    text = para.text.strip()
    if not text.endswith('...') and not text.endswith('…'):
        return False

    # Check if left-aligned (or None, which defaults to left)
    if para.alignment not in (WD_ALIGN_PARAGRAPH.LEFT, None):
        return False

    # Check if paragraph style has italic
    if para.style and hasattr(para.style, 'font') and para.style.font.italic:
        return True

    # Check if runs are italic
    if not para.runs:
        return False

    # Count italic vs non-italic characters
    italic_chars = 0
    total_chars = 0

    for run in para.runs:
        if run.text.strip():
            total_chars += len(run.text.strip())
            if run.italic:
                italic_chars += len(run.text.strip())

    # Consider it italic if >80% of characters are italic
    if total_chars > 0 and italic_chars / total_chars > 0.8:
        return True

    return False

def fix_paragraph_formatting(doc, content_start_idx, stats):
    """Fix first paragraph after taglines - make justified and indented."""
    paragraphs = doc.paragraphs

    for i in range(content_start_idx, len(paragraphs) - 1):
        para = paragraphs[i]

        # If current is tagline, find next non-empty paragraph
        if is_tagline_paragraph(para):
            # Find next non-empty paragraph
            next_para = None
            next_idx = i + 1
            while next_idx < len(paragraphs):
                candidate = paragraphs[next_idx]
                if candidate.text.strip():
                    next_para = candidate
                    break
                next_idx += 1

            if next_para is None:
                continue

            # Check if next paragraph needs fixing
            if (next_para.alignment == WD_ALIGN_PARAGRAPH.LEFT or
                next_para.alignment is None):

                # Fix alignment - set to justified
                next_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Fix indentation - set first line indent to 0.5 inches (36 points)
                if next_para.paragraph_format.first_line_indent is None or \
                   next_para.paragraph_format.first_line_indent.pt < 10:
                    next_para.paragraph_format.first_line_indent = Pt(36)
                    stats['para_fixes'] += 1

                    if stats['verbose']:
                        print(f"    Fixed para after tagline: {para.text[:50]}...")

def process_document(docx_path, verbose=False):
    """Process a document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'para_fixes': 0}

    stats = {
        'para_fixes': 0,
        'verbose': verbose
    }

    # Get content start index (skip front matter and TOC)
    content_start_idx = get_content_start_idx(doc)

    if verbose:
        print(f"  Content starts at paragraph index: {content_start_idx}")

    # Fix paragraph formatting
    fix_paragraph_formatting(doc, content_start_idx, stats)

    elapsed = time.time() - start_time
    fixes = stats['para_fixes']

    if fixes > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_tagline_fixed{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Fixed {fixes} paragraphs ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'para_fixes': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'para_fixes': fixes}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing paragraphs after taglines (italicized text ending with ...)...")
print()

docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-') and f.endswith('.docx')
    and '_tagline_fixed' not in f and not f.startswith('~$')
]

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_fixes = 0

for path in docx_files:
    result = process_document(path, verbose=False)
    if result['para_fixes'] > 0:
        total_updated += 1
        total_fixes += result['para_fixes']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Paragraph fixes: {total_fixes}")
print(f"\nUpdated files: {docs_dir}/*_tagline_fixed.docx")
