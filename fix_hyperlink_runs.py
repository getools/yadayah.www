"""
Fix Hebrew modifiers inside hyperlink elements (TOC entries).
These are not accessible via para.runs, only via XML.
"""

import sys, os, glob, time, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ALEPH = chr(0x02BE)
AYIN = chr(0x02BF)

APOSTROPHE_VARIANTS = [
    chr(0x0027), chr(0x2019), chr(0x2018),
    chr(0x0060), chr(0x00B4), chr(0x02BC),
]

# Import patterns from batch2 script
sys.path.insert(0, os.path.dirname(__file__))
from fix_hebrew_modifiers_batch2 import ALL_PATTERNS


def get_all_runs(para):
    """Get ALL w:r elements in paragraph, including those inside hyperlinks."""
    return para._element.findall('.//' + qn('w:r'))


def get_run_text(r_elem):
    """Get text from a w:r element."""
    t = r_elem.find(qn('w:t'))
    if t is None:
        return ''
    return t.text or ''


def set_run_text(r_elem, text):
    """Set text on a w:r element."""
    t = r_elem.find(qn('w:t'))
    if t is None:
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        r_elem.append(t)
    t.text = text
    t.set(qn('xml:space'), 'preserve')


def replace_in_text(text):
    """Apply all pattern replacements to text."""
    for old, new in ALL_PATTERNS:
        if old in text:
            text = text.replace(old, new)
    return text


def split_run_on_modifiers(r_elem):
    """Split a w:r element at modifier positions with Yada Towrah font."""
    text = get_run_text(r_elem)
    if not text or (ALEPH not in text and AYIN not in text):
        return False

    # If it's already just a single modifier, ensure font
    if text in (ALEPH, AYIN):
        ensure_font(r_elem)
        return False

    parent = r_elem.getparent()
    if parent is None:
        return False

    idx = list(parent).index(r_elem)
    rPr = r_elem.find(qn('w:rPr'))

    # Build segments
    segments = []
    buf = []
    for ch in text:
        if ch in (ALEPH, AYIN):
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            segments.append(('modifier', ch))
        else:
            buf.append(ch)
    if buf:
        segments.append(('text', ''.join(buf)))

    if len(segments) <= 1:
        return False

    def clone_rPr(exclude_font=False):
        new_rPr = OxmlElement('w:rPr')
        if rPr is not None:
            for child in rPr:
                if exclude_font and child.tag == qn('w:rFonts'):
                    continue
                new_rPr.append(copy.deepcopy(child))
        return new_rPr

    def make_run(content, is_modifier=False):
        r = OxmlElement('w:r')
        p = clone_rPr(exclude_font=is_modifier)
        if is_modifier:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Yada Towrah')
            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
            p.append(rFonts)
        r.append(p)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        r.append(t)
        return r

    parent.remove(r_elem)
    cur_idx = idx
    for seg_type, seg_text in segments:
        new_run = make_run(seg_text, is_modifier=(seg_type == 'modifier'))
        parent.insert(cur_idx, new_run)
        cur_idx += 1

    return True


def ensure_font(r_elem):
    """Ensure a modifier run has Yada Towrah font."""
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r_elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Yada Towrah')
    rFonts.set(qn('w:hAnsi'), 'Yada Towrah')


def process_paragraph_hyperlinks(para):
    """Process ALL runs in paragraph including those in hyperlinks."""
    fixes = 0
    p_elem = para._element

    # Find all hyperlink elements
    hyperlinks = p_elem.findall(qn('w:hyperlink'))

    for hl in hyperlinks:
        # Process each run inside the hyperlink
        runs = hl.findall(qn('w:r'))

        for r_elem in runs:
            text = get_run_text(r_elem)
            if not text:
                continue

            new_text = replace_in_text(text)
            if new_text != text:
                old_m = text.count(ALEPH) + text.count(AYIN)
                new_m = new_text.count(ALEPH) + new_text.count(AYIN)
                fixes += max(0, new_m - old_m)
                set_run_text(r_elem, new_text)

        # Cross-run within hyperlink
        for _ in range(50):
            runs = hl.findall(qn('w:r'))
            run_texts = [(r, get_run_text(r)) for r in runs]

            full_text = ''.join(t for _, t in run_texts)
            found = False

            for search, replacement in ALL_PATTERNS:
                pos = full_text.find(search)
                if pos == -1:
                    continue

                # Find spanning runs
                char_count = 0
                start_ri = end_ri = None
                for ri, (r, t) in enumerate(run_texts):
                    rs = char_count
                    re = char_count + len(t)

                    if start_ri is None and pos >= rs and pos < re:
                        start_ri = ri

                    if start_ri is not None and pos + len(search) <= re:
                        end_ri = ri
                        break

                    char_count = re

                if start_ri is None or end_ri is None:
                    continue

                if start_ri == end_ri:
                    r, t = run_texts[start_ri]
                    so = pos - sum(len(run_texts[i][1]) for i in range(start_ri))
                    eo = so + len(search)
                    new_t = t[:so] + replacement + t[eo:]
                    set_run_text(r, new_t)
                    old_m = t.count(ALEPH) + t.count(AYIN)
                    new_m = new_t.count(ALEPH) + new_t.count(AYIN)
                    fixes += max(0, new_m - old_m)
                else:
                    # Merge spanning runs
                    combined = ''.join(run_texts[i][1] for i in range(start_ri, end_ri + 1))
                    new_combined = combined.replace(search, replacement, 1)
                    set_run_text(run_texts[start_ri][0], new_combined)

                    for ri in range(end_ri, start_ri, -1):
                        hl.remove(run_texts[ri][0])

                    old_m = combined.count(ALEPH) + combined.count(AYIN)
                    new_m = new_combined.count(ALEPH) + new_combined.count(AYIN)
                    fixes += max(0, new_m - old_m)

                found = True
                break

            if not found:
                break

        # Split modifiers in hyperlink runs
        runs = hl.findall(qn('w:r'))
        i = 0
        while i < len(runs):
            r = runs[i]
            text = get_run_text(r)
            if text and (ALEPH in text or AYIN in text):
                if text in (ALEPH, AYIN):
                    ensure_font(r)
                    i += 1
                else:
                    if split_run_on_modifiers(r):
                        runs = hl.findall(qn('w:r'))
                        continue
                    i += 1
            else:
                i += 1

    return fixes


def process_document(docx_path):
    doc_name = os.path.basename(docx_path)
    print(f"Processing: {doc_name}")
    start = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    for para in doc.paragraphs:
        fixes += process_paragraph_hyperlinks(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph_hyperlinks(para)

    for section in doc.sections:
        for part in [section.header, section.footer]:
            if part:
                for para in part.paragraphs:
                    fixes += process_paragraph_hyperlinks(para)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                fixes += process_paragraph_hyperlinks(para)

    elapsed = time.time() - start

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  No changes ({elapsed:.1f}s)")

    return fixes


# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing hyperlink/TOC runs...")
print("=" * 80)

all_files = glob.glob(os.path.join(docs_dir, 'YY*.docx'))
docx_files = [f for f in all_files
              if not any(x in os.path.basename(f).lower()
                         for x in ['_updated', '_final', '_fmt', '_formatted',
                                   '_fixed', '_test', '_reverted', '_tagline',
                                   '_debug', '_backup'])]

print(f"Found {len(docx_files)} documents\n")

total = 0
for path in sorted(docx_files):
    total += process_document(path)

print("\n" + "=" * 80)
print(f"COMPLETE: {total} hyperlink fixes across {len(docx_files)} documents")
