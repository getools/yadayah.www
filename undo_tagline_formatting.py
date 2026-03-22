"""
Undo tagline paragraph formatting changes.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from lxml import etree
import time

def get_content_start_idx(doc):
    """Get paragraph index where main content starts."""
    body = doc._element.body
    sections = []

    # Find all section properties
    for i, para in enumerate(doc.paragraphs):
        para_element = para._element
        pPr = para_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            sectPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sectPr is not None:
                pgNumType = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
                if pgNumType is not None:
                    fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
                    start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
                    sections.append({'para_idx': i, 'fmt': fmt, 'start': start})

    sectPr = body.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
    if sectPr is not None:
        pgNumType = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
        if pgNumType is not None:
            fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
            start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
            sections.append({'para_idx': len(doc.paragraphs) - 1, 'fmt': fmt, 'start': start})

    content_start = 0
    for section in reversed(sections):
        if section['fmt'] == 'decimal' and section['start'] == '1':
            content_start = section['para_idx']
            break

    return content_start

def is_tagline_paragraph(para):
    """Detect if paragraph is a tagline."""
    if not para.text.strip():
        return False

    text = para.text.strip()
    if not text.endswith('...') and not text.endswith('…'):
        return False

    if para.alignment not in (WD_ALIGN_PARAGRAPH.LEFT, None):
        return False

    if para.style and hasattr(para.style, 'font') and para.style.font.italic:
        return True

    if not para.runs:
        return False

    italic_chars = 0
    total_chars = 0

    for run in para.runs:
        if run.text.strip():
            total_chars += len(run.text.strip())
            if run.italic:
                italic_chars += len(run.text.strip())

    if total_chars > 0 and italic_chars / total_chars > 0.8:
        return True

    return False

def undo_paragraph_formatting(doc, content_start_idx, stats):
    """Undo formatting on paragraphs after taglines."""
    paragraphs = doc.paragraphs

    for i in range(content_start_idx, len(paragraphs) - 1):
        para = paragraphs[i]

        if is_tagline_paragraph(para):
            # Find next non-empty paragraph
            next_para = None
            next_idx = i + 1
            while next_idx < len(paragraphs):
                candidate = paragraphs[next_idx]
                if candidate.text.strip():
                    next_para = candidate
                    break
                next_idx += 1

            if next_para is None:
                continue

            # Undo if justified with indent
            if next_para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                indent = next_para.paragraph_format.first_line_indent
                if indent and indent.pt >= 36:
                    # Revert to left-aligned with no indent
                    next_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    next_para.paragraph_format.first_line_indent = Pt(0)
                    stats['para_reverts'] += 1

def process_document(docx_path):
    """Process a document to undo formatting."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'para_reverts': 0}

    stats = {'para_reverts': 0}

    content_start_idx = get_content_start_idx(doc)
    undo_paragraph_formatting(doc, content_start_idx, stats)

    elapsed = time.time() - start_time
    reverts = stats['para_reverts']

    if reverts > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_reverted{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Reverted {reverts} paragraphs ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'para_reverts': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'para_reverts': reverts}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Undoing tagline paragraph formatting...")
print()

docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-') and f.endswith('.docx')
    and '_reverted' not in f and not f.startswith('~$')
]

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_reverts = 0

for path in docx_files:
    result = process_document(path)
    if result['para_reverts'] > 0:
        total_updated += 1
        total_reverts += result['para_reverts']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Paragraphs reverted: {total_reverts}")
print(f"\nReverted files: {docs_dir}/*_reverted.docx")
