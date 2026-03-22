"""
Format ALL content paragraphs (Normal style) to be justified with first-line indent.
Skip:
- Edition info page (Ver. 20260205)
- Headings/taglines
- Already formatted paragraphs
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time

def should_skip_paragraph(para):
    """Check if paragraph should be skipped."""
    text = para.text.strip()

    # Skip empty paragraphs
    if not text:
        return True

    # Skip edition info
    if text.startswith('Ver. 202'):
        return True

    # Skip headings (centered, chapter numbers, TOC entries)
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True

    if para.style:
        style_name = para.style.name.lower()
        # Skip special styles
        if any(x in style_name for x in ['heading', 'toc', 'title', 'chapter', 'tag', 'data']):
            return True

    # Skip very short lines (likely headings)
    if len(text) < 20:
        return True

    # Skip taglines (end with ellipsis)
    if text.endswith('...') or text.endswith('…'):
        return True

    return False

def find_about_author_start(doc):
    """Find where About the Author section starts."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if 'ABOUT' in text.upper() and 'AUTHOR' in text.upper():
            return i + 1  # Start after the heading
    return 0

def format_content_paragraphs(doc, stats):
    """Format all content paragraphs to be justified with indent."""
    paragraphs = doc.paragraphs

    # Find About the Author start
    about_author_idx = find_about_author_start(doc)

    if stats['verbose']:
        print(f"  About the Author starts at: {about_author_idx}")

    for i, para in enumerate(paragraphs):
        # Start from About the Author section
        if i < about_author_idx:
            continue

        # Skip paragraphs that shouldn't be formatted
        if should_skip_paragraph(para):
            continue

        # Check if Normal style
        if not para.style or para.style.name != 'Normal':
            continue

        # Check if already formatted correctly
        is_justified = para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        indent = para.paragraph_format.first_line_indent
        is_indented = indent and indent.pt >= 35

        if is_justified and is_indented:
            continue  # Already formatted

        # Format the paragraph
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Pt(36)
        stats['formatted'] += 1

        if stats['verbose'] and stats['formatted'] <= 5:
            print(f"    Formatted para {i}: {para.text[:60]}...")

def process_document(docx_path, verbose=False):
    """Process a document."""
    filename = os.path.basename(docx_path)
    print(f"\nProcessing: {filename}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'formatted': 0}

    stats = {
        'formatted': 0,
        'verbose': verbose
    }

    format_content_paragraphs(doc, stats)

    elapsed = time.time() - start_time
    formatted = stats['formatted']

    if formatted > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_formatted{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Formatted {formatted} paragraphs ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'formatted': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'formatted': formatted}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Formatting all content paragraphs (justified with first-line indent)...")
print()

docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-') and f.endswith('.docx')
    and '_formatted' not in f and not f.startswith('~$')
]

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_formatted = 0

for path in docx_files:
    result = process_document(path, verbose=False)
    if result['formatted'] > 0:
        total_updated += 1
        total_formatted += result['formatted']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Paragraphs formatted: {total_formatted}")
print(f"\nUpdated files: {docs_dir}/*_formatted.docx")
