"""
Fix YY Word documents:
1. Change any 'o' character in Yada Towrah font to Times New Roman (split runs)
2. Change 0.5" first-line indent to 0.3"
"""
import os, glob, copy
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from lxml import etree

DOCS_DIR = r'C:\Users\Joe\Work\dev\yada\docs'
YT_FONT = 'yada towrah'
TNR_FONT = 'Times New Roman'
HALF_INCH = Inches(0.5)
POINT_THREE = Inches(0.3)
TOLERANCE = Inches(0.02)


def get_run_font(r_elem):
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi')) or ''
    return ''


def get_run_text(r_elem):
    parts = []
    for t in r_elem.findall(qn('w:t')):
        parts.append(t.text or '')
    return ''.join(parts)


def set_run_text(r_elem, new_text):
    t_elems = r_elem.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = new_text
        if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
            t_elems[0].set(qn('xml:space'), 'preserve')
        for extra in t_elems[1:]:
            r_elem.remove(extra)
    elif new_text:
        t = etree.SubElement(r_elem, qn('w:t'))
        t.text = new_text
        if new_text[0] == ' ' or new_text[-1] == ' ':
            t.set(qn('xml:space'), 'preserve')


def set_run_font_name(r_elem, font_name):
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(r_elem, qn('w:rPr'))
        r_elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def get_all_runs(para):
    runs = []
    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            runs.append(child)
        elif tag == 'hyperlink':
            for r in child.findall(qn('w:r')):
                runs.append(r)
    return runs


def fix_o_in_yt_runs(para, stats):
    """Split runs to change 'o' chars in Yada Towrah font to Times New Roman."""
    runs = get_all_runs(para)
    for r_elem in runs:
        font = get_run_font(r_elem)
        if font.lower() != YT_FONT:
            continue
        text = get_run_text(r_elem)
        if 'o' not in text:
            continue

        # Split this run into segments: YT chars vs 'o' chars (TNR)
        segments = []  # list of (text, is_o)
        current = ''
        current_is_o = (text[0] == 'o') if text else False
        for ch in text:
            ch_is_o = (ch == 'o')
            if ch_is_o != current_is_o:
                if current:
                    segments.append((current, current_is_o))
                current = ch
                current_is_o = ch_is_o
            else:
                current += ch
        if current:
            segments.append((current, current_is_o))

        if len(segments) <= 1 and not any(is_o for _, is_o in segments):
            continue

        # If entire run is 'o' chars, just change font
        if len(segments) == 1 and segments[0][1]:
            set_run_font_name(r_elem, TNR_FONT)
            stats['o_font'] += len(segments[0][0])
            continue

        # Split into multiple runs
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)

        # First segment goes into original run
        set_run_text(r_elem, segments[0][0])
        if segments[0][1]:
            set_run_font_name(r_elem, TNR_FONT)
            stats['o_font'] += len(segments[0][0])

        # Remaining segments become new runs
        insert_after = idx
        for seg_text, is_o in segments[1:]:
            new_run = copy.deepcopy(r_elem)
            set_run_text(new_run, seg_text)
            if is_o:
                set_run_font_name(new_run, TNR_FONT)
                stats['o_font'] += len(seg_text)
            else:
                set_run_font_name(new_run, font)  # restore original YT font
            insert_after += 1
            parent.insert(insert_after, new_run)


def fix_indents(para, stats):
    """Change 0.5 inch first-line indent to 0.3 inch."""
    pf = para.paragraph_format
    if pf.first_line_indent and abs(pf.first_line_indent - HALF_INCH) < TOLERANCE:
        pf.first_line_indent = POINT_THREE
        stats['indent'] += 1


def process_doc(fpath):
    doc = Document(fpath)
    stats = {'o_font': 0, 'indent': 0}

    # Process body paragraphs
    for para in doc.paragraphs:
        fix_o_in_yt_runs(para, stats)
        fix_indents(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_o_in_yt_runs(para, stats)
                    fix_indents(para, stats)

    # Process headers/footers
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr:
                try:
                    linked = hdr.is_linked_to_previous
                except:
                    continue
                for para in hdr.paragraphs:
                    fix_o_in_yt_runs(para, stats)
        for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
            if ftr:
                try:
                    linked = ftr.is_linked_to_previous
                except:
                    continue
                for para in ftr.paragraphs:
                    fix_o_in_yt_runs(para, stats)

    doc.save(fpath)
    return stats


def main():
    files = sorted(glob.glob(os.path.join(DOCS_DIR, 'YY-*.docx')))
    files = [f for f in files if '_backup' not in f and '_updated' not in f]
    print(f'Found {len(files)} documents to process')

    totals = {'o_font': 0, 'indent': 0}
    for i, fpath in enumerate(files):
        fname = os.path.basename(fpath)
        print(f'  [{i+1}/{len(files)}] {fname}', end='', flush=True)
        stats = process_doc(fpath)
        changes = {k: v for k, v in stats.items() if v > 0}
        if changes:
            print(f'  {changes}')
        else:
            print('  (no changes)')
        for k in totals:
            totals[k] += stats[k]

    print()
    print('=' * 60)
    print('SUMMARY')
    print('=' * 60)
    print(f"  'o' chars: YT -> TNR:      {totals['o_font']}")
    print(f"  0.5\" -> 0.3\" indent:       {totals['indent']}")
    print(f"  Total changes:             {sum(totals.values())}")


if __name__ == '__main__':
    main()
