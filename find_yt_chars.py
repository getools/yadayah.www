"""
Scan YY Word documents for unusual characters in "Yada Towrah" font runs.
Goal: find a character that should be replaced with "8".

Findings:
- U+F067 (PUA character, decimal 61543, XML &#61543;) appears in Yada Towrah font
- It is a Private Use Area encoding of position 0x67 (= ASCII 'g')
- In the Yada Towrah font, 'g' renders as the ancient Hebrew letter "Ghah"
- The document describes Ghah as looking like "the numeral 8 with an upside-down u beneath it"
- This character appears once, in YY-s01v01 paragraph 1413
- The existing replace_unicode_chars() functions in parse_word_translations.py and
  fix_null_pages.py already handle \uf065->'e', \uf066->'f', \uf069->'i'
  but are MISSING the \uf067 mapping
"""

import os
import sys
import glob
import unicodedata
from collections import Counter, defaultdict
from docx import Document
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

DOCS_DIR = r"C:\Users\Joe\Work\dev\yada\docs"
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

all_files = glob.glob(os.path.join(DOCS_DIR, "YY-*.docx"))
doc_files = [
    f for f in all_files
    if "_backup" not in f and "_updated" not in f and ".XXX." not in f
]
doc_files.sort()

# Characters considered "normal" in Yada Towrah font
NORMAL_CHARS = set()
for c in range(ord('a'), ord('z') + 1):
    NORMAL_CHARS.add(chr(c))
for c in range(ord('A'), ord('Z') + 1):
    NORMAL_CHARS.add(chr(c))
for c in range(ord('0'), ord('9') + 1):
    NORMAL_CHARS.add(chr(c))
for ch in " .,;:!?'-\"()[]{}/@#$%^&*+=<>~`|\\\n\t\r":
    NORMAL_CHARS.add(ch)
for ch in '\u02BE\u02BF\u2018\u2019\u201C\u201D\u2013\u2014\u2026':
    NORMAL_CHARS.add(ch)

# Track findings
char_counter = Counter()
char_examples = defaultdict(list)

print(f"Processing {len(doc_files)} documents...")

for i, fpath in enumerate(doc_files):
    fname = os.path.basename(fpath)
    print(f"  [{i+1}/{len(doc_files)}] {fname}")

    try:
        doc = Document(fpath)
    except Exception as e:
        print(f"    ERROR: {e}")
        continue

    for pi, para in enumerate(doc.paragraphs):
        for ri, run in enumerate(para.runs):
            font_name = None
            if run.font and run.font.name:
                font_name = run.font.name
            if not font_name:
                rpr = run._element.find('.//w:rFonts', NSMAP)
                if rpr is not None:
                    font_name = (
                        rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii') or
                        rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi') or
                        rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs')
                    )

            if not font_name or "yada towrah" not in font_name.lower():
                continue

            text = run.text
            if not text:
                continue

            for j, ch in enumerate(text):
                if ch not in NORMAL_CHARS:
                    cp = ord(ch)
                    name = unicodedata.name(ch, "UNKNOWN")
                    key = (ch, f"U+{cp:04X}", name)
                    char_counter[key] += 1

                    if len(char_examples[key]) < 5:
                        # Get paragraph context
                        para_text = para.text
                        idx = para_text.find(ch)
                        if idx >= 0:
                            start = max(0, idx - 50)
                            end = min(len(para_text), idx + 51)
                            ctx = para_text[start:end]
                        else:
                            ctx = para_text[:100]
                        char_examples[key].append({
                            'file': fname,
                            'para': pi,
                            'font': font_name,
                            'context': ctx,
                        })

print("\n" + "=" * 80)
print("UNUSUAL CHARACTERS IN 'Yada Towrah' FONT RUNS")
print("=" * 80)

print(f"\nFound {len(char_counter)} unusual character types:\n")
for (ch, cp, name), count in char_counter.most_common():
    is_pua = 0xE000 <= ord(ch) <= 0xF8FF
    pua_note = ""
    if is_pua:
        ascii_equiv = chr(ord(ch) - 0xF000) if 0xF000 <= ord(ch) <= 0xF0FF else "?"
        pua_note = f" [PUA: maps to ASCII '{ascii_equiv}' (0x{ord(ch)-0xF000:02X})]"
    print(f"  {cp} {name} - {count} occurrence(s){pua_note}")
    for info in char_examples[(ch, cp, name)][:3]:
        print(f"    File: {info['file']}, Para: {info['para']}, Font: {info['font']}")
        print(f"    Context: ...{info['context']}...")
    print()

# Complete character inventory
print("=" * 80)
print("COMPLETE CHARACTER INVENTORY (Yada Towrah font)")
print("=" * 80)

all_chars = Counter()
for fpath in doc_files:
    try:
        doc = Document(fpath)
    except:
        continue
    for para in doc.paragraphs:
        for run in para.runs:
            font_name = None
            if run.font and run.font.name:
                font_name = run.font.name
            if not font_name:
                rpr = run._element.find('.//w:rFonts', NSMAP)
                if rpr is not None:
                    font_name = (
                        rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii') or
                        rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi') or
                        rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs')
                    )
            if not font_name or "yada towrah" not in font_name.lower():
                continue
            if run.text:
                for ch in run.text:
                    all_chars[ch] += 1

print(f"\nTotal unique characters: {len(all_chars)}\n")
for ch, count in all_chars.most_common():
    cp = ord(ch)
    name = unicodedata.name(ch, "UNKNOWN")
    marker = ""
    if ch not in NORMAL_CHARS:
        marker = " *** UNUSUAL ***"
    if 0xE000 <= cp <= 0xF8FF:
        marker = " *** PUA ***"
    print(f"  U+{cp:04X} {name:45s} {count:>7d}{marker}")

print("\n" + "=" * 80)
print("RECOMMENDATION")
print("=" * 80)
print("""
The character U+F067 (PUA, decimal 61543) should be replaced with '8'.

- It appears once in YY-s01v01 (An Intro to God-Dabarym-Words), paragraph 1413
- It is in a Yada Towrah font run
- In the XML it is stored as &#61543; which equals U+F067
- U+F067 is a Private Use Area character that maps to position 0x67 in the font
- 0x67 = ASCII 'g', so it renders the same glyph as lowercase 'g' in Yada Towrah
- The surrounding text describes the letter "Ghah" as looking like "the numeral 8"
- When extracted as plain text, this PUA character is invisible/unprintable

The existing replace_unicode_chars() function in parse_word_translations.py
and fix_null_pages.py should add:
    text = text.replace('\\uf067', '8')
alongside the existing \\uf065->'e', \\uf066->'f', \\uf069->'i' mappings.
""")
