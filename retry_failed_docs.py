#!/usr/bin/env python3
"""
Retry COM page extraction for documents that fell back to XML.
Uses a longer timeout and retries to get actual footer page numbers.
"""

import sys
import os
import logging
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from parse_word_translations import (
    extract_translations_from_doc,
    get_page_numbers_for_doc_com,
    setup_logging,
    DEFAULT_DIRECTORY,
)
from docx import Document

try:
    import psycopg2
except ImportError:
    print("ERROR: psycopg2 not installed")
    sys.exit(1)

try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: python-dotenv not installed")
    sys.exit(1)


# Documents that failed COM and used XML fallback
FAILED_DOCS = [
    "YY-s01v01-An Intro to God-Dabarym-Words",
    "YY-s02v05-Yada Yahowah-Qatsyr-Harvests",
]


def main():
    setup_logging(verbose=False)
    load_dotenv()

    logging.info("Retry COM Page Extraction for Failed Documents")
    logging.info("=" * 60)
    logging.info("Using Information(1) for footer page numbers via COM")
    logging.info("Extended timeout: 300s per document")

    conn = psycopg2.connect(
        host=os.getenv('POSTGRES_HOST', 'localhost'),
        port=os.getenv('POSTGRES_PORT', '5432'),
        user=os.getenv('POSTGRES_USER', 'postgres'),
        password=os.getenv('POSTGRES_PASSWORD'),
        database=os.getenv('POSTGRES_DB', 'yada'),
    )
    cur = conn.cursor()

    doc_dir = Path(DEFAULT_DIRECTORY)

    for book_stem in FAILED_DOCS:
        doc_path = doc_dir / f"{book_stem}.docx"
        if not doc_path.exists():
            logging.warning(f"File not found: {doc_path}")
            continue

        file_size_mb = doc_path.stat().st_size / (1024 * 1024)
        logging.info(f"\nProcessing {doc_path.name} ({file_size_mb:.1f}MB)")

        # Extract translations to get paragraph indices
        translations = extract_translations_from_doc(doc_path)
        if not translations:
            logging.warning(f"  No translations found")
            continue

        logging.info(f"  Found {len(translations)} translations")

        # Get paragraph text snippets for COM text-based Find
        doc = Document(str(doc_path))
        para_texts = {}
        for t in translations:
            idx = t["_para_idx"]
            if idx < len(doc.paragraphs):
                raw_text = doc.paragraphs[idx].text.strip()
                snippet = raw_text[:150] if len(raw_text) > 150 else raw_text
                if snippet:
                    para_texts[idx] = snippet

        # Try COM with extended timeout (300s)
        logging.info(f"  Attempting COM with 300s timeout...")
        page_map = get_page_numbers_for_doc_com(str(doc_path.absolute()), para_texts, timeout=300)

        if page_map:
            logging.info(f"  COM succeeded: {len(page_map)}/{len(translations)} page numbers")

            # Update database
            updated = 0
            for t in translations:
                para_idx = t["_para_idx"]
                new_page = page_map.get(para_idx)
                if new_page is None:
                    continue

                cur.execute("""
                    UPDATE translation
                    SET translation_page = %s
                    WHERE translation_book = %s
                      AND translation_text_word = %s
                      AND (translation_page IS DISTINCT FROM %s)
                """, (new_page, book_stem, t["text_word"], new_page))

                if cur.rowcount > 0:
                    updated += cur.rowcount

            conn.commit()
            logging.info(f"  Updated {updated} rows in database")
        else:
            logging.warning(f"  COM failed again for {doc_path.name}")
            logging.warning(f"  This document will keep XML-estimated page numbers")

    # Verify the specific translation the user asked about
    cur.execute("""
        SELECT translation_id, translation_page,
               LEFT(translation_text_word, 100) as text_preview
        FROM translation
        WHERE translation_id = 168
    """)
    row = cur.fetchone()
    if row:
        logging.info(f"\nVerification - translation_id 168:")
        logging.info(f"  Page: {row[1]}")
        logging.info(f"  Text: {row[2]}")

    cur.close()
    conn.close()
    logging.info("\nDone.")


if __name__ == "__main__":
    main()
