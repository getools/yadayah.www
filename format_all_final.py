import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

docs_dir = r'C:\users\joe\work\dev\yada\docs'

# Get only original YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
files = []
for f in all_files:
    name = os.path.basename(f)
    # Skip intermediate files
    if any(x in name for x in ['_fmt', '_formatted', '_fixed', '_test', '_tagline', '_reverted']):
        continue
    files.append(f)

print(f"Processing {len(files)} documents...")

total = 0
for docx_path in sorted(files):
    name = os.path.basename(docx_path)
    print(f"\\n{name}")

    doc = Document(docx_path)

    # Find About the Author
    start_idx = 0
    for i, para in enumerate(doc.paragraphs):
        if 'AUTHOR' in para.text.upper():
            start_idx = i + 1
            break

    formatted = 0
    for i in range(start_idx, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # Skip conditions
        if not text or text.startswith('Ver.') or len(text) < 20:
            continue
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue
        if text.endswith('...') or text.endswith('…'):
            continue
        if not para.style or para.style.name != 'Normal':
            continue

        # Skip if already formatted
        indent = para.paragraph_format.first_line_indent
        if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and indent and indent.pt >= 35:
            continue

        # Format it
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Pt(36)
        formatted += 1

    if formatted > 0:
        base, ext = os.path.splitext(docx_path)
        out = f"{base}_final{ext}"
        doc.save(out)
        print(f"  Formatted {formatted} paragraphs")
        total += formatted
    else:
        print(f"  No changes needed")

print(f"\\nTotal: {total} paragraphs formatted")
