"""
Apply Yada Towrah font ONLY to ʿ (Ayin) modifier characters.
Split runs so modifiers are isolated with their own font.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

AYIN = chr(0x02BF)  # ʿ

def split_run_for_modifiers(run):
    """Split a run so that ʿ modifiers are in separate runs with Yada Towrah font."""
    if not run.text or AYIN not in run.text:
        return False

    # Get run properties
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name or "Times New Roman"
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    # Get parent paragraph and run index
    para = run._element.getparent()
    run_idx = list(para).index(run._element)
    text = run.text
    run.text = ""

    def create_run(content, use_yada_font=False):
        """Create a new run with specified font."""
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Set font
        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else font_name
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        # Add text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        new_run.append(t)

        return new_run

    # Split text and create runs
    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_modifier=False):
        nonlocal current_run_idx
        if not buffer:
            return
        content = ''.join(buffer)
        new_run = create_run(content, use_yada_font=is_modifier)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    # Process each character
    for char in text:
        if char == AYIN:
            # Flush regular text buffer
            flush_buffer(is_modifier=False)
            # Add modifier to buffer and flush immediately
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)

    # Flush remaining regular text
    flush_buffer(is_modifier=False)

    return True

def process_paragraph_runs(para):
    """Process all runs in a paragraph."""
    fixes = 0
    for run in list(para.runs):
        if split_run_for_modifiers(run):
            fixes += 1
    return fixes

def process_document(docx_path):
    """Process a document to fix modifier fonts."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph_runs(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph_runs(para)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph_runs(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph_runs(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph_runs(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph_runs(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} runs ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing modifier fonts (ONLY ʿ characters use Yada Towrah)...")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total runs fixed: {total_fixes}")
