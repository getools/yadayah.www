#!/usr/bin/env python3
"""
Fix s01v01 page numbers by re-saving via python-docx and trying COM on the clean copy.
"""

import sys
import os
import logging
import tempfile
import subprocess
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from parse_word_translations import (
    extract_translations_from_doc,
    get_page_numbers_for_doc_com,
    setup_logging,
    DEFAULT_DIRECTORY,
)
from docx import Document

try:
    import psycopg2
except ImportError:
    print("ERROR: psycopg2 not installed")
    sys.exit(1)

try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: python-dotenv not installed")
    sys.exit(1)


def main():
    setup_logging(verbose=True)
    load_dotenv()

    doc_path = Path(DEFAULT_DIRECTORY) / "YY-s01v01-An Intro to God-Dabarym-Words.docx"
    book_name = doc_path.stem

    logging.info(f"Fixing page numbers for {doc_path.name}")
    logging.info("=" * 60)

    # Step 1: Re-save via python-docx to create clean copy
    logging.info("Step 1: Creating clean copy via python-docx...")
    clean_path = Path(tempfile.mkdtemp()) / "s01v01_clean.docx"
    doc = Document(str(doc_path))
    doc.save(str(clean_path))

    orig_size = doc_path.stat().st_size / (1024 * 1024)
    clean_size = clean_path.stat().st_size / (1024 * 1024)
    logging.info(f"  Original: {orig_size:.1f}MB, Clean copy: {clean_size:.1f}MB")

    # Step 2: Extract translations from ORIGINAL (to get correct para indices and text)
    logging.info("Step 2: Extracting translations from original...")
    translations = extract_translations_from_doc(doc_path)
    logging.info(f"  Found {len(translations)} translations")

    # Step 3: Get paragraph text for text-based Find
    doc = Document(str(doc_path))
    para_texts = {}
    for t in translations:
        idx = t["_para_idx"]
        if idx < len(doc.paragraphs):
            raw_text = doc.paragraphs[idx].text.strip()
            snippet = raw_text[:150] if len(raw_text) > 150 else raw_text
            if snippet:
                para_texts[idx] = snippet

    # Step 4: Try COM on the clean copy with extended timeout
    logging.info("Step 3: Attempting COM on clean copy (300s timeout)...")

    # Kill any lingering Word
    subprocess.run(["powershell", "-Command",
        "Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue"],
        capture_output=True, timeout=10)

    page_map = get_page_numbers_for_doc_com(str(clean_path.absolute()), para_texts, timeout=300)

    if page_map:
        logging.info(f"  COM succeeded on clean copy: {len(page_map)}/{len(translations)} page numbers")

        # Verify the 'wa hayah' translation
        for t in translations:
            if 'wa hayah' in t.get('text_word', '') and 'achar' in t.get('text_word', ''):
                pidx = t["_para_idx"]
                page = page_map.get(pidx)
                logging.info(f"  'wa hayah' translation: para_idx={pidx}, page={page}")
                break

        # Update database
        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST', 'localhost'),
            port=os.getenv('POSTGRES_PORT', '5432'),
            user=os.getenv('POSTGRES_USER', 'postgres'),
            password=os.getenv('POSTGRES_PASSWORD'),
            database=os.getenv('POSTGRES_DB', 'yada'),
        )
        cur = conn.cursor()
        updated = 0
        for t in translations:
            para_idx = t["_para_idx"]
            new_page = page_map.get(para_idx)
            if new_page is None:
                continue
            cur.execute("""
                UPDATE translation
                SET translation_page = %s
                WHERE translation_book = %s
                  AND translation_text_word = %s
                  AND (translation_page IS DISTINCT FROM %s)
            """, (new_page, book_name, t["text_word"], new_page))
            if cur.rowcount > 0:
                updated += cur.rowcount
        conn.commit()
        logging.info(f"  Updated {updated} rows in database")

        # Verify
        cur.execute("SELECT translation_id, translation_page FROM translation WHERE translation_id = 168")
        row = cur.fetchone()
        if row:
            logging.info(f"  Verification: translation_id 168, page={row[1]}")

        cur.close()
        conn.close()
    else:
        logging.warning("  COM failed on clean copy too")
        logging.info("  Setting s01v01 page numbers to NULL since footer pages cannot be read")

        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST', 'localhost'),
            port=os.getenv('POSTGRES_PORT', '5432'),
            user=os.getenv('POSTGRES_USER', 'postgres'),
            password=os.getenv('POSTGRES_PASSWORD'),
            database=os.getenv('POSTGRES_DB', 'yada'),
        )
        cur = conn.cursor()
        cur.execute("""
            UPDATE translation SET translation_page = NULL
            WHERE translation_book = %s
        """, (book_name,))
        logging.info(f"  Set {cur.rowcount} rows to NULL page")
        conn.commit()
        cur.close()
        conn.close()

    # Cleanup
    try:
        os.unlink(str(clean_path))
        os.rmdir(str(clean_path.parent))
    except:
        pass

    logging.info("\nDone.")


if __name__ == "__main__":
    main()
