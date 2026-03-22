"""
Fix YY Word documents: ensure headers, yy_header, and yy_break paragraphs have:
- First Line Indent = 0
- Hanging Indent = 0 (no negative first_line_indent)
- Left Indent = 0
- Alignment = CENTER
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os, glob

DOCS_DIR = r'C:\Users\Joe\Work\dev\yada\docs'


def fix_paragraph(para, stats, label):
    """Ensure paragraph has center alignment and zero indents."""
    pf = para.paragraph_format
    changed = False

    if pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        changed = True

    if pf.first_line_indent is None or pf.first_line_indent != 0:
        pf.first_line_indent = Pt(0)
        changed = True

    if pf.left_indent is None or pf.left_indent != 0:
        pf.left_indent = Pt(0)
        changed = True

    if changed:
        stats[label] = stats.get(label, 0) + 1


def process_doc(fpath):
    doc = Document(fpath)
    stats = {}

    # Fix yy_break style definition
    try:
        style = doc.styles['yy_break']
        if hasattr(style, 'paragraph_format'):
            pf = style.paragraph_format
            if pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if pf.first_line_indent is None or pf.first_line_indent != 0:
                pf.first_line_indent = Pt(0)
            if pf.left_indent is None or pf.left_indent != 0:
                pf.left_indent = Pt(0)
    except KeyError:
        pass

    # Fix body paragraphs with yy_break style
    for para in doc.paragraphs:
        sn = para.style.name if para.style else ''
        if sn == 'yy_break':
            fix_paragraph(para, stats, 'yy_break')

    # Fix section headers
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr:
                try:
                    linked = hdr.is_linked_to_previous
                except:
                    continue
                for para in hdr.paragraphs:
                    fix_paragraph(para, stats, 'header')

    doc.save(fpath)
    return stats


def main():
    files = sorted(glob.glob(os.path.join(DOCS_DIR, 'YY-*.docx')))
    files = [f for f in files if '_backup' not in f and '_updated' not in f]
    print(f'Found {len(files)} documents to process')

    totals = {}
    for i, fpath in enumerate(files):
        fname = os.path.basename(fpath)
        print(f'  [{i+1}/{len(files)}] {fname}', end='', flush=True)
        stats = process_doc(fpath)
        if stats:
            print(f'  {stats}')
        else:
            print('  (no changes)')
        for k, v in stats.items():
            totals[k] = totals.get(k, 0) + v

    print()
    print('=' * 60)
    print('SUMMARY')
    print('=' * 60)
    for k, v in sorted(totals.items()):
        print(f'  {k}: {v}')
    print(f'  Total: {sum(totals.values())}')


if __name__ == '__main__':
    main()
