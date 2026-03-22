"""Debug TOC paragraph structure to find where text lives."""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s01v03-An Intro to God-Towrah Mizmowr.docx')

# Check para 36 which has Yasha'yah
para = doc.paragraphs[36]
print(f"Para 36 text: {para.text[:100]}")
print(f"Number of runs (para.runs): {len(para.runs)}")

# Check if text is in runs
run_text = ''.join(r.text or '' for r in para.runs)
print(f"Run text: {run_text[:100]}")
print(f"Match: {run_text == para.text}")

# Check XML structure
p_elem = para._element
print(f"\nXML children of paragraph:")
for i, child in enumerate(p_elem):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    print(f"  [{i}] <{tag}>")

    if tag == 'hyperlink':
        # Check runs inside hyperlink
        for j, sub in enumerate(child):
            sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
            if sub_tag == 'r':
                t_elem = sub.find(qn('w:t'))
                text = t_elem.text if t_elem is not None else ''
                print(f"    [{j}] <{sub_tag}> text: {repr(text[:60])}")
    elif tag == 'r':
        t_elem = child.find(qn('w:t'))
        text = t_elem.text if t_elem is not None else ''
        print(f"    text: {repr(text[:60])}")

# Check Para 39 (Mal'aky)
print("\n" + "=" * 80)
para = doc.paragraphs[39]
print(f"Para 39 text: {para.text[:100]}")
print(f"Number of runs (para.runs): {len(para.runs)}")

p_elem = para._element
print(f"\nXML children:")
for i, child in enumerate(p_elem):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'hyperlink':
        for j, sub in enumerate(child):
            sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
            if sub_tag == 'r':
                t_elem = sub.find(qn('w:t'))
                text = t_elem.text if t_elem is not None else ''
                if text:
                    print(f"  hyperlink run [{j}]: {repr(text[:60])}")
