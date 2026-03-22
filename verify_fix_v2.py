import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')

print("Verification Report")
print("=" * 80)

# Check edition info page (should NOT be touched)
print("\n1. Checking edition info paragraphs (should NOT be justified/indented):\n")
for i in range(35, 46):
    para = doc.paragraphs[i]
    if para.text.strip():
        is_justified = para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        indent = para.paragraph_format.first_line_indent
        is_indented = indent and indent.pt >= 36

        status = "✗ INCORRECTLY MODIFIED" if (is_justified and is_indented) else "✓ Correctly left alone"
        print(f"  Para {i}: {status}")
        if i == 35:  # Show the version number
            print(f"    Text: {para.text[:50]}")

# Check the target paragraph (should BE touched)
print("\n2. Checking target paragraph (SHOULD be justified/indented):\n")
target_text = 'The fourth Miqra'
for i, para in enumerate(doc.paragraphs):
    if target_text in para.text:
        print(f"  Paragraph {i}: \"{para.text[:80]}...\"")
        print(f"    Alignment: {para.alignment}")
        indent = para.paragraph_format.first_line_indent
        if indent:
            print(f"    First-line indent: {indent.pt}pt")
        else:
            print(f"    First-line indent: None")

        is_justified = para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        is_indented = indent and indent.pt >= 36

        if is_justified and is_indented:
            print(f"\n  ✓ SUCCESS - Paragraph is correctly justified and indented!")
        else:
            print(f"\n  ✗ PROBLEM - Paragraph was not fixed")
        break

# Check a few more paragraphs after taglines
print("\n3. Checking other paragraphs after taglines:\n")
count = 0
for i in range(63, min(150, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()

    # Look for taglines
    if (text.endswith('...') or text.endswith('…')) and para.style:
        if hasattr(para.style, 'font') and para.style.font.italic:
            # This is a tagline - check next paragraph
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if next_para.text.strip():
                    is_justified = next_para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
                    indent = next_para.paragraph_format.first_line_indent
                    is_indented = indent and indent.pt >= 36

                    status = "✓ Fixed" if (is_justified and is_indented) else "✗ Not fixed"
                    print(f"  After tagline at {i}: {status}")
                    count += 1
                    if count >= 5:
                        break

print("\n" + "=" * 80)
