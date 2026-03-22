"""
Update Hebrew transliterations in YY Word documents (.docx files) to use
proper Unicode modifiers (Aleph ʾ and Ayin ʿ) instead of curly quotes.

Uses python-docx to read/write .docx files without needing Word COM.
"""

import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load refined conversions
conversions_file = r'C:\Users\Joe\Work\dev\yada\translations\hebrew_modifier_conversions_refined.txt'
conversions = {}

print("Loading refined conversions...")
with open(conversions_file, 'r', encoding='utf-8') as f:
    lines = f.readlines()
    for line in lines[2:]:  # Skip header
        if ', ' in line and not line.startswith('==='):
            parts = line.strip().split(', ', 1)
            if len(parts) == 2:
                orig, conv = parts
                conversions[orig] = conv

print(f"Loaded {len(conversions)} conversions\n")

# Build regex pattern for all words to replace
# Sort by length (longest first) to avoid partial replacements
sorted_words = sorted(conversions.keys(), key=len, reverse=True)

# Escape special regex characters
def escape_regex(s):
    return re.escape(s)

# Create pattern that matches whole words only (bounded by word boundaries or italic tags)
pattern_parts = [escape_regex(w) for w in sorted_words[:500]]  # Limit pattern size
REPLACE_PATTERN = re.compile('|'.join(pattern_parts))

def replace_text_in_paragraph(paragraph):
    """Replace Hebrew transliterations in a paragraph, preserving formatting."""
    replacements_made = 0

    # Process each run separately to preserve formatting
    for run in paragraph.runs:
        if not run.text:
            continue

        original_text = run.text
        new_text = original_text

        # Try direct replacements for words in our conversion dict
        for orig, conv in conversions.items():
            if orig in new_text:
                new_text = new_text.replace(orig, conv)
                if new_text != original_text:
                    replacements_made += 1

        if new_text != original_text:
            run.text = new_text

    return replacements_made

def process_document(docx_path):
    """Process a single Word document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: Could not open document: {e}")
        return 0

    total_replacements = 0

    # Process all paragraphs
    for para in doc.paragraphs:
        replacements = replace_text_in_paragraph(para)
        total_replacements += replacements

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replacements = replace_text_in_paragraph(para)
                    total_replacements += replacements

    if total_replacements > 0:
        # Save with _updated suffix
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Saved to: {os.path.basename(output_path)}")
            print(f"  Replacements made: {total_replacements}")
        except Exception as e:
            print(f"  ERROR: Could not save document: {e}")
            return 0
    else:
        print(f"  No replacements needed")

    return total_replacements

# Find all YY Word documents
docs_dir = r'C:\users\joe\work\dev\yada\docs'
if not os.path.exists(docs_dir):
    print(f"ERROR: Documents directory not found: {docs_dir}")
    sys.exit(1)

print(f"\nSearching for .docx files in: {docs_dir}")
docx_files = [os.path.join(docs_dir, f) for f in os.listdir(docs_dir)
              if f.startswith('YY-') and f.endswith('.docx') and not f.endswith('_updated.docx')]

print(f"Found {len(docx_files)} YY documents to process\n")
print("="*80)

total_files_updated = 0
total_replacements_overall = 0

for docx_path in docx_files:
    replacements = process_document(docx_path)
    if replacements > 0:
        total_files_updated += 1
        total_replacements_overall += replacements

print("\n" + "="*80)
print("WORD DOCUMENT UPDATE COMPLETE")
print("="*80)
print(f"  Files processed: {len(docx_files)}")
print(f"  Files updated: {total_files_updated}")
print(f"  Total replacements: {total_replacements_overall}")
print(f"\nUpdated files saved with '_updated.docx' suffix in: {docs_dir}")
