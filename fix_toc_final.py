"""
Fix remaining quotes in TOC entries by processing hyperlink runs.
TOC entries use w:hyperlink elements containing the runs, which
para.runs doesn't include. This script processes ALL w:r elements
(including those inside hyperlinks).

Also applies Yada Towrah font to any new half-ring characters.
"""
import re
import psycopg2
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

import sys
sys.path.insert(0, r'C:\Users\Joe\Work\dev\yada\translations')
from fix_half_rings import (
    normalize_quotes, normalize_blb, strip_nikkud, get_aleph_ayin_sequence,
    yy_to_hebrew_patterns, count_quotes, HALF_RING_ALEPH, HALF_RING_AYIN,
    QUOTE_CHARS, ALEPH, AYIN
)

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')
YT_FONT = 'Yada Towrah'
ALL_QUOTE_LIKE = QUOTE_CHARS | {HALF_RING_ALEPH, HALF_RING_AYIN}

ENGLISH_SKIP = {
    "allah's", "don't", "dowd's", "fool's", "god's", "heaven's",
    "islam's", "i'm", "lord's", "paul's", "satan's", "who's",
    "yahowah's", "yah's", "'bout", "noah",
}


def is_english(word):
    w = word.lower()
    for q in QUOTE_CHARS:
        w = w.replace(q, "'")
    w = w.replace(HALF_RING_ALEPH, "'").replace(HALF_RING_AYIN, "'")
    return w in ENGLISH_SKIP


def is_quote(ch):
    return ch in ALL_QUOTE_LIKE


def get_all_runs_from_para(para):
    """Get ALL w:r elements from paragraph, including inside hyperlinks."""
    runs = []
    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            runs.append(child)
        elif tag == 'hyperlink':
            for r in child.findall(qn('w:r')):
                runs.append(r)
    return runs


def get_run_text(run_elem):
    """Get text from a w:r element."""
    text = ''
    for t in run_elem.findall(qn('w:t')):
        if t.text:
            text += t.text
    return text


def set_run_text(run_elem, new_text):
    """Set text on a w:r element, preserving the first w:t element."""
    t_elems = run_elem.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = new_text
        t_elems[0].set(qn('xml:space'), 'preserve')
        # Remove extra w:t elements
        for t in t_elems[1:]:
            run_elem.remove(t)
    else:
        t = etree.SubElement(run_elem, qn('w:t'))
        t.text = new_text
        t.set(qn('xml:space'), 'preserve')


def process_paragraph_all_runs(para, patterns):
    """Process ALL runs in a paragraph (including hyperlink runs)."""
    all_runs = get_all_runs_from_para(para)
    if not all_runs:
        return 0

    # Build text and char-to-run mapping
    text = ''
    cmap = []  # (run_index, char_index_in_run_text)
    run_texts = []
    for ri, run in enumerate(all_runs):
        rtext = get_run_text(run)
        run_texts.append(rtext)
        for ci, ch in enumerate(rtext):
            text += ch
            cmap.append((ri, ci))

    if not text or not any(is_quote(ch) for ch in text):
        return 0

    count = 0
    used = set()

    for pat, char_reps in patterns:
        for m in pat.finditer(text):
            s, e = m.start(), m.end()
            if any(i in used for i in range(s, e)):
                continue
            for offset, new_ch in char_reps.items():
                pos = s + offset
                if pos < len(cmap):
                    ri, ci = cmap[pos]
                    t = list(run_texts[ri])
                    t[ci] = new_ch
                    run_texts[ri] = ''.join(t)
                    set_run_text(all_runs[ri], run_texts[ri])
            for i in range(s, e):
                used.add(i)
            count += 1

    return count


def apply_yt_font_all_runs(para):
    """Set Yada Towrah font on half-ring chars in ALL runs (including hyperlinks)."""
    all_runs = get_all_runs_from_para(para)
    if not all_runs:
        return 0

    full_text = ''.join(get_run_text(r) for r in all_runs)
    if HALF_RING_ALEPH not in full_text and HALF_RING_AYIN not in full_text:
        return 0

    hr_count = 0
    for run_elem in all_runs:
        text = get_run_text(run_elem)
        if not text:
            continue

        has_hr = any(ch in (HALF_RING_ALEPH, HALF_RING_AYIN) for ch in text)
        if not has_hr:
            continue

        all_hr = all(ch in (HALF_RING_ALEPH, HALF_RING_AYIN) for ch in text)
        if all_hr:
            # Set font on this run directly
            rPr = run_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(run_elem, qn('w:rPr'))
                run_elem.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:ascii'), YT_FONT)
            rFonts.set(qn('w:hAnsi'), YT_FONT)
            rFonts.set(qn('w:cs'), YT_FONT)
            hr_count += len(text)
            continue

        # Mixed content: split into segments
        segments = []
        current_text = ''
        current_is_hr = text[0] in (HALF_RING_ALEPH, HALF_RING_AYIN)
        for ch in text:
            ch_is_hr = ch in (HALF_RING_ALEPH, HALF_RING_AYIN)
            if ch_is_hr == current_is_hr:
                current_text += ch
            else:
                segments.append((current_text, current_is_hr))
                current_text = ch
                current_is_hr = ch_is_hr
        segments.append((current_text, current_is_hr))

        if len(segments) <= 1:
            continue

        # First segment stays on original run
        set_run_text(run_elem, segments[0][0])
        if segments[0][1]:
            rPr = run_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(run_elem, qn('w:rPr'))
                run_elem.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:ascii'), YT_FONT)
            rFonts.set(qn('w:hAnsi'), YT_FONT)
            rFonts.set(qn('w:cs'), YT_FONT)
            hr_count += len(segments[0][0])

        # Remaining segments: clone run and insert after
        prev_elem = run_elem
        for seg_text, seg_is_hr in segments[1:]:
            new_run = deepcopy(run_elem)
            set_run_text(new_run, seg_text)

            if seg_is_hr:
                rPr = new_run.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(new_run, qn('w:rPr'))
                    new_run.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                rFonts.set(qn('w:ascii'), YT_FONT)
                rFonts.set(qn('w:hAnsi'), YT_FONT)
                rFonts.set(qn('w:cs'), YT_FONT)
                hr_count += len(seg_text)

            prev_elem.addnext(new_run)
            prev_elem = new_run

    return hr_count


def lookup_word(word, spelling_map, blb_map, hebrew_lookup):
    """Look up a word and return the fixed version."""
    if is_english(word):
        return None

    n_quotes = count_quotes(word)
    if n_quotes == 0:
        return None

    normalized = normalize_quotes(word)
    hebrew = None

    # Strategy 1: YY spelling
    if normalized in spelling_map:
        hebrew = spelling_map[normalized]
    if not hebrew and '-' in normalized:
        for part in normalized.split('-'):
            if count_quotes(part) > 0 and part in spelling_map:
                hebrew = spelling_map[part]
                break

    # Strategy 2: BLB
    if not hebrew:
        blb_key = normalize_blb(word)
        if blb_key in blb_map:
            hebrew = blb_map[blb_key]
    if not hebrew and '-' in normalized:
        for part in normalized.split('-'):
            if count_quotes(part) > 0 and part in blb_map:
                hebrew = blb_map[part]
                break

    if hebrew:
        seq = get_aleph_ayin_sequence(hebrew)
        if len(seq) >= n_quotes:
            return replace_quotes(word, seq)
        elif len(seq) > 0:
            return replace_quotes(word, seq)

    # Strategy 3: Hebrew pattern
    patterns = yy_to_hebrew_patterns(normalized)
    for heb_pattern, _ in patterns:
        if heb_pattern in hebrew_lookup:
            db_seq = hebrew_lookup[heb_pattern]
            if len(db_seq) >= n_quotes:
                return replace_quotes(word, db_seq)

    if '-' in normalized:
        for part in normalized.split('-'):
            if count_quotes(part) > 0:
                part_patterns = yy_to_hebrew_patterns(part)
                for heb_pattern, _ in part_patterns:
                    if heb_pattern in hebrew_lookup:
                        db_seq = hebrew_lookup[heb_pattern]
                        if len(db_seq) >= count_quotes(part):
                            return replace_quotes(word, db_seq)

    # Strategy 4: Heuristic
    replacements = []
    for i, ch in enumerate(normalized):
        if ch == "'":
            if i == 0:
                replacements.append(HALF_RING_ALEPH)
            elif i == len(normalized) - 1:
                replacements.append(HALF_RING_ALEPH)
            else:
                replacements.append(HALF_RING_AYIN)
    return replace_quotes(word, replacements)


def replace_quotes(word, seq):
    result = ''
    idx = 0
    for ch in word:
        if ch in QUOTE_CHARS or ch == HALF_RING_ALEPH or ch == HALF_RING_AYIN:
            if idx < len(seq):
                result += seq[idx]
                idx += 1
            else:
                result += ch
        else:
            result += ch
    return result


def build_pattern(word):
    QUOTE_CLASS = "['\u2018\u2019`\u00B4\u02BC\u02BE\u02BF]"
    parts = []
    for ch in word:
        if is_quote(ch):
            parts.append(QUOTE_CLASS)
        else:
            parts.append(re.escape(ch))
    return re.compile(r'(?<![a-zA-Z])' + ''.join(parts) + r'(?![a-zA-Z])')


def get_char_replacements(original, fixed):
    return {i: f for i, (o, f) in enumerate(zip(original, fixed)) if o != f}


def main():
    # Load DB maps
    conn = psycopg2.connect(
        host='localhost', port=5433, dbname='yada',
        user='postgres', password='yada_password'
    )
    cur = conn.cursor()

    cur.execute('''
        SELECT ws.word_translit_text, w.word_hebrew
        FROM yy_word_translit ws
        JOIN yy_word w ON ws.word_id = w.word_key
        WHERE w.word_hebrew IS NOT NULL AND w.word_hebrew != ''
    ''')
    spelling_map = {}
    for spelling, hebrew in cur.fetchall():
        key = normalize_quotes(spelling.strip())
        if key not in spelling_map:
            spelling_map[key] = hebrew.strip()

    cur.execute('''
        SELECT word_import_translit, word_import_hebrew
        FROM yy_word_import
        WHERE word_import_translit IS NOT NULL AND word_import_hebrew IS NOT NULL
    ''')
    blb_map = {}
    for translit, hebrew in cur.fetchall():
        key = normalize_blb(translit.strip())
        if key not in blb_map:
            blb_map[key] = hebrew.strip()

    hebrew_lookup = {}
    cur.execute('''
        SELECT word_import_hebrew FROM yy_word_import
        WHERE word_import_hebrew IS NOT NULL AND word_import_hebrew != ''
    ''')
    for (heb_raw,) in cur.fetchall():
        heb = strip_nikkud(heb_raw.strip())
        if heb and heb not in hebrew_lookup:
            seq = get_aleph_ayin_sequence(heb_raw.strip())
            if seq:
                hebrew_lookup[heb] = seq

    conn.close()

    # First pass: find all unique words with quotes in TOC (hyperlink) paragraphs
    files = sorted(DOCS_DIR.glob('YY*.docx'))
    files = [f for f in files if '_backup' not in f.name and '_updated' not in f.name
             and '_halfring' not in f.name]

    toc_words = set()
    for fp in files:
        doc = Document(str(fp))
        for para in doc.paragraphs[:80]:
            has_field = bool(para._element.findall('.//' + qn('w:instrText')))
            style_name = para.style.name if para.style else ''
            is_toc = has_field or 'toc' in style_name.lower()
            if not is_toc:
                continue
            # Get text from all runs including hyperlinks
            all_runs = get_all_runs_from_para(para)
            text = ''.join(get_run_text(r) for r in all_runs)
            for w in text.split():
                clean = w.strip('\t0123456789 ')
                if clean and any(ch in QUOTE_CHARS for ch in clean):
                    toc_words.add(clean)

    # Also grab the full replacement map from italic_quote_words_fixed.txt
    fixed_file = Path(r'C:\Users\Joe\Work\dev\yada\translations\italic_quote_words_fixed.txt')
    existing_map = {}
    lines = open(fixed_file, encoding='utf-8').readlines()
    for line in lines[2:]:
        if not line.strip():
            continue
        fixed = line[:40].strip()
        original = line[41:71].strip()
        if original and fixed and original != fixed:
            if original not in existing_map:
                existing_map[original] = fixed

    # Build combined replacement map
    replacement_map = dict(existing_map)  # start with existing

    # Add TOC-specific words
    print(f'TOC words with quotes: {len(toc_words)}')
    for word in sorted(toc_words):
        if is_english(word):
            print(f'  {word} -> [skip: English]')
            continue
        # Check if already in existing map
        if word in replacement_map:
            print(f'  {word} -> {replacement_map[word]} [existing]')
            continue
        fixed = lookup_word(word, spelling_map, blb_map, hebrew_lookup)
        if fixed and fixed != word:
            replacement_map[word] = fixed
            print(f'  {word} -> {fixed} [new]')
        else:
            print(f'  {word} -> [no match]')

    # Build patterns (longest first)
    patterns = []
    for original, fixed in sorted(replacement_map.items(), key=lambda x: -len(x[0])):
        pat = build_pattern(original)
        reps = get_char_replacements(original, fixed)
        if reps:
            patterns.append((pat, reps))
    print(f'\n{len(patterns)} total patterns')

    # Process files - using ALL runs (including hyperlinks)
    print(f'\nProcessing {len(files)} files...\n')
    grand_rep = 0
    grand_font = 0

    for i, fp in enumerate(files):
        print(f'  [{i+1}/{len(files)}] {fp.name}', end='', flush=True)
        try:
            doc = Document(str(fp))
            count = 0

            # Process ALL paragraphs using the hyperlink-aware function
            for para in doc.paragraphs:
                count += process_paragraph_all_runs(para, patterns)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            count += process_paragraph_all_runs(para, patterns)

            for section in doc.sections:
                try:
                    if not section.header.is_linked_to_previous:
                        for para in section.header.paragraphs:
                            count += process_paragraph_all_runs(para, patterns)
                except Exception:
                    pass

            # Apply YT font to half-rings (also hyperlink-aware)
            font_count = 0
            for para in doc.paragraphs:
                font_count += apply_yt_font_all_runs(para)
            for section in doc.sections:
                try:
                    if not section.header.is_linked_to_previous:
                        for para in section.header.paragraphs:
                            font_count += apply_yt_font_all_runs(para)
                except Exception:
                    pass

            doc.save(str(fp))
            print(f' -> {count} replacements, {font_count} font fixes')
            grand_rep += count
            grand_font += font_count
        except Exception as e:
            print(f' ERROR: {e}')

    print(f'\nDone! {grand_rep} replacements, {grand_font} font fixes')

    # Verify
    print('\n--- Verification: remaining quotes in TOC ---')
    remaining = 0
    for fp in files:
        doc = Document(str(fp))
        for i, para in enumerate(doc.paragraphs[:80]):
            has_field = bool(para._element.findall('.//' + qn('w:instrText')))
            style_name = para.style.name if para.style else ''
            is_toc = has_field or 'toc' in style_name.lower()
            if not is_toc:
                continue
            all_runs = get_all_runs_from_para(para)
            text = ''.join(get_run_text(r) for r in all_runs)
            quotes_in_hebrew = []
            words = text.split()
            for w in words:
                clean = w.strip('\t0123456789 ')
                if clean and any(ch in QUOTE_CHARS for ch in clean) and not is_english(clean):
                    quotes_in_hebrew.append(clean)
            if quotes_in_hebrew:
                remaining += 1
                print(f'  {fp.name} para[{i}]: {quotes_in_hebrew}')

    if remaining == 0:
        print('  All Hebrew transliteration quotes in TOC have been replaced!')
    else:
        print(f'\n  {remaining} TOC entries still have Hebrew quotes')


if __name__ == '__main__':
    main()
