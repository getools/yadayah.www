"""
Update all YY Word documents with proper Hebrew modifiers (Aleph ʾ and Ayin ʿ).
Processes all .docx files in the docs directory.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import time

# Load conversions
conversions_file = r'C:\Users\Joe\Work\dev\yada\translations\hebrew_modifier_conversions_refined.txt'
conversions = {}

print("Loading conversions...")
with open(conversions_file, 'r', encoding='utf-8') as f:
    lines = f.readlines()
    for line in lines[2:]:
        if ', ' in line and not line.startswith('===') and not line.startswith('Score'):
            parts = line.strip().split(', ', 1)
            if len(parts) == 2:
                orig, conv = parts
                conversions[orig] = conv

print(f"Loaded {len(conversions)} conversions\n")

# Find all YY documents
docs_dir = r'C:\users\joe\work\dev\yada\docs'
files = [f for f in os.listdir(docs_dir)
         if f.startswith('YY-') and f.endswith('.docx') and not '_updated' in f]

print(f"Found {len(files)} YY documents to process")
print("="*80 + "\n")

total_files_updated = 0
total_replacements = 0

for idx, filename in enumerate(files, 1):
    docx_path = os.path.join(docs_dir, filename)
    print(f"[{idx}/{len(files)}] {filename}")

    start_time = time.time()

    try:
        doc = Document(docx_path)
        replacements = 0

        # Process paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                if not run.text:
                    continue

                original = run.text
                new_text = original

                for orig, conv in conversions.items():
                    if orig in new_text:
                        new_text = new_text.replace(orig, conv)

                if new_text != original:
                    run.text = new_text
                    replacements += 1

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text:
                                continue

                            original = run.text
                            new_text = original

                            for orig, conv in conversions.items():
                                if orig in new_text:
                                    new_text = new_text.replace(orig, conv)

                            if new_text != original:
                                run.text = new_text
                                replacements += 1

        elapsed = time.time() - start_time

        if replacements > 0:
            output_path = docx_path.replace('.docx', '_updated.docx')
            doc.save(output_path)
            print(f"  ✓ {replacements} replacements, saved to {os.path.basename(output_path)} ({elapsed:.1f}s)")
            total_files_updated += 1
            total_replacements += replacements
        else:
            print(f"  - No changes needed ({elapsed:.1f}s)")

    except Exception as e:
        print(f"  ✗ ERROR: {e}")

    print()

print("="*80)
print("COMPLETE")
print("="*80)
print(f"Files processed: {len(files)}")
print(f"Files updated: {total_files_updated}")
print(f"Total replacements: {total_replacements}")
print(f"\nUpdated files saved with '_updated.docx' suffix in: {docs_dir}")
