"""
Replace ALL apostrophe variants in Hebrew transliterations with proper ALEPH (ʾ) and AYIN (ʿ) modifiers.
Handles: ' ' ' ` ´ ʼ and any other apostrophe-like characters.
Then splits modifiers into separate runs with Yada Towrah font.
Processes body text, tables, TOC, headers, and footers.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

ALEPH = chr(0x02BE)  # ʾ
AYIN = chr(0x02BF)   # ʿ

# ALL possible apostrophe variations to handle (use chr() to avoid encoding issues)
APOSTROPHE_VARIANTS = [
    chr(0x0027),  # Regular apostrophe '
    chr(0x2019),  # Right single quote ' (curly apostrophe)
    chr(0x2018),  # Left single quote ' (left curly quote)
    chr(0x0060),  # Backtick `
    chr(0x00B4),  # Acute accent ´
    chr(0x02BC),  # Modifier letter apostrophe ʼ
]

# Base replacements (using ' as placeholder for any apostrophe variant)
# Format: (pattern_with_apostrophe, replacement_with_modifier)
BASE_REPLACEMENTS = [
    # ALEPH replacements (ʾ)
    ("Yisrae'l", "Yisraʾel"),
    ("'Adam", "ʾAdam"),
    ("'Abraham", "ʾAbraham"),
    ("Miqra'", "Miqraʾ"),
    ("'el", "ʾel"),
    ("'Arek", "ʾArek"),
    ("'arek", "ʾarek"),
    ("'Akad", "ʾAkad"),
    ("'akad", "ʾakad"),
    ("'Ashuwr", "ʾAshuwr"),
    ("'ashuwr", "ʾashuwr"),
    ("'elowah", "ʾelowah"),
    ("'elowahym", "ʾelowahym"),
    ("'Abraham's", "ʾAbraham's"),
    ("'Abshalowm's", "ʾAbshalowm's"),
    ("'Ayil", "ʾAyil"),
    ("'Abshalowm", "ʾAbshalowm"),
    ("'Amal", "ʾAmal"),
    ("'amal", "ʾamal"),
    ("'atem", "ʾatem"),
    ("'any", "ʾany"),
    ("'Edowm", "ʾEdowm"),
    ("'Amown", "ʾAmown"),
    ("'atah", "ʾatah"),
    ("'ElYah", "ʾElYah"),
    ("'anachnuw", "ʾanachnuw"),
    ("'Elowym", "ʾElowym"),
    ("'ozen", "ʾozen"),
    ("'Abyb", "ʾAbyb"),
    ("'ishah", "ʾishah"),
    ("'chasydy", "ʾchasydy"),

    # AYIN replacements (ʿ)
    ("Yisra'el", "Yisraʿel"),
    ("Ya'aqob", "Yaʿaqob"),
    ("Mow'ab", "Mowʿab"),
    ("Mow'edym", "Mowʿedym"),
    ("Yahowsha'", "Yahowshaʿ"),
    ("Yow'ab", "Yowʿab"),
    ("Ga'al", "Gaʿal"),
    ("yasha'", "yashaʿ"),
    ("Yisra'elites", "Yisraʿelites"),
    ("Mow'abites", "Mowʿabites"),
    ("Yada'", "Yadaʿ"),
    ("Zarowa'", "Zarowaʿ"),
    ("Yow'el", "Yowʿel"),
    ("Yasha'", "Yashaʿ"),
    ("Howsha'", "Howshaʿ"),
    ("Sha'uwl", "Shaʿuwl"),
    ("Shin'ar", "Shinʿar"),
    ("shin'ar", "shinʿar"),
    ("'arabah", "ʿarabah"),
    ("'owd", "ʿowd"),
    ("'Eden", "ʿEden"),
    ("'aleichem", "ʿaleichem"),
    ("Naby'", "Nabyʿ"),
    ("Ma'achach", "Maʿachach"),
    ("Yahowshuwa'", "Yahowshuwaʿ"),
    ("mari'yth", "mariʿyth"),
    ("mal'akym", "malʿakym"),
    ("yesha'y", "yeshaʿy"),
    ("arba'ym", "arbaʿym"),
    ("sha'b", "shaʿb"),
    ("mits'ar", "mitsʿar"),
]

def generate_all_variants():
    """Generate replacement pairs for each apostrophe variant."""
    all_replacements = []

    for old_pattern, new_pattern in BASE_REPLACEMENTS:
        # Create version for each apostrophe variant
        for apos in APOSTROPHE_VARIANTS:
            variant_old = old_pattern.replace("'", apos)
            if variant_old != new_pattern:  # Skip if already correct
                all_replacements.append((variant_old, new_pattern))

    return all_replacements

ALL_REPLACEMENTS = generate_all_variants()

def split_run_at_modifier(run_elem, text, modifier_char):
    """Split run to isolate modifier with Yada Towrah font."""
    if modifier_char not in text:
        return False

    # Get parent paragraph
    para = run_elem.getparent()
    if para is None:
        return False

    run_idx = list(para).index(run_elem)

    # Get original formatting
    rPr = run_elem.find(qn('w:rPr'))

    def clone_rPr(exclude_font=False):
        new_rPr = OxmlElement('w:rPr')
        if rPr is not None:
            for child in rPr:
                if exclude_font and child.tag == qn('w:rFonts'):
                    continue
                import copy
                new_rPr.append(copy.deepcopy(child))
        return new_rPr

    def create_run(content, is_modifier=False):
        new_run = OxmlElement('w:r')
        new_rPr = clone_rPr(exclude_font=is_modifier)

        if is_modifier:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Yada Towrah')
            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
            new_rPr.append(rFonts)

        new_run.append(new_rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        new_run.append(t)

        return new_run

    # Remove original run
    para.remove(run_elem)

    # Split text and create runs
    current_idx = run_idx
    parts = text.split(modifier_char)

    for i, part in enumerate(parts):
        if i > 0:  # Insert modifier before each part except first
            modifier_run = create_run(modifier_char, is_modifier=True)
            para.insert(current_idx, modifier_run)
            current_idx += 1

        if part:  # Insert text part if not empty
            part_run = create_run(part, is_modifier=False)
            para.insert(current_idx, part_run)
            current_idx += 1

    return True

def process_paragraph(para):
    """Process paragraph to replace all apostrophe variants and split modifiers."""
    fixes = 0

    for run in list(para.runs):  # Iterate over copy
        if not run.text:
            continue

        original_text = run.text
        new_text = original_text

        # Apply all replacements
        for old_pattern, new_pattern in ALL_REPLACEMENTS:
            if old_pattern in new_text:
                new_text = new_text.replace(old_pattern, new_pattern)

        if new_text != original_text:
            # Text was changed, update run and split modifiers
            run.text = new_text

            # Split both ALEPH and AYIN modifiers
            if ALEPH in new_text:
                if split_run_at_modifier(run._element, new_text, ALEPH):
                    fixes += new_text.count(ALEPH)
            if AYIN in new_text:
                if split_run_at_modifier(run._element, new_text, AYIN):
                    fixes += new_text.count(AYIN)

            break  # Run modified, move to next

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables (including TOC)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Replacing ALL apostrophe variants with proper modifiers...")
print(f"Handling apostrophes: {', '.join(repr(a) for a in APOSTROPHE_VARIANTS)}")
print("(Processing body, tables, TOC, headers, footers)")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in sorted(docx_files):
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total apostrophe variants replaced: {total_fixes}")
