"""Debug: inspect XML structure of TOC paragraphs to understand why replacements fail."""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from pathlib import Path

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')
doc_path = DOCS_DIR / 'YY-s01v02-An Intro to God-Mitswah-Instructions.docx'
doc = Document(str(doc_path))

# Para 29 has: "1\tYatsa' ~ Delivered I Am Yahowah \t  1"
para = doc.paragraphs[29]

print(f'para.text = "{para.text[:80]}"')
print(f'Number of runs: {len(para.runs)}')
print()

# Show each run's text
for i, run in enumerate(para.runs):
    print(f'  run[{i}]: text="{run.text}" italic={run.font.italic} bold={run.font.bold}')

print()

# Show full XML of this paragraph
xml_str = etree.tostring(para._element, pretty_print=True).decode('utf-8')
print('Full XML:')
print(xml_str[:3000])

# Also check paragraph 30 (Mala'kah)
print('\n\n--- Para 30 ---')
para30 = doc.paragraphs[30]
print(f'para.text = "{para30.text[:80]}"')
print(f'Number of runs: {len(para30.runs)}')
for i, run in enumerate(para30.runs):
    print(f'  run[{i}]: text="{run.text}"')

# Check: does para.text concatenation match run texts?
run_text = ''.join(r.text for r in para.runs)
print(f'\npara.text == run concat: {para.text == run_text}')
print(f'para.text length: {len(para.text)}')
print(f'run concat length: {len(run_text)}')
if para.text != run_text:
    print(f'para.text: {repr(para.text[:100])}')
    print(f'run concat: {repr(run_text[:100])}')

    # Find where they differ
    for j in range(min(len(para.text), len(run_text))):
        if para.text[j] != run_text[j]:
            print(f'  First diff at pos {j}: para={repr(para.text[j])} run={repr(run_text[j])}')
            break

    # Check if there's text in non-run elements (hyperlinks, field codes)
    all_text_elements = para._element.findall('.//' + qn('w:t'))
    all_text = ''.join(t.text for t in all_text_elements if t.text)
    print(f'\nAll w:t elements text: "{all_text[:100]}"')
    print(f'Count of w:t elements: {len(all_text_elements)}')

    # Check for hyperlinks
    hyperlinks = para._element.findall('.//' + qn('w:hyperlink'))
    print(f'Hyperlinks: {len(hyperlinks)}')
    for hl in hyperlinks:
        hl_runs = hl.findall('.//' + qn('w:r'))
        for r in hl_runs:
            t = r.find(qn('w:t'))
            if t is not None and t.text:
                print(f'  Hyperlink run text: "{t.text}"')
