"""
Simple test: Update one Word document with Hebrew modifiers.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

print("Starting script...")

try:
    from docx import Document
    print("✓ python-docx imported")
except ImportError as e:
    print(f"ERROR: Could not import python-docx: {e}")
    sys.exit(1)

# Load conversions
conversions_file = r'C:\Users\Joe\Work\dev\yada\translations\hebrew_modifier_conversions_refined.txt'
conversions = {}

print(f"Loading conversions from: {conversions_file}")
try:
    with open(conversions_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        for line in lines[2:]:
            if ', ' in line and not line.startswith('===') and not line.startswith('Score'):
                parts = line.strip().split(', ', 1)
                if len(parts) == 2:
                    orig, conv = parts
                    conversions[orig] = conv
    print(f"✓ Loaded {len(conversions)} conversions")
except Exception as e:
    print(f"ERROR loading conversions: {e}")
    sys.exit(1)

# Find first YY document
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print(f"\nLooking for documents in: {docs_dir}")

if not os.path.exists(docs_dir):
    print(f"ERROR: Directory not found")
    sys.exit(1)

files = [f for f in os.listdir(docs_dir) if f.startswith('YY-') and f.endswith('.docx') and not '_updated' in f]
print(f"Found {len(files)} YY documents")

if not files:
    print("No documents to process")
    sys.exit(0)

# Process first document only
test_file = files[0]
docx_path = os.path.join(docs_dir, test_file)
print(f"\nProcessing TEST file: {test_file}")

try:
    print("  Opening document...")
    doc = Document(docx_path)
    print(f"  ✓ Opened ({len(doc.paragraphs)} paragraphs)")

    replacements = 0
    para_count = 0

    print("  Processing paragraphs...")
    for para in doc.paragraphs[:100]:  # First 100 paragraphs only
        para_count += 1
        for run in para.runs:
            if not run.text:
                continue

            original = run.text
            new_text = original

            # Simple replacement
            for orig, conv in conversions.items():
                if orig in new_text:
                    new_text = new_text.replace(orig, conv)

            if new_text != original:
                run.text = new_text
                replacements += 1

        if para_count % 20 == 0:
            print(f"    Processed {para_count} paragraphs, {replacements} replacements...")

    print(f"  ✓ Processed {para_count} paragraphs, made {replacements} replacements")

    if replacements > 0:
        output_path = os.path.join(docs_dir, test_file.replace('.docx', '_test_updated.docx'))
        print(f"  Saving to: {os.path.basename(output_path)}")
        doc.save(output_path)
        print(f"  ✓ Saved successfully")
    else:
        print("  No changes needed")

except Exception as e:
    print(f"  ERROR: {e}")
    import traceback
    traceback.print_exc()

print("\nDone!")
