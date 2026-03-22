"""
Update YY Word documents with proper Hebrew modifiers (Aleph ʾ and Ayin ʿ)
and apply "Yada Towrah" font to the modifier characters.

This script:
1. Loads conversion mappings from hebrew_modifier_conversions_refined.txt
2. Processes all YY*.docx files
3. Replaces curly quotes with Hebrew modifiers in italic text
4. Sets font to "Yada Towrah" for modifier characters only
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

# Hebrew modifier characters
ALEPH = chr(0x02BE)  # ʾ (modifier letter right half ring)
AYIN = chr(0x02BF)   # ʿ (modifier letter left half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)

# Load conversions
conversions_file = r'C:\Users\Joe\Work\dev\yada\translations\hebrew_modifier_conversions_refined.txt'
conversions = {}

print("Loading conversions...")
with open(conversions_file, 'r', encoding='utf-8') as f:
    lines = f.readlines()
    for line in lines[2:]:  # Skip header
        if ', ' in line and not line.startswith('===') and not line.startswith('Score'):
            parts = line.strip().split(', ', 1)
            if len(parts) == 2:
                orig, conv = parts
                conversions[orig] = conv

print(f"Loaded {len(conversions)} conversions\n")

def split_run_at_modifiers(run, new_text):
    """
    Replace run text with new_text, applying "Yada Towrah" font to modifier chars.
    Creates multiple runs if needed to handle font changes.
    """
    if new_text == run.text:
        # No change
        return

    # Save original formatting
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name
    font_size = run.font.size

    # Clear the run text
    run.text = ""

    # Get the parent paragraph
    para = run._element.getparent()
    run_idx = list(para).index(run._element)

    # Build runs for each character, grouping non-modifiers together
    i = 0
    buffer = []

    def flush_buffer(is_modifier=False):
        """Helper to create a run from buffered characters."""
        if not buffer:
            return
        text = ''.join(buffer)

        # Create new run element
        new_run = OxmlElement('w:r')

        # Copy run properties
        rPr = OxmlElement('w:rPr')

        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Set font - "Yada Towrah" for modifiers, original for others
        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if is_modifier else (font_name or "Calibri")
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        # Set font size if available
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size.pt * 2))  # Half-points
            rPr.append(sz)

        new_run.append(rPr)

        # Add text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        # Insert after original run
        nonlocal run_idx
        para.insert(run_idx + 1, new_run)
        run_idx += 1

        buffer.clear()

    # Process each character
    while i < len(new_text):
        char = new_text[i]
        if char in (ALEPH, AYIN):
            # Flush any pending non-modifier text
            flush_buffer(is_modifier=False)
            # Add modifier char to buffer and flush
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)
        i += 1

    # Flush remaining buffer
    flush_buffer(is_modifier=False)

def process_paragraph(para):
    """Process a paragraph, replacing Hebrew words in italic runs."""
    replacements = 0

    # Process runs in reverse so we can safely modify/remove them
    for run in list(para.runs):
        if not run.text or not run.italic:
            continue

        original = run.text
        new_text = original

        # Try each conversion
        for orig, conv in conversions.items():
            if orig in new_text:
                new_text = new_text.replace(orig, conv)

        if new_text != original:
            # Replace with formatted version
            split_run_at_modifiers(run, new_text)
            replacements += 1

    return replacements

def process_document(docx_path):
    """Process a single Word document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")

    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: Could not open document: {e}")
        return 0

    total_replacements = 0

    # Process all paragraphs
    for para in doc.paragraphs:
        replacements = process_paragraph(para)
        total_replacements += replacements

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replacements = process_paragraph(para)
                    total_replacements += replacements

    elapsed = time.time() - start_time

    if total_replacements > 0:
        # Save with _updated_fonts suffix
        base, ext = os.path.splitext(docx_path)
        # Remove old _updated suffix if present
        if base.endswith('_updated'):
            base = base[:-8]
        output_path = f"{base}_updated_fonts{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ {total_replacements} replacements, saved to {os.path.basename(output_path)} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: Could not save document: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return total_replacements

# Find all YY Word documents
docs_dir = r'C:\users\joe\work\dev\yada\docs'
if not os.path.exists(docs_dir):
    print(f"ERROR: Documents directory not found: {docs_dir}")
    sys.exit(1)

print(f"\nSearching for .docx files in: {docs_dir}")
docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-')
    and f.endswith('.docx')
    and not f.endswith('_updated.docx')
    and not f.endswith('_updated_fonts.docx')
    and not f.startswith('~$')
]

print(f"Found {len(docx_files)} YY documents to process\n")
print("=" * 80)

total_files_updated = 0
total_replacements_overall = 0

for docx_path in docx_files:
    replacements = process_document(docx_path)
    if replacements > 0:
        total_files_updated += 1
        total_replacements_overall += replacements

print("\n" + "=" * 80)
print("WORD DOCUMENT UPDATE COMPLETE")
print("=" * 80)
print(f"  Files processed: {len(docx_files)}")
print(f"  Files updated: {total_files_updated}")
print(f"  Total replacements: {total_replacements_overall}")
print(f"\nUpdated files saved with '_updated_fonts.docx' suffix in: {docs_dir}")
