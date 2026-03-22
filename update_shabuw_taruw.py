"""
Update Shabuw'ah and Taruw'ah with proper Ayin modifiers in all YY Word documents.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

# Character definitions
AYIN = chr(0x02BF)   # ʿ (modifier letter left half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

# Replacement patterns
REPLACEMENTS = [
    # Shabuw'ah → Shabuwʿah
    (re.compile(r"Shabuw['\\u2018\\u2019\\u2032]ah\\b", re.IGNORECASE),
     lambda m: f'Shabuw{AYIN}ah' if m.group()[0].isupper() else f'shabuw{AYIN}ah'),

    # Taruw'ah → Taruwʿah
    (re.compile(r"Taruw['\\u2018\\u2019\\u2032]ah\\b", re.IGNORECASE),
     lambda m: f'Taruw{AYIN}ah' if m.group()[0].isupper() else f'taruw{AYIN}ah'),
]

def apply_replacements(text):
    """Apply all replacement patterns."""
    original = text
    for pattern, replacement_func in REPLACEMENTS:
        text = pattern.sub(replacement_func, text)
    return text if text != original else None

def create_formatted_runs(run, new_text):
    """Create runs with proper font formatting for modifiers."""
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name or "Times New Roman"
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    para = run._element.getparent()
    run_idx = list(para).index(run._element)
    run.text = ""

    def create_run(text, use_yada_font=False):
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else font_name
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_modifier=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_modifier)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    for char in new_text:
        if char == AYIN:
            flush_buffer(is_modifier=False)
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)

    flush_buffer(is_modifier=False)

def ensure_modifier_fonts(run):
    """Ensure existing modifiers use Yada Towrah font."""
    if not run.text:
        return False
    if AYIN not in run.text:
        return False
    create_formatted_runs(run, run.text)
    return True

def process_paragraph(para, stats):
    """Process paragraph for word replacements."""
    for run in list(para.runs):
        if not run.text:
            continue

        new_text = apply_replacements(run.text)
        if new_text:
            create_formatted_runs(run, new_text)
            stats['replacements'] += 1
        else:
            if ensure_modifier_fonts(run):
                stats['font_fixes'] += 1

def process_document(docx_path):
    """Process a document."""
    print(f"\\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'replacements': 0, 'font_fixes': 0}

    stats = {
        'replacements': 0,
        'font_fixes': 0
    }

    # Process paragraphs
    for para in doc.paragraphs:
        process_paragraph(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    repl = stats['replacements']
    fonts = stats['font_fixes']

    if repl > 0 or fonts > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Words: {repl}, Fonts: {fonts} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'replacements': 0, 'font_fixes': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'replacements': repl, 'font_fixes': fonts}

# Main
docs_dir = r'C:\\users\\joe\\work\\dev\\yada\\docs'
print("Updating Shabuw'ah and Taruw'ah with proper modifiers...")
print()

# Get all YY files (excluding intermediate files)
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\\n" + "=" * 80)

total_updated = 0
total_repl = 0
total_fonts = 0

for path in docx_files:
    result = process_document(path)
    if result['replacements'] > 0 or result['font_fixes'] > 0:
        total_updated += 1
        total_repl += result['replacements']
        total_fonts += result['font_fixes']

print("\\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Word replacements: {total_repl}")
print(f"Font corrections: {total_fonts}")
print(f"\\nUpdated files: {docs_dir}/*_updated.docx")
