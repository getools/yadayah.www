"""
Split text elements containing ʿ (Ayin) and apply Yada Towrah font to the modifier.
Processes all parts: body, tables, TOC, headers, footers.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

AYIN = chr(0x02BF)  # ʿ

def split_and_fix_element(t_elem, run_elem):
    """Split text element if it contains ʿ, and apply Yada Towrah to the modifier."""
    if not t_elem.text or AYIN not in t_elem.text:
        return False

    # Get run properties to preserve formatting
    rPr = run_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')

    # Get parent of the run
    parent = run_elem.getparent()
    run_idx = list(parent).index(run_elem)
    text = t_elem.text

    # Clear the original text
    t_elem.text = ""

    def create_run(content, is_modifier=False):
        """Create a new run with proper formatting."""
        new_run = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')

        # Copy formatting properties
        if rPr is not None:
            # Copy bold
            bold = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
            if bold is not None:
                new_rPr.append(OxmlElement('w:b'))

            # Copy italic
            italic = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i')
            if italic is not None:
                new_rPr.append(OxmlElement('w:i'))

            # Copy size
            sz = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
            if sz is not None:
                import copy
                new_rPr.append(copy.deepcopy(sz))

            szCs = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs')
            if szCs is not None:
                import copy
                new_rPr.append(copy.deepcopy(szCs))

            # Copy color
            color = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            if color is not None:
                import copy
                new_rPr.append(copy.deepcopy(color))

            # Copy original font if not modifier
            if not is_modifier:
                rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                if rFonts is not None:
                    import copy
                    new_rPr.append(copy.deepcopy(rFonts))

        # Set font for modifier or ensure it's set
        if is_modifier:
            rFonts_elem = OxmlElement('w:rFonts')
            rFonts_elem.set(qn('w:ascii'), 'Yada Towrah')
            rFonts_elem.set(qn('w:hAnsi'), 'Yada Towrah')
            new_rPr.append(rFonts_elem)

        new_run.append(new_rPr)

        # Add text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        new_run.append(t)

        return new_run

    # Split and create runs
    current_idx = run_idx
    buffer = []

    def flush_buffer(is_modifier=False):
        nonlocal current_idx
        if not buffer:
            return
        content = ''.join(buffer)
        new_run = create_run(content, is_modifier=is_modifier)
        parent.insert(current_idx + 1, new_run)
        current_idx += 1
        buffer.clear()

    # Process each character
    for char in text:
        if char == AYIN:
            flush_buffer(is_modifier=False)
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)

    flush_buffer(is_modifier=False)

    return True

def process_paragraph(para):
    """Process all text elements in a paragraph."""
    fixes = 0

    for element in list(para._element.iter()):
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            run_elem = element.getparent()
            if run_elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                if split_and_fix_element(element, run_elem):
                    fixes += 1

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Split and fixed {fixes} elements ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Splitting all ʿ (Ayin) modifiers into separate runs with Yada Towrah font...")
print("(Processing body, tables, TOC, headers, footers)")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total elements split and fixed: {total_fixes}")
