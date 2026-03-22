"""
Update all YY Word documents:
1. Replace 'Eden or 'eden with ʾEden (ʾ in Yada Towrah font)
2. Replace all Calibri font with Times New Roman
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

# Character definitions
ALEPH = chr(0x02BE)  # ʾ (modifier letter right half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

def normalize_quotes(s):
    """Normalize all quote variations to straight apostrophe."""
    return s.replace(LEFT_CURLY, "'").replace(RIGHT_CURLY, "'").replace(chr(0x2032), "'")

def should_replace_eden(text):
    """Check if text contains 'Eden or 'eden (any quote variation)."""
    normalized = normalize_quotes(text).lower()
    return "'eden" in normalized

def apply_eden_capitalization(original):
    """
    Determine capitalization for ʾEden based on original.
    'Eden -> ʾEden, 'eden -> ʾeden, 'EDEN -> ʾEDEN
    """
    # Normalize quotes in original
    normalized = normalize_quotes(original).lower()

    if "'eden" not in normalized:
        return "ʾEden"  # Default

    # Find the 'eden substring
    pattern = re.compile(r"['\u2018\u2019\u2032]eden", re.IGNORECASE)
    match = pattern.search(original)
    if not match:
        return "ʾEden"

    source = match.group()

    # Check if all caps
    source_letters = [c for c in source if c.isalpha()]
    if source_letters and all(c.isupper() for c in source_letters):
        return "ʾEDEN"

    # Check if first letter is capital
    if source_letters and source_letters[0].isupper():
        return "ʾEden"
    else:
        return "ʾeden"

def create_eden_runs(run, original_text, replacement_text):
    """
    Create new runs for ʾEden with proper font formatting.
    ʾ uses "Yada Towrah", other letters preserve original font.
    """
    # Save original formatting
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name or "Calibri"
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    # Get parent paragraph and run position
    para = run._element.getparent()
    run_idx = list(para).index(run._element)

    # Clear original run
    run.text = ""

    def create_run(text, use_yada_font=False):
        """Create a run with specified formatting."""
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        # Italic
        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Font - Yada Towrah for ʾ, Times New Roman for regular text (or original if not Calibri)
        rFonts = OxmlElement('w:rFonts')
        if use_yada_font:
            target_font = "Yada Towrah"
        else:
            # Replace Calibri with Times New Roman
            target_font = "Times New Roman" if font_name == "Calibri" else font_name
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        # Font size
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        # Color
        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    # Split replacement text at ʾ character
    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_aleph=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_aleph)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    for char in replacement_text:
        if char == ALEPH:
            # Flush regular text
            flush_buffer(is_aleph=False)
            # Add ʾ and flush
            buffer.append(char)
            flush_buffer(is_aleph=True)
        else:
            buffer.append(char)

    # Flush remaining
    flush_buffer(is_aleph=False)

def replace_calibri_font(run):
    """Replace Calibri font with Times New Roman in a run."""
    if run.font.name == "Calibri":
        run.font.name = "Times New Roman"
        return True
    return False

def process_paragraph(para, stats):
    """Process a paragraph: replace 'Eden and change fonts."""
    for run in list(para.runs):
        if not run.text:
            continue

        # Task 1: Replace 'Eden with ʾEden
        if should_replace_eden(run.text):
            replacement = apply_eden_capitalization(run.text)

            # Find and replace 'eden pattern
            new_text = run.text
            pattern = re.compile(r"['\u2018\u2019\u2032]eden", re.IGNORECASE)
            match = pattern.search(new_text)

            if match:
                before = new_text[:match.start()]
                after = new_text[match.end():]
                final_text = before + replacement + after

                if final_text != run.text:
                    create_eden_runs(run, run.text, final_text)
                    stats['eden_replacements'] += 1
                    if stats['verbose']:
                        print(f"    Eden: {run.text} -> {final_text}")

        # Task 2: Replace Calibri font with Times New Roman
        if replace_calibri_font(run):
            stats['font_replacements'] += 1

def process_document(docx_path, verbose=False):
    """Process a single Word document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: Could not open document: {e}")
        return {'eden': 0, 'fonts': 0}

    stats = {
        'eden_replacements': 0,
        'font_replacements': 0,
        'verbose': verbose
    }

    # Process paragraphs
    for para in doc.paragraphs:
        process_paragraph(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    eden_count = stats['eden_replacements']
    font_count = stats['font_replacements']

    if eden_count > 0 or font_count > 0:
        # Save with _updated suffix
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Eden: {eden_count}, Fonts: {font_count}, saved to {os.path.basename(output_path)} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: Could not save document: {e}")
            return {'eden': 0, 'fonts': 0}
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return {'eden': eden_count, 'fonts': font_count}

# Main execution
docs_dir = r'C:\users\joe\work\dev\yada\docs'
if not os.path.exists(docs_dir):
    print(f"ERROR: Documents directory not found: {docs_dir}")
    sys.exit(1)

print("Updating YY Word documents...")
print("1. Replacing 'Eden with ʾEden (ʾ in Yada Towrah font)")
print("2. Replacing Calibri font with Times New Roman")
print()

# Find all YY documents
docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-')
    and f.endswith('.docx')
    and not '_updated' in f
    and not '_yashayah' in f
    and not f.startswith('~$')
]

print(f"Found {len(docx_files)} YY documents to process\n")
print("=" * 80)

total_files_updated = 0
total_eden_replacements = 0
total_font_replacements = 0

for docx_path in docx_files:
    result = process_document(docx_path, verbose=False)
    if result['eden'] > 0 or result['fonts'] > 0:
        total_files_updated += 1
        total_eden_replacements += result['eden']
        total_font_replacements += result['fonts']

print("\n" + "=" * 80)
print("UPDATE COMPLETE")
print("=" * 80)
print(f"  Files processed: {len(docx_files)}")
print(f"  Files updated: {total_files_updated}")
print(f"  'Eden → ʾEden replacements: {total_eden_replacements}")
print(f"  Calibri → Times New Roman replacements: {total_font_replacements}")
print(f"\nUpdated files saved with '_updated.docx' suffix in: {docs_dir}")
