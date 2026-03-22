"""
Replace remaining curly apostrophes (') with proper ALEPH (ʾ) or AYIN (ʿ) modifiers.
Then split modifiers into separate runs with Yada Towrah font.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

ALEPH = chr(0x02BE)  # ʾ
AYIN = chr(0x02BF)   # ʿ
CURLY_APOS = chr(0x2019)  # '

# Replacements with curly apostrophe specifically
CURLY_REPLACEMENTS = [
    # ALEPH
    (f"Yisrae{CURLY_APOS}l", "Yisraʾel"),
    (f"{CURLY_APOS}Adam", "ʾAdam"),
    (f"{CURLY_APOS}Abraham", "ʾAbraham"),
    (f"Miqra{CURLY_APOS}", "Miqraʾ"),
    (f"{CURLY_APOS}el", "ʾel"),

    # AYIN
    (f"Yisra{CURLY_APOS}el", "Yisraʿel"),
    (f"Ya{CURLY_APOS}aqob", "Yaʿaqob"),
    (f"Mow{CURLY_APOS}ab", "Mowʿab"),
    (f"Mow{CURLY_APOS}edym", "Mowʿedym"),
    (f"Yahowsha{CURLY_APOS}", "Yahowshaʿ"),
    (f"Yow{CURLY_APOS}ab", "Yowʿab"),
    (f"Ga{CURLY_APOS}al", "Gaʿal"),
    (f"yasha{CURLY_APOS}", "yashaʿ"),
    (f"Yisra{CURLY_APOS}elites", "Yisraʿelites"),
    (f"Mow{CURLY_APOS}abites", "Mowʿabites"),
    (f"Yada{CURLY_APOS}", "Yadaʿ"),
    (f"Zarowa{CURLY_APOS}", "Zarowaʿ"),
    (f"Yow{CURLY_APOS}el", "Yowʿel"),
    (f"Yasha{CURLY_APOS}", "Yashaʿ"),
    (f"Howsha{CURLY_APOS}", "Howshaʿ"),
    (f"Sha{CURLY_APOS}uwl", "Shaʿuwl"),
]

def split_run_at_modifier(run_elem, text, modifier_char):
    """Split run to isolate modifier with Yada Towrah font."""
    if modifier_char not in text:
        return False

    # Get parent paragraph
    para = run_elem.getparent()
    if para is None:
        return False

    run_idx = list(para).index(run_elem)

    # Get original formatting
    rPr = run_elem.find(qn('w:rPr'))

    def clone_rPr(exclude_font=False):
        new_rPr = OxmlElement('w:rPr')
        if rPr is not None:
            for child in rPr:
                if exclude_font and child.tag == qn('w:rFonts'):
                    continue
                import copy
                new_rPr.append(copy.deepcopy(child))
        return new_rPr

    def create_run(content, is_modifier=False):
        new_run = OxmlElement('w:r')
        new_rPr = clone_rPr(exclude_font=is_modifier)

        if is_modifier:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Yada Towrah')
            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
            new_rPr.append(rFonts)

        new_run.append(new_rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        new_run.append(t)

        return new_run

    # Remove original run
    para.remove(run_elem)

    # Split text and create runs
    current_idx = run_idx
    parts = text.split(modifier_char)

    for i, part in enumerate(parts):
        if i > 0:  # Insert modifier before each part except first
            modifier_run = create_run(modifier_char, is_modifier=True)
            para.insert(current_idx, modifier_run)
            current_idx += 1

        if part:  # Insert text part if not empty
            part_run = create_run(part, is_modifier=False)
            para.insert(current_idx, part_run)
            current_idx += 1

    return True

def process_paragraph(para):
    """Process paragraph to replace curly apostrophes and split modifiers."""
    fixes = 0

    for run in list(para.runs):  # Iterate over copy
        if not run.text or CURLY_APOS not in run.text:
            continue

        original_text = run.text
        new_text = original_text

        # Apply replacements
        for old_pattern, new_pattern in CURLY_REPLACEMENTS:
            new_text = new_text.replace(old_pattern, new_pattern)

        if new_text != original_text:
            # Text was changed, split modifiers
            run.text = new_text

            # Check which modifier is present
            if ALEPH in new_text:
                if split_run_at_modifier(run._element, new_text, ALEPH):
                    fixes += new_text.count(ALEPH)
            elif AYIN in new_text:
                if split_run_at_modifier(run._element, new_text, AYIN):
                    fixes += new_text.count(AYIN)

            break  # Run modified, move to next

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables (including TOC)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Replacing curly apostrophes (') with proper modifiers...")
print("(Processing body, tables, TOC, headers, footers)")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in sorted(docx_files):
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total curly apostrophes replaced: {total_fixes}")
