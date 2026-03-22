"""Check all YY docs for remaining quote chars in TOC entries and chapter headings."""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')
QUOTE_CHARS = set("'\u2018\u2019`\u00B4\u02BC")
HALF_RING_ALEPH = '\u02BE'
HALF_RING_AYIN = '\u02BF'

files = sorted(DOCS_DIR.glob('YY*.docx'))
files = [f for f in files if '_backup' not in f.name and '_updated' not in f.name
         and '_halfring' not in f.name]

toc_words_with_quotes = set()

for fp in files:
    doc = Document(str(fp))
    for i, para in enumerate(doc.paragraphs[:80]):
        text = para.text
        # Detect TOC entries: they have field codes (PAGEREF or TOC)
        has_field = bool(para._element.findall('.//' + qn('w:instrText')))
        # Also check paragraph style
        style_name = para.style.name if para.style else ''
        is_toc = has_field or 'toc' in style_name.lower()

        if not is_toc:
            continue

        quotes = [ch for ch in text if ch in QUOTE_CHARS]
        if quotes:
            # Extract the transliterated words with quotes
            words = text.split()
            for w in words:
                if any(ch in QUOTE_CHARS for ch in w):
                    # Clean off tab chars and numbers
                    clean = w.strip('\t0123456789 ')
                    if clean and any(ch in QUOTE_CHARS for ch in clean):
                        toc_words_with_quotes.add(clean)
            print(f'{fp.name} para[{i}]: "{text[:80]}"')

print(f'\nUnique TOC words still with quotes: {len(toc_words_with_quotes)}')
for w in sorted(toc_words_with_quotes):
    print(f'  {w}')

# Also check chapter heading paragraphs (Heading 1/2 styles) for remaining quotes
print(f'\nHeading paragraphs with remaining quotes:')
heading_words = set()
for fp in files:
    doc = Document(str(fp))
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if 'heading' in style_name.lower() or 'Heading' in style_name:
            text = para.text
            if any(ch in QUOTE_CHARS for ch in text):
                words = text.split()
                for w in words:
                    if any(ch in QUOTE_CHARS for ch in w):
                        clean = w.strip('\t0123456789.,;:!? ')
                        if clean:
                            heading_words.add(clean)
                print(f'  {fp.name}: "{text[:80]}" [{style_name}]')

print(f'\nUnique heading words with quotes: {len(heading_words)}')
for w in sorted(heading_words):
    print(f'  {w}')
