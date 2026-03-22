import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')

# Find the paragraph mentioned
target_text = 'The fourth Miqra'
for i, para in enumerate(doc.paragraphs):
    if target_text in para.text:
        print(f'Found at index {i}:')
        print(f'  Text: {para.text[:100]}...')
        print(f'  Alignment: {para.alignment}')
        print(f'  First-line indent: {para.paragraph_format.first_line_indent}')

        # Check previous paragraphs
        if i > 0:
            prev_para = doc.paragraphs[i-1]
            print(f'\nPrevious paragraph (index {i-1}):')
            print(f'  Text: {prev_para.text[:150]}')
            print(f'  Ends with ...: {prev_para.text.strip().endswith("...")}')
            print(f'  Ends with …: {prev_para.text.strip().endswith("…")}')
            print(f'  Alignment: {prev_para.alignment}')

            # Check italic
            italic_chars = 0
            total_chars = 0
            for r in prev_para.runs:
                if r.text.strip():
                    total_chars += len(r.text.strip())
                    if r.italic:
                        italic_chars += len(r.text.strip())

            if total_chars > 0:
                print(f'  Italic percentage: {italic_chars}/{total_chars} = {100*italic_chars/total_chars:.1f}%')

            # Check each run
            print(f'  Runs:')
            for j, run in enumerate(prev_para.runs):
                if run.text.strip():
                    print(f'    Run {j}: italic={run.italic}, text="{run.text.strip()[:40]}"')

        break
