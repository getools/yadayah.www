"""Check for Abraham and Adam in s04v01 document."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s04v01-Coming Home-Qowl-A Voice.docx')

CURLY_APOS = chr(0x2019)  # '
ALEPH = chr(0x02BE)  # ʾ
REGULAR_APOS = "'"
BACKTICK = "`"
LEFT_SINGLE = chr(0x2018)  # '

print('Searching ALL sections for Abraham and Adam...\n')

def check_text(text, location):
    if 'Abraham' in text or 'Adam' in text:
        # Check what character precedes these words
        for word in ['Abraham', 'Adam']:
            if word in text:
                idx = text.index(word)
                if idx > 0:
                    prev_char = text[idx-1]
                    print(f'{location} - Found {word}')
                    print(f'  Previous char: {repr(prev_char)} (ord={ord(prev_char)})')
                    context = text[max(0, idx-25):min(len(text), idx+30)]
                    print(f'  Context: ...{context}...')

                    # Check if it's one of the expected apostrophes
                    if prev_char == REGULAR_APOS:
                        print(f'  → Has REGULAR apostrophe')
                    elif prev_char == CURLY_APOS:
                        print(f'  → Has CURLY apostrophe')
                    elif prev_char == ALEPH:
                        print(f'  → Has ALEPH modifier (CORRECT!)')
                    elif prev_char == LEFT_SINGLE:
                        print(f'  → Has LEFT SINGLE quote')
                    elif prev_char == BACKTICK:
                        print(f'  → Has BACKTICK')
                    else:
                        print(f'  → Has other character')
                    print()

# Check all paragraphs
for i, para in enumerate(doc.paragraphs):
    check_text(para.text, f'Body para {i}')

# Check tables
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, para in enumerate(cell.paragraphs):
                check_text(para.text, f'Table {t_idx} row {r_idx} cell {c_idx}')

# Check headers
for s_idx, section in enumerate(doc.sections):
    if section.header:
        for p_idx, para in enumerate(section.header.paragraphs):
            check_text(para.text, f'Header section {s_idx} para {p_idx}')

    if section.footer:
        for p_idx, para in enumerate(section.footer.paragraphs):
            check_text(para.text, f'Footer section {s_idx} para {p_idx}')

print("Search complete.")
