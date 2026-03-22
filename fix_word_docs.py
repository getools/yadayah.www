"""
Fix YY Word documents:
1. Remove comma from yy_break centered HWHY, paragraphs
2. Replace standalone 'why' with 'hwhy' in Yada Towrah font runs
3. Replace 'hwhy,,' with 'hwhy,'
4. Replace U+F067 (PUA char) with '8'
5. Replace Towrahʿs with Towrah's (set to Times New Roman)
"""
import os, sys, glob, re, copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOCS_DIR = r'C:\Users\Joe\Work\dev\yada\docs'
YT_FONT = 'Yada Towrah'
TNR_FONT = 'Times New Roman'

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}


def get_all_runs(para):
    """Get ALL w:r elements including inside hyperlinks."""
    runs = []
    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            runs.append(child)
        elif tag == 'hyperlink':
            for r in child.findall(qn('w:r')):
                runs.append(r)
    return runs


def get_run_font(r_elem):
    """Get font name from a w:r element."""
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi')) or ''
    return ''


def get_run_text(r_elem):
    """Get text from a w:r element."""
    parts = []
    for t in r_elem.findall(qn('w:t')):
        parts.append(t.text or '')
    return ''.join(parts)


def set_run_text(r_elem, new_text):
    """Set text on a w:r element, preserving space."""
    t_elems = r_elem.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = new_text
        if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
            t_elems[0].set(qn('xml:space'), 'preserve')
        for extra in t_elems[1:]:
            r_elem.remove(extra)
    elif new_text:
        t = etree.SubElement(r_elem, qn('w:t'))
        t.text = new_text
        if new_text[0] == ' ' or new_text[-1] == ' ':
            t.set(qn('xml:space'), 'preserve')


def set_run_font(r_elem, font_name):
    """Set font on a w:r element."""
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(r_elem, qn('w:rPr'))
        r_elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def build_cmap(runs):
    """Build character-to-run map: list of (char, run_elem, char_index_in_run, font)."""
    cmap = []
    for r in runs:
        text = get_run_text(r)
        font = get_run_font(r)
        for ci, ch in enumerate(text):
            cmap.append((ch, r, ci, font))
    return cmap


def split_run_at(r_elem, char_idx):
    """Split a run at char_idx, return (before_run, after_run). Original is modified to before."""
    text = get_run_text(r_elem)
    before_text = text[:char_idx]
    after_text = text[char_idx:]

    after_run = copy.deepcopy(r_elem)
    set_run_text(r_elem, before_text)
    set_run_text(after_run, after_text)

    parent = r_elem.getparent()
    parent.insert(list(parent).index(r_elem) + 1, after_run)
    return r_elem, after_run


def process_paragraphs(paragraphs, stats, is_header=False):
    """Process a list of paragraph objects."""
    for para in paragraphs:
        all_runs = get_all_runs(para)
        if not all_runs:
            continue

        style_name = para.style.name if para.style else ''

        # --- Task 1: yy_break paragraphs with comma ---
        if style_name == 'yy_break':
            full_text = ''.join(get_run_text(r) for r in all_runs)
            if ',' in full_text:
                for r in all_runs:
                    txt = get_run_text(r)
                    if ',' in txt:
                        set_run_text(r, txt.replace(',', ''))
                        stats['yy_break_comma'] += txt.count(',')

        # --- Task 4: Replace U+F067 with '8' ---
        for r in all_runs:
            txt = get_run_text(r)
            if '\uf067' in txt:
                count = txt.count('\uf067')
                set_run_text(r, txt.replace('\uf067', '8'))
                stats['pua_f067'] += count

        # --- Task 3: Replace hwhy,, with hwhy, ---
        # Need cross-run matching
        all_runs = get_all_runs(para)  # refresh after modifications
        cmap = build_cmap(all_runs)
        full_text = ''.join(c[0] for c in cmap)

        pos = 0
        while True:
            idx = full_text.find('hwhy,,', pos)
            if idx < 0:
                break
            # Check if in YT font (at least the hwhy part)
            yt_chars = sum(1 for i in range(idx, idx+4) if cmap[i][3].lower() == 'yada towrah')
            if yt_chars >= 2:
                # Remove one comma (the 5th character = idx+5, which is the second comma)
                comma_entry = cmap[idx + 5]
                r_elem = comma_entry[1]
                ci = comma_entry[2]
                txt = get_run_text(r_elem)
                new_txt = txt[:ci] + txt[ci+1:]
                set_run_text(r_elem, new_txt)
                stats['double_comma'] += 1
                # Rebuild cmap
                all_runs = get_all_runs(para)
                cmap = build_cmap(all_runs)
                full_text = ''.join(c[0] for c in cmap)
            pos = idx + 4

        # --- Task 2: Replace standalone 'why' with 'hwhy' in YT font ---
        # Standalone = not preceded by another YT font character
        all_runs = get_all_runs(para)
        cmap = build_cmap(all_runs)
        full_text = ''.join(c[0] for c in cmap)

        replacements = []
        pos = 0
        while True:
            idx = full_text.find('why', pos)
            if idx < 0:
                break
            # Check all 3 chars are in YT font
            if (idx + 2 < len(cmap) and
                cmap[idx][3].lower() == 'yada towrah' and
                cmap[idx+1][3].lower() == 'yada towrah' and
                cmap[idx+2][3].lower() == 'yada towrah'):
                # Check NOT preceded by YT font alpha char
                preceded_by_yt = False
                if idx > 0 and cmap[idx-1][3].lower() == 'yada towrah' and cmap[idx-1][0].isalpha():
                    preceded_by_yt = True
                # Check NOT followed by YT font alpha char (beyond 'y')
                followed_by_yt = False
                if idx + 3 < len(cmap) and cmap[idx+3][3].lower() == 'yada towrah' and cmap[idx+3][0].isalpha():
                    followed_by_yt = True
                if not preceded_by_yt and not followed_by_yt:
                    replacements.append(idx)
            pos = idx + 1

        # Apply replacements in reverse order
        for idx in reversed(replacements):
            # Insert 'h' before 'w' at position idx
            entry = cmap[idx]
            r_elem = entry[1]
            ci = entry[2]
            txt = get_run_text(r_elem)
            new_txt = txt[:ci] + 'h' + txt[ci:]
            set_run_text(r_elem, new_txt)
            stats['why_to_hwhy'] += 1

        # --- Task 5: Replace Towrahʿs with Towrah's in Times New Roman ---
        all_runs = get_all_runs(para)
        cmap = build_cmap(all_runs)
        full_text = ''.join(c[0] for c in cmap)

        target = 'Towrah\u02BFs'
        pos = 0
        while True:
            idx = full_text.find(target, pos)
            if idx < 0:
                break
            # Find the ʿ character (at idx+6) and 's' (at idx+7)
            half_ring_entry = cmap[idx + 6]  # ʿ
            s_entry = cmap[idx + 7]  # s

            # Replace ʿ with ' (straight apostrophe)
            hr_r = half_ring_entry[1]
            hr_ci = half_ring_entry[2]
            hr_txt = get_run_text(hr_r)
            new_txt = hr_txt[:hr_ci] + "'" + hr_txt[hr_ci+1:]
            set_run_text(hr_r, new_txt)

            # Set font to Times New Roman for the ʿ->apostrophe run
            set_run_font(hr_r, TNR_FONT)

            # Set font to TNR for the 's' if it's in YT font
            s_r = s_entry[1]
            if get_run_font(s_r).lower() == 'yada towrah':
                set_run_font(s_r, TNR_FONT)

            stats['towrah_possessive'] += 1

            # Rebuild cmap
            all_runs = get_all_runs(para)
            cmap = build_cmap(all_runs)
            full_text = ''.join(c[0] for c in cmap)
            pos = idx + 8


def process_doc(fpath, stats):
    """Process a single document."""
    doc = Document(fpath)

    # Process body paragraphs
    process_paragraphs(doc.paragraphs, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, stats)

    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                process_paragraphs(header.paragraphs, stats)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                process_paragraphs(footer.paragraphs, stats)

    doc.save(fpath)


def main():
    files = sorted(glob.glob(os.path.join(DOCS_DIR, 'YY-*.docx')))
    files = [f for f in files if '_backup' not in f and '_updated' not in f]
    print(f'Found {len(files)} documents to process')

    totals = {
        'yy_break_comma': 0,
        'why_to_hwhy': 0,
        'double_comma': 0,
        'pua_f067': 0,
        'towrah_possessive': 0,
    }

    for i, fpath in enumerate(files):
        fname = os.path.basename(fpath)
        stats = {k: 0 for k in totals}
        print(f'  [{i+1}/{len(files)}] {fname}', end='', flush=True)
        process_doc(fpath, stats)
        changes = {k: v for k, v in stats.items() if v > 0}
        if changes:
            print(f'  {changes}')
        else:
            print('  (no changes)')
        for k in totals:
            totals[k] += stats[k]

    print()
    print('=' * 60)
    print('SUMMARY')
    print('=' * 60)
    print(f"  yy_break comma removed:    {totals['yy_break_comma']}")
    print(f"  why -> hwhy:               {totals['why_to_hwhy']}")
    print(f"  hwhy,, -> hwhy,:           {totals['double_comma']}")
    print(f"  U+F067 -> 8:               {totals['pua_f067']}")
    print(f"  Towrahʿs -> Towrah's:      {totals['towrah_possessive']}")
    print(f"  Total changes:             {sum(totals.values())}")


if __name__ == '__main__':
    main()
