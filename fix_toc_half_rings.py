"""
Fix remaining single quotes in TOC entries and throughout all YY docs.
Looks up Hebrew transliteration words in DB to determine correct half-ring.
Skips English possessives/contractions.
"""
import psycopg2
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
from lxml import etree

# Reuse functions from fix_half_rings.py
import sys
sys.path.insert(0, r'C:\Users\Joe\Work\dev\yada\translations')
from fix_half_rings import (
    normalize_quotes, normalize_blb, strip_nikkud, get_aleph_ayin_sequence,
    yy_to_hebrew_patterns, count_quotes, HALF_RING_ALEPH, HALF_RING_AYIN,
    QUOTE_CHARS, ALEPH, AYIN
)

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')
YT_FONT = 'Yada Towrah'

# English words to skip (possessives, contractions, non-Hebrew)
ENGLISH_SKIP = {
    "allah's", "don't", "dowd's", "fool's", "god's", "heaven's",
    "islam's", "i'm", "lord's", "paul's", "satan's", "who's",
    "yahowah's", "yah's", "'bout", "shaʿuwl's",
}


def is_english(word):
    """Check if word is English (possessive/contraction) not Hebrew transliteration."""
    w = word.lower().replace(HALF_RING_ALEPH, "'").replace(HALF_RING_AYIN, "'")
    for q in QUOTE_CHARS:
        w = w.replace(q, "'")
    return w in ENGLISH_SKIP


def lookup_word(word, spelling_map, blb_map, hebrew_lookup):
    """Look up a word and return the fixed version with correct half-rings."""
    if is_english(word):
        return None

    n_quotes = count_quotes(word)
    if n_quotes == 0:
        return None

    normalized = normalize_quotes(word)
    hebrew = None

    # Strategy 1: YY spelling match
    if normalized in spelling_map:
        hebrew = spelling_map[normalized]

    # Strategy 1b: hyphen parts
    if not hebrew and '-' in normalized:
        for part in normalized.split('-'):
            if count_quotes(part) > 0 and part in spelling_map:
                hebrew = spelling_map[part]
                break

    # Strategy 2: BLB translit match
    if not hebrew:
        blb_key = normalize_blb(word)
        if blb_key in blb_map:
            hebrew = blb_map[blb_key]

    # Strategy 2b: BLB hyphen parts
    if not hebrew and '-' in normalized:
        for part in normalized.split('-'):
            if count_quotes(part) > 0 and part in blb_map:
                hebrew = blb_map[part]
                break

    if hebrew:
        seq = get_aleph_ayin_sequence(hebrew)
        if len(seq) >= n_quotes:
            return replace_quotes_with_seq(word, seq)
        elif len(seq) > 0:
            return replace_quotes_with_seq(word, seq)

    # Strategy 3: Hebrew consonant pattern matching
    patterns = yy_to_hebrew_patterns(normalized)
    for heb_pattern, aleph_ayin_seq in patterns:
        if heb_pattern in hebrew_lookup:
            db_seq = hebrew_lookup[heb_pattern]
            if len(db_seq) >= n_quotes:
                return replace_quotes_with_seq(word, db_seq)

    # Strategy 3b: pattern match on hyphen parts
    if '-' in normalized:
        for part in normalized.split('-'):
            if count_quotes(part) > 0:
                part_patterns = yy_to_hebrew_patterns(part)
                for heb_pattern, _ in part_patterns:
                    if heb_pattern in hebrew_lookup:
                        db_seq = hebrew_lookup[heb_pattern]
                        part_quotes = count_quotes(part)
                        if len(db_seq) >= part_quotes:
                            return replace_quotes_with_seq(word, db_seq)

    # Strategy 4: Heuristic
    replacements = []
    in_word = normalized
    for i, ch in enumerate(in_word):
        if ch == "'":
            if i == 0:
                replacements.append(HALF_RING_ALEPH)  # Initial: usually Aleph
            elif i == len(in_word) - 1:
                replacements.append(HALF_RING_ALEPH)  # Final: default Aleph
            else:
                replacements.append(HALF_RING_AYIN)   # Internal: usually Ayin
    return replace_quotes_with_seq(word, replacements)


def replace_quotes_with_seq(word, seq):
    """Replace quote chars in word with the half-ring sequence."""
    result = ''
    idx = 0
    for ch in word:
        if ch in QUOTE_CHARS or ch == HALF_RING_ALEPH or ch == HALF_RING_AYIN:
            if idx < len(seq):
                result += seq[idx]
                idx += 1
            else:
                result += ch
        else:
            result += ch
    return result


def is_quote(ch):
    return ch in QUOTE_CHARS or ch == HALF_RING_ALEPH or ch == HALF_RING_AYIN


def build_pattern(word):
    """Build regex matching word with any quote variant."""
    QUOTE_CLASS = "['\u2018\u2019`\u00B4\u02BC\u02BE\u02BF]"
    parts = []
    for ch in word:
        if is_quote(ch):
            parts.append(QUOTE_CLASS)
        else:
            parts.append(re.escape(ch))
    return re.compile(r'(?<![a-zA-Z])' + ''.join(parts) + r'(?![a-zA-Z])')


def get_char_replacements(original, fixed):
    return {i: f for i, (o, f) in enumerate(zip(original, fixed)) if o != f}


def process_paragraph(para, patterns):
    """Apply replacements. Returns count."""
    if not para.runs:
        return 0
    text = ''
    cmap = []
    for ri, run in enumerate(para.runs):
        for ci, ch in enumerate(run.text):
            text += ch
            cmap.append((ri, ci))
    if not text or not any(is_quote(ch) for ch in text):
        return 0

    count = 0
    used = set()
    for pat, char_reps in patterns:
        for m in pat.finditer(text):
            s, e = m.start(), m.end()
            if any(i in used for i in range(s, e)):
                continue
            for offset, new_ch in char_reps.items():
                pos = s + offset
                if pos < len(cmap):
                    ri, ci = cmap[pos]
                    t = list(para.runs[ri].text)
                    t[ci] = new_ch
                    para.runs[ri].text = ''.join(t)
            for i in range(s, e):
                used.add(i)
            count += 1
    return count


def apply_yt_font_to_half_rings(para):
    """Set Yada Towrah font on half-ring characters, splitting runs as needed."""
    if not para.runs:
        return 0
    full_text = ''.join(r.text for r in para.runs)
    if HALF_RING_ALEPH not in full_text and HALF_RING_AYIN not in full_text:
        return 0

    hr_count = 0
    runs_to_process = list(para.runs)
    for run in runs_to_process:
        text = run.text
        if not text:
            continue
        has_hr = any(ch in (HALF_RING_ALEPH, HALF_RING_AYIN) for ch in text)
        if not has_hr:
            continue

        all_hr = all(ch in (HALF_RING_ALEPH, HALF_RING_AYIN) for ch in text)
        if all_hr:
            run.font.name = YT_FONT
            hr_count += len(text)
            continue

        # Split run into segments
        segments = []
        current_text = ''
        current_is_hr = text[0] in (HALF_RING_ALEPH, HALF_RING_AYIN)
        for ch in text:
            ch_is_hr = ch in (HALF_RING_ALEPH, HALF_RING_AYIN)
            if ch_is_hr == current_is_hr:
                current_text += ch
            else:
                segments.append((current_text, current_is_hr))
                current_text = ch
                current_is_hr = ch_is_hr
        segments.append((current_text, current_is_hr))

        if len(segments) <= 1:
            if current_is_hr:
                run.font.name = YT_FONT
                hr_count += len(text)
            continue

        run.text = segments[0][0]
        if segments[0][1]:
            run.font.name = YT_FONT
            hr_count += len(segments[0][0])

        run_element = run._element
        for seg_text, seg_is_hr in segments[1:]:
            new_run_elem = deepcopy(run_element)
            t_elem = new_run_elem.find(qn('w:t'))
            if t_elem is None:
                t_elem = etree.SubElement(new_run_elem, qn('w:t'))
            t_elem.text = seg_text
            t_elem.set(qn('xml:space'), 'preserve')

            if seg_is_hr:
                rPr = new_run_elem.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(new_run_elem, qn('w:rPr'))
                    new_run_elem.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                rFonts.set(qn('w:ascii'), YT_FONT)
                rFonts.set(qn('w:hAnsi'), YT_FONT)
                rFonts.set(qn('w:cs'), YT_FONT)
                hr_count += len(seg_text)

            run_element.addnext(new_run_elem)
            run_element = new_run_elem

    return hr_count


def main():
    conn = psycopg2.connect(
        host='localhost', port=5433, dbname='yada',
        user='postgres', password='yada_password'
    )
    cur = conn.cursor()

    # Load DB maps (same as fix_half_rings.py)
    cur.execute('''
        SELECT ws.word_translit_text, w.word_hebrew
        FROM yy_word_translit ws
        JOIN yy_word w ON ws.word_id = w.word_key
        WHERE w.word_hebrew IS NOT NULL AND w.word_hebrew != ''
    ''')
    spelling_map = {}
    for spelling, hebrew in cur.fetchall():
        key = normalize_quotes(spelling.strip())
        if key not in spelling_map:
            spelling_map[key] = hebrew.strip()

    cur.execute('''
        SELECT word_import_translit, word_import_hebrew
        FROM yy_word_import
        WHERE word_import_translit IS NOT NULL AND word_import_hebrew IS NOT NULL
    ''')
    blb_map = {}
    for translit, hebrew in cur.fetchall():
        key = normalize_blb(translit.strip())
        if key not in blb_map:
            blb_map[key] = hebrew.strip()

    hebrew_lookup = {}
    cur.execute('''
        SELECT word_import_hebrew FROM yy_word_import
        WHERE word_import_hebrew IS NOT NULL AND word_import_hebrew != ''
    ''')
    for (heb_raw,) in cur.fetchall():
        heb = strip_nikkud(heb_raw.strip())
        if heb and heb not in hebrew_lookup:
            seq = get_aleph_ayin_sequence(heb_raw.strip())
            if seq:
                hebrew_lookup[heb] = seq

    conn.close()
    print(f'DB: {len(spelling_map)} spellings, {len(blb_map)} BLB, {len(hebrew_lookup)} Hebrew patterns')

    # Collect all remaining words with quotes from TOC across all files
    files = sorted(DOCS_DIR.glob('YY*.docx'))
    files = [f for f in files if '_backup' not in f.name and '_updated' not in f.name
             and '_halfring' not in f.name]

    # Gather unique words with quotes from TOC + body text
    toc_words = set()
    for fp in files:
        doc = Document(str(fp))
        for para in doc.paragraphs[:80]:
            has_field = bool(para._element.findall('.//' + qn('w:instrText')))
            style_name = para.style.name if para.style else ''
            is_toc = has_field or 'toc' in style_name.lower()
            if not is_toc:
                continue
            words = para.text.split()
            for w in words:
                clean = w.strip('\t0123456789 ')
                if clean and any(ch in QUOTE_CHARS for ch in clean):
                    toc_words.add(clean)

    print(f'\nTOC words with remaining quotes: {len(toc_words)}')

    # Look up each word and build replacement map
    replacement_map = {}
    for word in sorted(toc_words):
        fixed = lookup_word(word, spelling_map, blb_map, hebrew_lookup)
        if fixed and fixed != word:
            replacement_map[word] = fixed
            print(f'  {word} -> {fixed}')
        elif is_english(word):
            print(f'  {word} -> [skip: English]')
        else:
            print(f'  {word} -> [no change]')

    if not replacement_map:
        print('\nNo replacements needed.')
        return

    # Build patterns (longest first)
    patterns = []
    for original, fixed in sorted(replacement_map.items(), key=lambda x: -len(x[0])):
        pat = build_pattern(original)
        reps = get_char_replacements(original, fixed)
        if reps:
            patterns.append((pat, reps))
    print(f'\nApplying {len(patterns)} patterns to {len(files)} files...\n')

    grand_rep = 0
    grand_font = 0
    for i, fp in enumerate(files):
        print(f'  [{i+1}/{len(files)}] {fp.name}', end='', flush=True)
        try:
            doc = Document(str(fp))
            count = 0

            # Process ALL paragraphs (body includes TOC)
            for para in doc.paragraphs:
                count += process_paragraph(para, patterns)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            count += process_paragraph(para, patterns)

            # Headers
            for section in doc.sections:
                try:
                    if not section.header.is_linked_to_previous:
                        for para in section.header.paragraphs:
                            count += process_paragraph(para, patterns)
                except Exception:
                    pass

            # Apply YT font to any new half-rings
            font_count = 0
            for para in doc.paragraphs:
                font_count += apply_yt_font_to_half_rings(para)
            for section in doc.sections:
                try:
                    if not section.header.is_linked_to_previous:
                        for para in section.header.paragraphs:
                            font_count += apply_yt_font_to_half_rings(para)
                except Exception:
                    pass

            doc.save(str(fp))
            print(f' -> {count} replacements, {font_count} font fixes')
            grand_rep += count
            grand_font += font_count
        except Exception as e:
            print(f' ERROR: {e}')

    print(f'\nDone! {grand_rep} replacements, {grand_font} font fixes')


if __name__ == '__main__':
    main()
