"""Check if apostrophes and names are split across runs."""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

LEFT_SINGLE = chr(0x2018)  # '

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s01v02-An Intro to God-Mitswah-Instructions.docx')

# Check para 596 which has 'Abram
para = doc.paragraphs[596]
print(f"Para 596 full text: {para.text[:100]}...\n")

print(f"Number of runs: {len(para.runs)}\n")

for i, run in enumerate(para.runs):
    if run.text and (LEFT_SINGLE in run.text or 'Abram' in run.text or 'Abraham' in run.text):
        print(f"Run {i}: {repr(run.text)}")

# Check para 990 which has 'Abraham
print("\n" + "="*80 + "\n")
para = doc.paragraphs[990]
print(f"Para 990 full text: {para.text[:100]}...\n")

print(f"Number of runs: {len(para.runs)}\n")

for i, run in enumerate(para.runs):
    if run.text and (LEFT_SINGLE in run.text or 'Abraham' in run.text):
        print(f"Run {i}: {repr(run.text)}")
