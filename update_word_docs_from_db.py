"""
Update YY Word documents with Hebrew modifiers from database.
Replaces italic Hebrew words that have quotes/apostrophes with their
database versions that use proper modifiers (ʾ and ʿ), applying
"Yada Towrah" font to the modifier characters.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from dotenv import load_dotenv
load_dotenv(r'C:\Users\Joe\Work\dev\yada\translations\.env')
import psycopg2

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

# Character definitions
ALEPH = chr(0x02BE)  # ʾ (modifier letter right half ring)
AYIN = chr(0x02BF)   # ʿ (modifier letter left half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

def normalize_quotes(s):
    """Normalize all quote variations to straight apostrophe."""
    return s.replace(LEFT_CURLY, "'").replace(RIGHT_CURLY, "'").replace(chr(0x2032), "'")

def has_modifier(s):
    """Check if string contains Hebrew modifiers."""
    return ALEPH in s or AYIN in s

def has_quote(s):
    """Check if string contains any quote/apostrophe."""
    return LEFT_CURLY in s or RIGHT_CURLY in s or STRAIGHT_QUOTE in s

def apply_capitalization(source, target):
    """
    Apply capitalization pattern from source to target string.
    Preserves:
    - All caps (BA'AL -> BAʿAL)
    - Title case (Ba'al -> Baʿal)
    - First cap (Ba'al -> Baʿal)
    - Lower case (ba'al -> baʿal)
    """
    if not source or not target:
        return target

    # Check if source is all uppercase (ignoring non-letter chars)
    source_letters = [c for c in source if c.isalpha()]
    if source_letters and all(c.isupper() for c in source_letters):
        # Apply all caps to target
        return target.upper()

    # Apply character-by-character capitalization
    result = []
    source_idx = 0
    for target_char in target:
        # Find next source character (skip non-matching chars like modifiers)
        while source_idx < len(source) and not source[source_idx].isalpha():
            source_idx += 1

        if source_idx < len(source) and target_char.isalpha():
            if source[source_idx].isupper():
                result.append(target_char.upper())
            else:
                result.append(target_char.lower())
            source_idx += 1
        else:
            result.append(target_char)

    return ''.join(result)

# Load word mappings from database
print("Loading Hebrew word mappings from database...")
conn = psycopg2.connect(
    host=os.getenv('POSTGRES_HOST', 'localhost'),
    port=os.getenv('POSTGRES_PORT', '5432'),
    user=os.getenv('POSTGRES_USER', 'postgres'),
    password=os.getenv('POSTGRES_PASSWORD'),
    dbname=os.getenv('POSTGRES_DB', 'yada')
)
cur = conn.cursor()

# Build mapping: normalized_form_without_modifiers -> form_with_modifiers
# This handles both yy_word.word_translit and yy_word_translit.word_translit_text
word_mappings = {}  # normalized (no modifiers) -> with modifiers

# Get words with modifiers from yy_word
cur.execute("""
    SELECT DISTINCT word_translit
    FROM yy_word
    WHERE word_translit LIKE %s OR word_translit LIKE %s
""", (f'%{ALEPH}%', f'%{AYIN}%'))

for (word_translit,) in cur.fetchall():
    if has_modifier(word_translit):
        # Create normalized version (replace modifiers with straight apostrophe)
        normalized = word_translit.replace(ALEPH, "'").replace(AYIN, "'")
        word_mappings[normalized.lower()] = word_translit

# Get spellings with modifiers from yy_word_translit
cur.execute("""
    SELECT DISTINCT word_translit_text
    FROM yy_word_translit
    WHERE word_translit_text LIKE %s OR word_translit_text LIKE %s
""", (f'%{ALEPH}%', f'%{AYIN}%'))

for (spelling,) in cur.fetchall():
    if has_modifier(spelling):
        normalized = spelling.replace(ALEPH, "'").replace(AYIN, "'")
        key = normalized.lower()
        if key not in word_mappings:
            word_mappings[key] = spelling

cur.close()
conn.close()

print(f"Loaded {len(word_mappings)} Hebrew words with modifiers")
print(f"Sample mappings:")
for i, (old, new) in enumerate(list(word_mappings.items())[:10]):
    print(f"  {old} -> {new}")
print()

def find_word_replacement(text):
    """
    Find if text (or any word in text) should be replaced with modifier version.
    Preserves original capitalization.
    Returns (original_text, replacement_text) or None if no replacement needed.
    """
    # Normalize quotes in the input text
    normalized = normalize_quotes(text).lower()

    # Check for exact match first
    if normalized in word_mappings:
        base_replacement = word_mappings[normalized]
        # Apply capitalization from original text
        replacement = apply_capitalization(text, base_replacement)
        return (text, replacement)

    # Check if any token in the text matches
    # Split on whitespace and common punctuation
    tokens = re.split(r'[\s,;:!?()\[\]{}]+', text)
    for token in tokens:
        if not token:
            continue
        # Remove trailing/leading punctuation
        cleaned = token.strip('.,;:!?()-/[]{}"\u2013\u2014\u2015')
        if cleaned:
            normalized_token = normalize_quotes(cleaned).lower()
            if normalized_token in word_mappings:
                # Get base replacement from database
                base_replacement = word_mappings[normalized_token]
                # Apply capitalization from the original cleaned token
                replacement = apply_capitalization(cleaned, base_replacement)

                # Replace in the original text, preserving position
                new_text = text

                # Build variations to try (with different quote styles)
                token_variations = [
                    token,
                    token.replace("'", LEFT_CURLY),
                    token.replace("'", RIGHT_CURLY),
                ]

                for variation in token_variations:
                    if variation in new_text and variation != replacement:
                        # For the replacement, we need to preserve any punctuation
                        # that was stripped from the cleaned version
                        if token != cleaned:
                            # Token has extra punctuation - preserve it
                            # Find what was stripped
                            prefix = token[:len(token) - len(token.lstrip('.,;:!?()-/[]{}"\u2013\u2014\u2015'))]
                            suffix = token[len(token.rstrip('.,;:!?()-/[]{}"\u2013\u2014\u2015')):]
                            full_replacement = prefix + replacement + suffix
                        else:
                            full_replacement = replacement

                        new_text = new_text.replace(variation, full_replacement)
                        if new_text != text:
                            return (text, new_text)

    return None

def split_run_at_modifiers(run, new_text):
    """
    Replace run text with new_text, applying "Yada Towrah" font to modifier chars.
    Creates multiple runs if needed to handle font changes.
    """
    if new_text == run.text or not new_text:
        return

    # Save original formatting
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    # Get the parent paragraph and run position
    para = run._element.getparent()
    run_idx = list(para).index(run._element)

    # Clear the original run
    run.text = ""

    # Build new runs
    def create_run(text, use_yada_font=False):
        """Helper to create a run with specified formatting."""
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        # Italic
        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Font - "Yada Towrah" for modifiers, "Times New Roman" for regular text
        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else "Times New Roman"
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        # Font size
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        # Color
        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    # Split text into segments (regular text vs modifier chars)
    i = 0
    buffer = []
    current_run_idx = run_idx

    def flush_buffer(is_modifier=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_modifier)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    # Process each character
    while i < len(new_text):
        char = new_text[i]
        if char in (ALEPH, AYIN):
            # Flush any pending regular text
            flush_buffer(is_modifier=False)
            # Add modifier and flush
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)
        i += 1

    # Flush remaining
    flush_buffer(is_modifier=False)

def process_paragraph(para, stats):
    """Process a paragraph, replacing Hebrew words in italic runs."""
    for run in list(para.runs):
        if not run.text or not run.italic:
            continue

        # Check if this run needs replacement
        result = find_word_replacement(run.text)
        if result:
            original, replacement = result
            if original != replacement and has_modifier(replacement):
                split_run_at_modifiers(run, replacement)
                stats['replacements'] += 1
                if stats['verbose'] and stats['replacements'] <= 20:
                    print(f"    Replaced: {original} -> {replacement}")

def process_document(docx_path, verbose=False):
    """Process a single Word document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: Could not open document: {e}")
        return 0

    stats = {'replacements': 0, 'verbose': verbose}

    # Process paragraphs
    for para in doc.paragraphs:
        process_paragraph(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    replacements = stats['replacements']

    if replacements > 0:
        # Save with _updated_modifiers suffix
        base, ext = os.path.splitext(docx_path)
        # Remove old suffixes if present
        for old_suffix in ['_updated', '_updated_fonts', '_test_updated', '_updated_modifiers']:
            if base.endswith(old_suffix):
                base = base[:-len(old_suffix)]
        output_path = f"{base}_updated_modifiers{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ {replacements} replacements, saved to {os.path.basename(output_path)} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: Could not save document: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return replacements

# Main execution
docs_dir = r'C:\users\joe\work\dev\yada\docs'
if not os.path.exists(docs_dir):
    print(f"ERROR: Documents directory not found: {docs_dir}")
    sys.exit(1)

print(f"\nSearching for YY documents in: {docs_dir}")

# Find all original YY documents (not previously updated versions)
docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-')
    and f.endswith('.docx')
    and not any(suffix in f for suffix in ['_updated', '_test_updated'])
    and not f.startswith('~$')
]

print(f"Found {len(docx_files)} YY documents to process\n")
print("=" * 80)

total_files_updated = 0
total_replacements = 0

for docx_path in docx_files:
    replacements = process_document(docx_path, verbose=False)
    if replacements > 0:
        total_files_updated += 1
        total_replacements += replacements

print("\n" + "=" * 80)
print("WORD DOCUMENT UPDATE COMPLETE")
print("=" * 80)
print(f"  Files processed: {len(docx_files)}")
print(f"  Files updated: {total_files_updated}")
print(f"  Total replacements: {total_replacements}")
print(f"\nUpdated files saved with '_updated_modifiers.docx' suffix in: {docs_dir}")
