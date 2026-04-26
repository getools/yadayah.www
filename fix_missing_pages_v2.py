#!/usr/bin/env python3
"""
Fix translations with missing page numbers using a page-by-page approach.

Instead of iterating all paragraphs (slow) or text-based Find (fails for Unicode),
this navigates page by page via COM GoTo, collects text from each page, then
matches translation paragraph text to pages on the Python side.
"""

import sys
import os
import re
import logging
import tempfile
import subprocess
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from parse_word_translations import (
    extract_translations_from_doc,
    setup_logging,
    DEFAULT_DIRECTORY,
)
from docx import Document

try:
    import psycopg2
except ImportError:
    print("ERROR: psycopg2 not installed")
    sys.exit(1)

try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: python-dotenv not installed")
    sys.exit(1)


def get_page_texts(doc_path: str, timeout: int = 300) -> dict:
    """
    Navigate page-by-page via COM GoTo and collect text from each page.
    Returns {footer_page_number: [text_chunk1, text_chunk2, ...], ...}
    where each text chunk is the cleaned text from one visit to that page.
    """
    out = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8')
    out.close()

    # VBScript navigates page by page, outputting footer_page|text for each page
    vbs_content = f'''
Dim word, fso, outputFile
Set fso = CreateObject("Scripting.FileSystemObject")
Set word = CreateObject("Word.Application")
word.Visible = False
word.DisplayAlerts = 0

Set outputFile = fso.CreateTextFile("{out.name}", True, True)

WScript.StdErr.WriteLine "Opening document..."
Dim doc
Set doc = word.Documents.Open("{doc_path}", , True)

Dim totalPages
totalPages = doc.ComputeStatistics(2)
WScript.StdErr.WriteLine "Opened, pages: " & totalPages

Dim p, footerPage, pageText
Dim rng, rngEnd

For p = 1 To totalPages
    On Error Resume Next

    Set rng = doc.GoTo(1, 1, p)

    If p < totalPages Then
        Set rngEnd = doc.GoTo(1, 1, p + 1)
        rng.End = rngEnd.Start
    Else
        rng.End = doc.Content.End
    End If

    footerPage = rng.Information(1)
    pageText = rng.Text

    If Len(pageText) > 4000 Then
        pageText = Left(pageText, 4000)
    End If

    pageText = Replace(pageText, vbCr, " ")
    pageText = Replace(pageText, vbLf, " ")
    pageText = Replace(pageText, vbTab, " ")
    pageText = Replace(pageText, "|", " ")

    If Err.Number = 0 Then
        outputFile.WriteLine footerPage & "|" & pageText
    End If
    On Error GoTo 0
Next

doc.Close False
outputFile.Close
word.Quit
WScript.StdErr.WriteLine "Done - " & totalPages & " pages"
'''

    vbs_file = tempfile.NamedTemporaryFile(mode='w', suffix='.vbs', delete=False, encoding='utf-8')
    vbs_file.write(vbs_content)
    vbs_file.close()

    try:
        result = subprocess.run(
            ["cscript", "//NoLogo", vbs_file.name],
            capture_output=True, text=True, timeout=timeout
        )

        if result.stderr:
            for line in result.stderr.strip().split('\n'):
                if line.strip():
                    logging.info(f"  VBS: {line.strip()}")

        if result.returncode != 0:
            logging.warning(f"  VBS exited with code {result.returncode}")
            if result.stdout:
                logging.warning(f"  VBS stdout: {result.stdout[:200]}")

        page_texts = {}
        if os.path.exists(out.name):
            with open(out.name, 'r', encoding='utf-16', errors='replace') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('\ufeff'):
                        line = line.lstrip('\ufeff').strip()
                    if not line:
                        continue
                    pipe = line.find('|')
                    if pipe > 0:
                        try:
                            page = int(line[:pipe])
                            text = line[pipe+1:]
                            if page not in page_texts:
                                page_texts[page] = []
                            page_texts[page].append(text)
                        except ValueError:
                            pass

        return page_texts

    except subprocess.TimeoutExpired:
        logging.warning(f"  COM timed out after {timeout}s")
        try:
            subprocess.run(["powershell", "-Command",
                "Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue"],
                capture_output=True, timeout=10)
        except:
            pass
        return {}
    except Exception as e:
        logging.warning(f"  COM failed: {e}")
        return {}
    finally:
        for fp in [vbs_file.name, out.name]:
            try:
                os.unlink(fp)
            except:
                pass


def norm(text: str) -> str:
    """Normalize text for comparison: collapse whitespace, strip."""
    return re.sub(r'\s+', ' ', text).strip()


def find_page_for_text(para_text: str, page_texts: dict) -> int:
    """
    Find which page contains the given paragraph text.
    Uses substring matching with normalized text.
    Returns page number or None.
    """
    normed = norm(para_text)
    if len(normed) < 10:
        return None

    # Try progressively shorter snippets for matching
    for snippet_len in [80, 50, 30]:
        if len(normed) < snippet_len:
            snippet = normed
        else:
            snippet = normed[:snippet_len]

        for page, texts in page_texts.items():
            for page_text in texts:
                if snippet in norm(page_text):
                    return page

    return None


def main():
    setup_logging(verbose=False)
    load_dotenv()

    logging.info("Fix Missing Page Numbers v2 (page-by-page approach)")
    logging.info("=" * 60)

    conn = psycopg2.connect(
        host=os.getenv('POSTGRES_HOST', 'localhost'),
        port=os.getenv('POSTGRES_PORT', '5432'),
        user=os.getenv('POSTGRES_USER', 'postgres'),
        password=os.getenv('POSTGRES_PASSWORD'),
        database=os.getenv('POSTGRES_DB', 'yada'),
    )
    cur = conn.cursor()

    # Get books with missing pages
    cur.execute("""
        SELECT translation_book, COUNT(*) as missing
        FROM yy_cite_translation WHERE translation_page IS NULL
        GROUP BY translation_book ORDER BY missing DESC
    """)
    missing_books = cur.fetchall()
    total_missing = sum(row[1] for row in missing_books)
    logging.info(f"{total_missing} missing pages across {len(missing_books)} books\n")

    doc_dir = Path(DEFAULT_DIRECTORY)
    total_fixed = 0

    for book_name, missing_count in missing_books:
        doc_path = doc_dir / f"{book_name}.docx"
        if not doc_path.exists():
            logging.warning(f"  File not found: {doc_path}")
            continue

        file_size_mb = doc_path.stat().st_size / (1024 * 1024)
        logging.info(f"{book_name} ({file_size_mb:.1f}MB, {missing_count} missing)")

        # Step 1: Extract translations to get para_idx for missing ones
        translations = extract_translations_from_doc(doc_path)

        # Step 2: Get missing translation IDs from DB
        cur.execute("""
            SELECT translation_id, translation_text_word
            FROM yy_cite_translation
            WHERE translation_book = %s AND translation_page IS NULL
        """, (book_name,))
        missing_db = {row[0]: row[1] for row in cur.fetchall()}

        # Step 3: Map missing translations to their paragraph indices
        missing_para_indices = {}  # translation_id -> para_idx
        for t in translations:
            text_word = t.get('text_word', '')
            para_idx = t.get('_para_idx')
            for tid, db_text in missing_db.items():
                if tid not in missing_para_indices and db_text == text_word:
                    missing_para_indices[tid] = para_idx
                    break

        if not missing_para_indices:
            logging.warning(f"  Could not match any missing translations to paragraphs")
            continue

        # Step 4: Get python-docx paragraph texts for the needed indices
        doc = Document(str(doc_path))
        pdx_texts = {}  # para_idx -> raw text
        for tid, pidx in missing_para_indices.items():
            if pidx < len(doc.paragraphs):
                pdx_texts[pidx] = doc.paragraphs[pidx].text

        # Step 5: Kill Word, then get page texts
        subprocess.run(["powershell", "-Command",
            "Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue"],
            capture_output=True, timeout=10)

        timeout = int(120 + file_size_mb * 30)
        page_texts = {}

        # Try clean copy first (faster and more reliable for all sizes)
        for attempt_label in ["clean copy", "original"]:
            if attempt_label == "clean copy":
                logging.info(f"  Creating clean copy...")
                try:
                    clean_dir = tempfile.mkdtemp()
                    clean_path = os.path.join(clean_dir, "clean.docx")
                    Document(str(doc_path)).save(clean_path)
                    attempt_path = clean_path
                except Exception as e:
                    logging.warning(f"  Clean copy failed: {e}")
                    continue
            else:
                attempt_path = str(doc_path)

            logging.info(f"  COM page-by-page ({attempt_label}, timeout {timeout}s)...")
            page_texts = get_page_texts(attempt_path, timeout=timeout)

            if attempt_label == "clean copy":
                try:
                    os.unlink(attempt_path)
                    os.rmdir(os.path.dirname(attempt_path))
                except:
                    pass

            if page_texts:
                logging.info(f"  Got {len(page_texts)} pages from COM")
                break
            else:
                logging.warning(f"  COM failed ({attempt_label})")

        if not page_texts:
            logging.warning(f"  Could not get page texts for {book_name}")
            continue

        # Step 6: Match each missing translation's paragraph text to a page
        fixed = 0
        for tid, pidx in missing_para_indices.items():
            para_text = pdx_texts.get(pidx, '')
            if not para_text:
                continue

            page = find_page_for_text(para_text, page_texts)
            if page is not None:
                cur.execute("""
                    UPDATE yy_cite_translation SET translation_page = %s
                    WHERE translation_id = %s AND translation_page IS NULL
                """, (page, tid))
                if cur.rowcount > 0:
                    fixed += 1

        conn.commit()
        total_fixed += fixed
        logging.info(f"  Fixed {fixed}/{missing_count}")

    # Final check
    cur.execute("SELECT COUNT(*) FROM yy_cite_translation WHERE translation_page IS NULL")
    still_missing = cur.fetchone()[0]

    logging.info(f"\n{'=' * 60}")
    logging.info(f"SUMMARY")
    logging.info(f"Total fixed: {total_fixed}/{total_missing}")
    logging.info(f"Still missing: {still_missing}")

    cur.close()
    conn.close()
    logging.info("Done.")


if __name__ == "__main__":
    main()
