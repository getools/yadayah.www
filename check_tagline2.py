import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')

# Find the paragraph mentioned
target_text = 'The fourth Miqra'
for i, para in enumerate(doc.paragraphs):
    if target_text in para.text:
        print(f'Found target at index {i}')
        print(f'  Alignment: {para.alignment}')
        print(f'  First-line indent: {para.paragraph_format.first_line_indent}')
        print()

        # Check previous 5 paragraphs
        start = max(0, i - 5)
        for j in range(start, i):
            p = doc.paragraphs[j]
            text = p.text.strip()
            if text:
                print(f'\nParagraph {j}:')
                print(f'  Text: {text[:100]}{"..." if len(text) > 100 else ""}')
                print(f'  Ends with ...: {text.endswith("...")}')
                print(f'  Ends with …: {text.endswith("…")}')
                print(f'  Alignment: {p.alignment}')

                # Check italic
                italic_chars = 0
                total_chars = 0
                for r in p.runs:
                    if r.text.strip():
                        total_chars += len(r.text.strip())
                        if r.italic:
                            italic_chars += len(r.text.strip())

                if total_chars > 0:
                    print(f'  Italic: {100*italic_chars/total_chars:.1f}%')

        break
