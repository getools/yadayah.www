"""
Apply Yada Towrah font to BOTH ʿ (AYIN) and ʾ (ALEPH) modifier characters.
This ensures every modifier has the proper font, regardless of current state.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

AYIN = chr(0x02BF)   # ʿ (left half ring)
ALEPH = chr(0x02BE)  # ʾ (right half ring)

def process_document(docx_path):
    """Apply Yada Towrah font to all ʿ and ʾ modifiers in a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process all paragraphs (body, tables, headers, footers)
    all_paras = list(doc.paragraphs)

    # Add table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    # Add header/footer paragraphs
    for section in doc.sections:
        if section.header:
            all_paras.extend(section.header.paragraphs)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_paras.extend(cell.paragraphs)

        if section.footer:
            all_paras.extend(section.footer.paragraphs)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_paras.extend(cell.paragraphs)

    # Process each paragraph
    for para in all_paras:
        for element in para._element.iter():
            if element.tag == qn('w:t'):
                # Check if element contains ONLY a modifier character
                if element.text in [AYIN, ALEPH]:
                    # Get parent run
                    run_elem = element.getparent()
                    if run_elem.tag == qn('w:r'):
                        # Get or create rPr
                        rPr = run_elem.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            run_elem.insert(0, rPr)

                        # Get or create rFonts
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.append(rFonts)

                        # Check if already Yada Towrah
                        current = rFonts.get(qn('w:ascii'))
                        if current != 'Yada Towrah':
                            rFonts.set(qn('w:ascii'), 'Yada Towrah')
                            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
                            fixes += 1

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} modifiers ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - Already correct ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Applying Yada Towrah font to ALL ʿ (AYIN) and ʾ (ALEPH) modifiers...")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in sorted(docx_files):
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total ʿ + ʾ modifiers set to Yada Towrah: {total_fixes}")
