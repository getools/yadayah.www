"""
Final comprehensive update for Hebrew words in YY Word documents.
Includes corrections to previous updates (Howsha', Yahowsha' now use Ayin not Aleph).
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

# Character definitions
ALEPH = chr(0x02BE)  # ʾ (modifier letter right half ring)
AYIN = chr(0x02BF)   # ʿ (modifier letter left half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

def normalize_quotes(s):
    """Normalize all quote variations to straight apostrophe."""
    return s.replace(LEFT_CURLY, "'").replace(RIGHT_CURLY, "'").replace(chr(0x2032), "'")

# Replacement patterns - order matters, most specific first
REPLACEMENTS = [
    # Fix previous Aleph to Ayin: Howshaʾ → Howshaʿ
    (re.compile(rf'Howsha{ALEPH}(?!s)\b', re.IGNORECASE),
     lambda m: f'Howsha{AYIN}' if m.group()[0].isupper() else f'howsha{AYIN}'),

    # Fix previous Aleph to Ayin: Howshaʾ's → Howshaʿ's
    (re.compile(rf"Howsha{ALEPH}'s\b", re.IGNORECASE),
     lambda m: f"Howsha{AYIN}'s" if m.group()[0].isupper() else f"howsha{AYIN}'s"),

    # Fix previous Aleph to Ayin: Yahowshaʾ → Yahowshaʿ
    (re.compile(rf'Yahowsha{ALEPH}\b', re.IGNORECASE),
     lambda m: f'Yahowsha{AYIN}' if m.group()[0].isupper() else f'yahowsha{AYIN}'),

    # 'Ely'ab → ʾElyʿab (both upper and lower, capitalize E)
    (re.compile(r"['\u2018\u2019\u2032][Ee]ly['\u2018\u2019\u2032]ab\b", re.IGNORECASE),
     lambda m: f'{ALEPH}Ely{AYIN}ab'),

    # Zarowa's → Zarowaʾ's
    (re.compile(r"Zarowa['\u2018\u2019\u2032]s\b", re.IGNORECASE),
     lambda m: f'Zarowa{ALEPH}\'s' if m.group()[0].isupper() else f'zarowa{ALEPH}\'s'),

    # Zarowa' → Zarowaʾ
    (re.compile(r"Zarowa['\u2018\u2019\u2032]\b", re.IGNORECASE),
     lambda m: f'Zarowa{ALEPH}' if m.group()[0].isupper() else f'zarowa{ALEPH}'),

    # She'owl → Sheʿowl
    (re.compile(r"She['\u2018\u2019\u2032]owl\b", re.IGNORECASE),
     lambda m: f'She{AYIN}owl' if m.group()[0].isupper() else f'she{AYIN}owl'),

    # Shamuw'el → Shamuwʿel
    (re.compile(r"Shamuw['\u2018\u2019\u2032]el\b", re.IGNORECASE),
     lambda m: f'Shamuw{AYIN}el' if m.group()[0].isupper() else f'shamuw{AYIN}el'),

    # Howsha' → Howshaʿ (without 's)
    (re.compile(r"Howsha['\u2018\u2019\u2032](?!s)\b", re.IGNORECASE),
     lambda m: f'Howsha{AYIN}' if m.group()[0].isupper() else f'howsha{AYIN}'),

    # Howsha's → Howshaʿ's
    (re.compile(r"Howsha['\u2018\u2019\u2032]s\b", re.IGNORECASE),
     lambda m: f'Howsha{AYIN}\'s' if m.group()[0].isupper() else f'howsha{AYIN}\'s'),

    # Yirma'yah → Yirmaʿyah
    (re.compile(r"Yirma['\u2018\u2019\u2032]yah\b", re.IGNORECASE),
     lambda m: f'Yirma{AYIN}yah' if m.group()[0].isupper() else f'yirma{AYIN}yah'),

    # Mal'aky → Malʿaky
    (re.compile(r"Mal['\u2018\u2019\u2032]aky\b", re.IGNORECASE),
     lambda m: f'Mal{AYIN}aky' if m.group()[0].isupper() else f'mal{AYIN}aky'),

    # Yasha'yah → Yashaʿyah (if not already done)
    (re.compile(r"Yasha['\u2018\u2019\u2032]yah\b", re.IGNORECASE),
     lambda m: f'Yasha{AYIN}yah' if m.group()[0].isupper() else f'yasha{AYIN}yah'),

    # Yashaʿyah → Yashaʿyah (ensure correct, case insensitive)
    (re.compile(rf'Yasha{AYIN}yah\b'),
     lambda m: f'Yasha{AYIN}yah'),  # Already correct, ensure capitalization

    # Yahowsha' → Yahowshaʿ
    (re.compile(r"Yahowsha['\u2018\u2019\u2032]\b", re.IGNORECASE),
     lambda m: f'Yahowsha{AYIN}' if m.group()[0].isupper() else f'yahowsha{AYIN}'),

    # 'Ely → ʾEly
    (re.compile(r"['\u2018\u2019\u2032]Ely\b"),
     lambda m: f'{ALEPH}Ely'),

    # 'ely → ʾely
    (re.compile(r"['\u2018\u2019\u2032]ely\b"),
     lambda m: f'{ALEPH}ely'),

    # 'Azab → ʾAzab
    (re.compile(r"['\u2018\u2019\u2032]Azab\b"),
     lambda m: f'{ALEPH}Azab'),

    # 'azab → ʾazab
    (re.compile(r"['\u2018\u2019\u2032]azab\b"),
     lambda m: f'{ALEPH}azab'),

    # 'Asher → ʾAsher
    (re.compile(r"['\u2018\u2019\u2032]Asher\b"),
     lambda m: f'{ALEPH}Asher'),

    # 'asher → ʾasher
    (re.compile(r"['\u2018\u2019\u2032]asher\b"),
     lambda m: f'{ALEPH}asher'),

    # 'Any → ʾAny
    (re.compile(r"['\u2018\u2019\u2032]Any\b"),
     lambda m: f'{ALEPH}Any'),

    # 'any → ʾany
    (re.compile(r"['\u2018\u2019\u2032]any\b"),
     lambda m: f'{ALEPH}any'),

    # 'Ysh → Yesh (no modifier, just fix)
    (re.compile(r"['\u2018\u2019\u2032]Ysh\b"),
     lambda m: 'Yesh'),

    # 'ysh → yesh (no modifier, just fix)
    (re.compile(r"['\u2018\u2019\u2032]ysh\b"),
     lambda m: 'yesh'),

    # 'Ayn → ʾAyn
    (re.compile(r"['\u2018\u2019\u2032]Ayn\b"),
     lambda m: f'{ALEPH}Ayn'),

    # 'ayn → ʾayn
    (re.compile(r"['\u2018\u2019\u2032]ayn\b"),
     lambda m: f'{ALEPH}ayn'),

    # Miqra'ey → Miqraʿey
    (re.compile(r"Miqra['\u2018\u2019\u2032]ey\b", re.IGNORECASE),
     lambda m: f'Miqra{AYIN}ey' if m.group()[0].isupper() else f'miqra{AYIN}ey'),

    # 'Eythan → ʾEythan
    (re.compile(r"['\u2018\u2019\u2032]Eythan\b", re.IGNORECASE),
     lambda m: f'{ALEPH}Eythan' if 'E' in m.group() else f'{ALEPH}eythan'),
]

def apply_replacements(text):
    """Apply all replacement patterns to text."""
    original = text
    for pattern, replacement_func in REPLACEMENTS:
        text = pattern.sub(replacement_func, text)
    return text if text != original else None

def create_formatted_runs(run, new_text):
    """Create new runs with proper font formatting for modifiers."""
    # Save original formatting
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name or "Times New Roman"
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    # Get parent paragraph and run position
    para = run._element.getparent()
    run_idx = list(para).index(run._element)

    # Clear original run
    run.text = ""

    def create_run(text, use_yada_font=False):
        """Create a run with specified formatting."""
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Font
        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else font_name
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        # Font size
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        # Color
        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    # Split text at modifier characters
    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_modifier=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_modifier)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    for char in new_text:
        if char in (ALEPH, AYIN):
            flush_buffer(is_modifier=False)
            buffer.append(char)
            flush_buffer(is_modifier=True)
        else:
            buffer.append(char)

    flush_buffer(is_modifier=False)

def ensure_modifier_fonts(run):
    """Ensure existing modifiers use Yada Towrah font."""
    if not run.text:
        return False
    if ALEPH not in run.text and AYIN not in run.text:
        return False
    create_formatted_runs(run, run.text)
    return True

def process_paragraph(para, stats):
    """Process paragraph: apply replacements and ensure fonts."""
    for run in list(para.runs):
        if not run.text:
            continue

        new_text = apply_replacements(run.text)
        if new_text:
            create_formatted_runs(run, new_text)
            stats['replacements'] += 1
            if stats['verbose']:
                print(f"    {run.text} -> {new_text}")
        else:
            if ensure_modifier_fonts(run):
                stats['font_fixes'] += 1

def process_document(docx_path, verbose=False):
    """Process a single document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'replacements': 0, 'font_fixes': 0}

    stats = {'replacements': 0, 'font_fixes': 0, 'verbose': verbose}

    for para in doc.paragraphs:
        process_paragraph(para, stats)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    repl = stats['replacements']
    fonts = stats['font_fixes']

    if repl > 0 or fonts > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Replacements: {repl}, Fonts: {fonts} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'replacements': 0, 'font_fixes': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'replacements': repl, 'font_fixes': fonts}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Final Hebrew word updates...")
print()

docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-') and f.endswith('.docx')
    and '_updated' not in f and not f.startswith('~$')
]

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_repl = 0
total_fonts = 0

for path in docx_files:
    result = process_document(path)
    if result['replacements'] > 0 or result['font_fixes'] > 0:
        total_updated += 1
        total_repl += result['replacements']
        total_fonts += result['font_fixes']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Word replacements: {total_repl}")
print(f"Font corrections: {total_fonts}")
print(f"\nUpdated files: {docs_dir}/*_updated.docx")
