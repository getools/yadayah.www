import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')

print("Checking paragraphs around the target...")
print("=" * 80)

# Show paragraphs from 35-70
for i in range(35, min(75, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "None"
        print(f"\n{i}: [{style}]")
        print(f"    {text[:100]}{'...' if len(text) > 100 else ''}")

        # Check if it's a tagline
        if text.endswith('...') or text.endswith('…'):
            if para.style and hasattr(para.style, 'font') and para.style.font.italic:
                print(f"    >>> TAGLINE <<<")
