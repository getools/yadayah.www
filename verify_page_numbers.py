#!/usr/bin/env python3
"""
Verify that build_footer_page_map() XML-based page numbers match
Word COM footer page numbers (rng.Information(1)).

Picks sample paragraphs from several documents and compares.
"""

import sys
import logging
from pathlib import Path
from docx import Document

from parse_word_translations import (
    build_footer_page_map,
    get_page_numbers_for_doc_com,
    replace_unicode_chars,
)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')

DOCS_DIR = Path(r"C:\users\joe\work\dev\yada\docs")

# Pick a representative set of docs (different series, known reliable + unreliable)
TEST_DOCS = [
    "YY-s01v01-An Intro to God-Dabarym-Words",
    "YY-s02v04-Yada Yahowah-Miqra'ey-Invitations",
    "YY-s04v03-Coming Home-Dowd-Beloved",
    "YY-s05v02-Babel-Tow'ebah-Abominable",
    "YY-s06v04-Twistianity-Incredible",
    "YY-s07v01-God Damn Religion-Snake",
]

# Number of sample paragraphs to check per document
SAMPLES_PER_DOC = 20


def pick_sample_indices(page_map, content_start, last_page, total_paras, n=SAMPLES_PER_DOC):
    """Pick evenly-spaced paragraph indices across content range."""
    content_indices = [
        idx for idx, pg in sorted(page_map.items())
        if idx >= content_start and pg < last_page
    ]
    if not content_indices:
        return []
    step = max(1, len(content_indices) // n)
    return content_indices[::step][:n]


def main():
    if sys.stdout.encoding != 'utf-8':
        sys.stdout.reconfigure(encoding='utf-8')

    total_checked = 0
    total_match = 0
    total_mismatch = 0
    all_mismatches = []

    for stem in TEST_DOCS:
        doc_path = DOCS_DIR / f"{stem}.docx"
        if not doc_path.exists():
            logging.warning(f"Skipping {stem}: file not found")
            continue

        logging.info(f"\n{'='*70}")
        logging.info(f"Document: {stem}")

        # Phase 1: XML-based footer page map
        page_map, content_start, last_page = build_footer_page_map(doc_path)
        total_paras = len(page_map)
        logging.info(f"  XML: {total_paras} paragraphs, content starts at idx {content_start}, last page {last_page}")

        # Pick sample indices
        sample_indices = pick_sample_indices(page_map, content_start, last_page, total_paras)
        logging.info(f"  Sampling {len(sample_indices)} paragraphs for COM comparison")

        if not sample_indices:
            logging.warning(f"  No content paragraphs to sample")
            continue

        # Get paragraph text snippets for COM search
        doc = Document(str(doc_path))
        para_texts = {}
        for idx in sample_indices:
            if idx < len(doc.paragraphs):
                text = replace_unicode_chars(doc.paragraphs[idx].text)
                # Use first 120 chars for search (Word Find has limits)
                snippet = text[:120].strip()
                if snippet:
                    para_texts[idx] = snippet

        logging.info(f"  Found {len(para_texts)} paragraphs with searchable text")

        # Phase 2: COM-based footer page numbers
        com_page_map = get_page_numbers_for_doc_com(
            str(doc_path.absolute()), para_texts, timeout=180
        )
        logging.info(f"  COM returned {len(com_page_map)} page numbers")

        # Phase 3: Compare
        doc_match = 0
        doc_mismatch = 0
        doc_com_only = 0

        for idx in sorted(sample_indices):
            xml_page = page_map.get(idx)
            com_page = com_page_map.get(idx)

            if com_page is None:
                doc_com_only += 1
                continue

            total_checked += 1
            if xml_page == com_page:
                doc_match += 1
                total_match += 1
            else:
                doc_mismatch += 1
                total_mismatch += 1
                text_preview = para_texts.get(idx, '')[:60]
                all_mismatches.append({
                    'doc': stem,
                    'para_idx': idx,
                    'xml_page': xml_page,
                    'com_page': com_page,
                    'text': text_preview,
                })

        logging.info(f"  Results: {doc_match} match, {doc_mismatch} mismatch, {doc_com_only} COM-not-found")

    # Summary
    print(f"\n{'='*70}")
    print("VERIFICATION SUMMARY")
    print(f"{'='*70}")
    print(f"Total paragraphs checked: {total_checked}")
    print(f"  Match (XML == COM):     {total_match}")
    print(f"  Mismatch (XML != COM):  {total_mismatch}")

    if total_checked > 0:
        pct = total_match / total_checked * 100
        print(f"  Accuracy:               {pct:.1f}%")

    if all_mismatches:
        print(f"\nMISMATCHES ({len(all_mismatches)}):")
        print(f"{'Doc':<45s} {'Para':>6s} {'XML':>5s} {'COM':>5s} {'Diff':>5s}  Text")
        print("-" * 110)
        for m in all_mismatches:
            diff = (m['xml_page'] or 0) - (m['com_page'] or 0)
            print(f"{m['doc']:<45s} {m['para_idx']:>6d} {m['xml_page'] or '?':>5} {m['com_page']:>5d} {diff:>+5d}  {m['text']}")
    else:
        print("\nAll page numbers match! XML footer extraction is accurate.")


if __name__ == "__main__":
    main()
