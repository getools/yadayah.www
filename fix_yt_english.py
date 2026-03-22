"""
Fix YY Word documents: change non-Hebrew text in Yada Towrah font to Times New Roman.
Keeps single characters and valid Hebrew transliteration patterns in YT font.
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import os, glob, copy

DOCS_DIR = r'C:\Users\Joe\Work\dev\yada\docs'
YT_FONT = 'yada towrah'
TNR_FONT = 'Times New Roman'

# Valid single chars and short patterns that should stay in Yada Towrah font
HALF_RINGS = {'\u02BE', '\u02BF'}
# Hebrew transliteration short words (up to ~5 chars) that are valid YT
VALID_YT_WORDS = {
    'hwhy', 'hw', 'hy', 'wh', 'wy', 'yh', 'yw',
    'h', 'w', 'y', 'k', 'l', 'm', 'n', 'b', 'd', 'g',
    'p', 'q', 'r', 's', 't', 'x', 'z',
    'ts', 'sh', 'th', 'kh', 'ch',
}


def get_run_font(r_elem):
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi')) or ''
    return ''


def get_run_text(r_elem):
    parts = []
    for t in r_elem.findall(qn('w:t')):
        parts.append(t.text or '')
    return ''.join(parts)


def set_run_text(r_elem, new_text):
    t_elems = r_elem.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = new_text
        if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
            t_elems[0].set(qn('xml:space'), 'preserve')
        for extra in t_elems[1:]:
            r_elem.remove(extra)
    elif new_text:
        t = etree.SubElement(r_elem, qn('w:t'))
        t.text = new_text
        if new_text[0] == ' ' or new_text[-1] == ' ':
            t.set(qn('xml:space'), 'preserve')


def set_run_font_name(r_elem, font_name):
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(r_elem, qn('w:rPr'))
        r_elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def is_valid_yt(text):
    """Check if text should remain in Yada Towrah font."""
    if not text:
        return True
    stripped = text.strip()
    if not stripped:
        return True
    # Single character: always valid
    if len(stripped) <= 1:
        return True
    # Half-ring + single char or single char + half-ring
    non_hr = stripped.replace('\u02BE', '').replace('\u02BF', '')
    if len(non_hr) <= 1:
        return True
    # Known valid Hebrew transliteration patterns
    if stripped.lower() in VALID_YT_WORDS:
        return True
    return False


def should_be_yt(ch):
    """Check if a single character should be in Yada Towrah font."""
    return ch in HALF_RINGS


def get_all_runs(para):
    runs = []
    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            runs.append(child)
        elif tag == 'hyperlink':
            for r in child.findall(qn('w:r')):
                runs.append(r)
    return runs


def fix_yt_runs(para, stats):
    """Fix non-Hebrew text in YT font runs by switching to TNR."""
    runs = get_all_runs(para)
    for r_elem in runs:
        font = get_run_font(r_elem)
        if font.lower() != YT_FONT:
            continue

        text = get_run_text(r_elem)
        if not text:
            continue

        # If entire run is valid YT, skip
        if is_valid_yt(text):
            continue

        # Need to split: keep half-ring chars in YT, switch rest to TNR
        # Build segments: (text, should_be_yt_font)
        segments = []
        current = ''
        current_yt = should_be_yt(text[0])

        for ch in text:
            ch_yt = should_be_yt(ch)
            if ch_yt != current_yt:
                if current:
                    segments.append((current, current_yt))
                current = ch
                current_yt = ch_yt
            else:
                current += ch
        if current:
            segments.append((current, current_yt))

        # If no YT chars at all, just switch the whole run
        if not any(is_yt for _, is_yt in segments):
            set_run_font_name(r_elem, TNR_FONT)
            stats['runs'] += 1
            stats['chars'] += len(text)
            continue

        # Split into multiple runs
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)

        # First segment in original run
        set_run_text(r_elem, segments[0][0])
        if not segments[0][1]:
            set_run_font_name(r_elem, TNR_FONT)
            stats['chars'] += len(segments[0][0])

        insert_after = idx
        for seg_text, seg_yt in segments[1:]:
            new_run = copy.deepcopy(r_elem)
            set_run_text(new_run, seg_text)
            if not seg_yt:
                set_run_font_name(new_run, TNR_FONT)
                stats['chars'] += len(seg_text)
            else:
                # Restore YT font on the half-ring run
                set_run_font_name(new_run, 'Yada Towrah')
            insert_after += 1
            parent.insert(insert_after, new_run)

        stats['runs'] += 1


def process_doc(fpath):
    doc = Document(fpath)
    stats = {'runs': 0, 'chars': 0}

    for para in doc.paragraphs:
        sn = para.style.name if para.style else ''
        if sn == 'yy_break':
            continue
        fix_yt_runs(para, stats)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_yt_runs(para, stats)

    # Headers/footers
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr:
                try:
                    hdr.is_linked_to_previous
                except:
                    continue
                for para in hdr.paragraphs:
                    fix_yt_runs(para, stats)

    if stats['runs'] > 0:
        doc.save(fpath)
    return stats


def main():
    files = sorted(glob.glob(os.path.join(DOCS_DIR, 'YY-*.docx')))
    files = [f for f in files if '_backup' not in f and '_updated' not in f]
    print(f'Found {len(files)} documents')

    total_runs = 0
    total_chars = 0
    for i, fpath in enumerate(files):
        fname = os.path.basename(fpath)
        print(f'  [{i+1}/{len(files)}] {fname}', end='', flush=True)
        stats = process_doc(fpath)
        if stats['runs']:
            print(f'  {stats["runs"]} runs, {stats["chars"]} chars')
        else:
            print('  (no changes)')
        total_runs += stats['runs']
        total_chars += stats['chars']

    print()
    print('=' * 60)
    print(f'Total runs fixed:  {total_runs}')
    print(f'Total chars fixed: {total_chars}')


if __name__ == '__main__':
    main()
