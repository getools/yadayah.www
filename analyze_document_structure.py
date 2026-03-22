import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from lxml import etree

doc = Document(r'C:\users\joe\work\dev\yada\docs\YY-s02v05-Yada Yahowah-Qatsyr-Harvests.docx')

print("Document Structure Analysis")
print("=" * 80)

# Get sections and their page numbering
body = doc._element.body
sections = []

for i, para in enumerate(doc.paragraphs):
    para_element = para._element
    pPr = para_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        sectPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
        if sectPr is not None:
            pgNumType = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
            if pgNumType is not None:
                fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
                start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
                sections.append({
                    'para_idx': i,
                    'para_text': doc.paragraphs[i].text[:60],
                    'fmt': fmt,
                    'start': start
                })

# Document-level section
sectPr = body.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
if sectPr is not None:
    pgNumType = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
    if pgNumType is not None:
        fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
        start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
        sections.append({
            'para_idx': len(doc.paragraphs) - 1,
            'para_text': 'END',
            'fmt': fmt,
            'start': start
        })

print(f"\nFound {len(sections)} sections:\n")
for i, section in enumerate(sections):
    print(f"Section {i}:")
    print(f"  Starts at paragraph: {section['para_idx']}")
    print(f"  Text: {section['para_text']}")
    print(f"  Page number format: {section['fmt']}")
    print(f"  Page number start: {section['start']}")
    print()

# Find taglines and show their locations
print("\nTaglines found:\n")
tagline_count = 0
for i, para in enumerate(doc.paragraphs[:100]):  # First 100 paragraphs
    text = para.text.strip()
    if text.endswith('...') or text.endswith('…'):
        if para.style and hasattr(para.style, 'font') and para.style.font.italic:
            # Determine which section this is in
            section_num = 0
            for j, section in enumerate(sections):
                if i >= section['para_idx']:
                    section_num = j

            print(f"Paragraph {i} (Section {section_num}):")
            print(f"  Text: {text[:70]}")
            print(f"  Style: {para.style.name}")
            tagline_count += 1

            if tagline_count >= 10:
                print("\n... (showing first 10 only)")
                break

print(f"\nTotal taglines found in first 100 paragraphs: {tagline_count}")

# Find the specific paragraph mentioned
print("\n" + "=" * 80)
print("Looking for target paragraph...")
target_text = 'The fourth Miqra'
for i, para in enumerate(doc.paragraphs):
    if target_text in para.text:
        # Determine section
        section_num = 0
        for j, section in enumerate(sections):
            if i >= section['para_idx']:
                section_num = j

        print(f"\nFound at paragraph {i} (Section {section_num}):")
        print(f"  Text: {para.text[:100]}...")
        print(f"  Should be in content area with page numbers")
        break
