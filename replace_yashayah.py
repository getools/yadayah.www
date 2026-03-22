"""
Replace Yasha'yah with Yashaʿyah in all YY Word documents.
The ʿ (Ayin) character uses "Yada Towrah" font, other letters preserve original formatting.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time
import re

# Character definitions
AYIN = chr(0x02BF)   # ʿ (modifier letter left half ring)
LEFT_CURLY = chr(0x2018)
RIGHT_CURLY = chr(0x2019)
STRAIGHT_QUOTE = "'"

# Patterns to match (case-insensitive)
SEARCH_PATTERNS = [
    "Yasha'yah",
    "yasha'yah",
    "YASHA'YAH",
]

def normalize_quotes(s):
    """Normalize all quote variations to straight apostrophe."""
    return s.replace(LEFT_CURLY, "'").replace(RIGHT_CURLY, "'").replace(chr(0x2032), "'")

def should_replace(text):
    """Check if text contains Yasha'yah (any quote variation, any case)."""
    normalized = normalize_quotes(text).lower()
    return "yasha'yah" in normalized

def apply_capitalization_pattern(original, template="yashaʿyah"):
    """
    Apply capitalization from original to template.
    Handles: Yasha'yah -> Yashaʿyah, yasha'yah -> yashaʿyah, YASHA'YAH -> YASHAʿYAH
    """
    if not original:
        return template

    # Normalize quotes in original
    normalized_orig = normalize_quotes(original).lower()

    # Find the "yasha'yah" part in normalized original
    if "yasha'yah" not in normalized_orig:
        return template

    # Get the actual "yasha'yah" substring from original (preserving case)
    pattern = re.compile(r"yasha['\u2018\u2019\u2032]yah", re.IGNORECASE)
    match = pattern.search(original)
    if not match:
        return template

    source = match.group()

    # Check if all caps
    source_letters = [c for c in source if c.isalpha()]
    if source_letters and all(c.isupper() for c in source_letters):
        return template.upper()

    # Apply character-by-character capitalization
    result = []
    source_idx = 0
    for target_char in template:
        # Skip non-letter chars in source
        while source_idx < len(source) and not source[source_idx].isalpha():
            source_idx += 1

        if source_idx < len(source) and target_char.isalpha():
            if source[source_idx].isupper():
                result.append(target_char.upper())
            else:
                result.append(target_char.lower())
            source_idx += 1
        else:
            result.append(target_char)

    return ''.join(result)

def create_replacement_runs(run, original_text, replacement_text):
    """
    Create new runs for replacement text with proper font formatting.
    ʿ uses "Yada Towrah", other letters preserve original font.
    """
    # Save original formatting
    is_bold = run.bold
    is_italic = run.italic
    font_name = run.font.name
    font_size = run.font.size
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = run.font.color.rgb

    # Get parent paragraph and run position
    para = run._element.getparent()
    run_idx = list(para).index(run._element)

    # Clear original run
    run.text = ""

    def create_run(text, use_yada_font=False):
        """Create a run with specified formatting."""
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        # Italic
        if is_italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)

        # Font
        rFonts = OxmlElement('w:rFonts')
        target_font = "Yada Towrah" if use_yada_font else (font_name or "Calibri")
        rFonts.set(qn('w:ascii'), target_font)
        rFonts.set(qn('w:hAnsi'), target_font)
        rPr.append(rFonts)

        # Font size
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(szCs)

        # Color
        if font_color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), str(font_color))
            rPr.append(color)

        new_run.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        return new_run

    # Split replacement text at ʿ character
    current_run_idx = run_idx
    buffer = []

    def flush_buffer(is_ayin=False):
        nonlocal current_run_idx
        if not buffer:
            return
        text = ''.join(buffer)
        new_run = create_run(text, use_yada_font=is_ayin)
        para.insert(current_run_idx + 1, new_run)
        current_run_idx += 1
        buffer.clear()

    for char in replacement_text:
        if char == AYIN:
            # Flush regular text
            flush_buffer(is_ayin=False)
            # Add ʿ and flush
            buffer.append(char)
            flush_buffer(is_ayin=True)
        else:
            buffer.append(char)

    # Flush remaining
    flush_buffer(is_ayin=False)

def process_paragraph(para, stats):
    """Process a paragraph, replacing Yasha'yah with Yashaʿyah."""
    for run in list(para.runs):
        if not run.text:
            continue

        if should_replace(run.text):
            # Get the capitalization-matched replacement
            replacement = apply_capitalization_pattern(run.text, "yashaʿyah")

            # Replace all quote variations
            new_text = run.text
            pattern = re.compile(r"yasha['\u2018\u2019\u2032]yah", re.IGNORECASE)

            # Find and replace
            match = pattern.search(new_text)
            if match:
                # Replace the matched portion
                before = new_text[:match.start()]
                after = new_text[match.end():]
                final_text = before + replacement + after

                if final_text != run.text:
                    create_replacement_runs(run, run.text, final_text)
                    stats['replacements'] += 1
                    if stats['verbose']:
                        print(f"    {run.text} -> {final_text}")

def process_document(docx_path, verbose=False):
    """Process a single Word document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: Could not open document: {e}")
        return 0

    stats = {'replacements': 0, 'verbose': verbose}

    # Process paragraphs
    for para in doc.paragraphs:
        process_paragraph(para, stats)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, stats)

    elapsed = time.time() - start_time
    replacements = stats['replacements']

    if replacements > 0:
        # Save with _yashayah_updated suffix
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_yashayah_updated{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ {replacements} replacements, saved to {os.path.basename(output_path)} ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: Could not save document: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return replacements

# Main execution
docs_dir = r'C:\users\joe\work\dev\yada\docs'
if not os.path.exists(docs_dir):
    print(f"ERROR: Documents directory not found: {docs_dir}")
    sys.exit(1)

print("Replacing Yasha'yah with Yashaʿyah (ʿ in Yada Towrah font)...")
print()

# Find all YY documents
docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-')
    and f.endswith('.docx')
    and not '_yashayah_updated' in f
    and not f.startswith('~$')
]

print(f"Found {len(docx_files)} YY documents to process\n")
print("=" * 80)

total_files_updated = 0
total_replacements = 0

for docx_path in docx_files:
    replacements = process_document(docx_path, verbose=False)
    if replacements > 0:
        total_files_updated += 1
        total_replacements += replacements

print("\n" + "=" * 80)
print("REPLACEMENT COMPLETE")
print("=" * 80)
print(f"  Files processed: {len(docx_files)}")
print(f"  Files updated: {total_files_updated}")
print(f"  Total replacements: {total_replacements}")
print(f"\nUpdated files saved with '_yashayah_updated.docx' suffix in: {docs_dir}")
