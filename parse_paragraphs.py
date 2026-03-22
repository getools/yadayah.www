#!/usr/bin/env python3
"""
Parse ALL YY Word documents and populate yy_paragraph table.

Extracts every paragraph from each document with:
  - Full formatting metadata (paragraph_text_raw as JSON)
  - HTML+CSS rendering (paragraph_text_html)
  - Word footer page number (via COM/VBScript)
  - Chapter assignment (via paragraph style detection)

Also generates a global CSS stylesheet from all formatting encountered.

Usage:
    python parse_paragraphs.py
    python parse_paragraphs.py --dry-run --verbose
    python parse_paragraphs.py --skip-pages
    python parse_paragraphs.py --single "YY-s01v01-An Intro to God-Dabarym-Words"
"""

import os
import sys
import json
import re
import argparse
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import psycopg2
    from psycopg2.extras import execute_values
except ImportError:
    print("ERROR: psycopg2 not installed. Run: pip install psycopg2-binary")
    sys.exit(1)

try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: python-dotenv not installed. Run: pip install python-dotenv")
    sys.exit(1)

from parse_word_translations import replace_unicode_chars, consolidate_html, build_footer_page_map, get_com_page_map_iter

DEFAULT_DIRECTORY = r"C:\users\joe\work\dev\yada\docs"

# ── Database connection ──────────────────────────────────────────────

def get_db_connection():
    load_dotenv()
    host = os.getenv('POSTGRES_HOST', 'localhost')
    port = os.getenv('POSTGRES_PORT', '5432')
    user = os.getenv('POSTGRES_USER', 'postgres')
    password = os.getenv('POSTGRES_PASSWORD')
    db = os.getenv('POSTGRES_DB', 'yada')
    if not password:
        print("ERROR: POSTGRES_PASSWORD not set in .env")
        sys.exit(1)
    return psycopg2.connect(host=host, port=port, user=user, password=password, dbname=db)


# ── Phase 1: FK Lookups ─────────────────────────────────────────────

def load_volume_mapping(conn):
    """Load filename stem -> (yy_volume_key, yy_series_key) mapping."""
    cur = conn.cursor()
    cur.execute("SELECT volume_key, volume_file, series_key FROM yy_volume WHERE volume_file IS NOT NULL")
    mapping = {}
    for vol_key, filename, series_key in cur.fetchall():
        if filename:
            mapping[filename] = (vol_key, series_key)
            norm = filename.replace('\u2019', "'").replace('\u2018', "'")
            if norm != filename:
                mapping[norm] = (vol_key, series_key)
    cur.close()
    return mapping


def load_yy_chapter_mapping(conn):
    """Load (yy_volume_key, chapter_number) -> yy_chapter_key."""
    cur = conn.cursor()
    cur.execute("SELECT chapter_key, volume_key, chapter_number FROM yy_chapter")
    mapping = {}
    for ch_key, vol_key, ch_num in cur.fetchall():
        mapping[(vol_key, ch_num)] = ch_key
    cur.close()
    return mapping


# ── Formatting helpers ───────────────────────────────────────────────

def emu_to_pt(emu_value):
    """Convert EMU to points. Returns None if input is None."""
    if emu_value is None:
        return None
    return round(emu_value / 12700, 1)


def alignment_to_str(alignment):
    """Convert WD_ALIGN_PARAGRAPH enum to string."""
    if alignment is None:
        return None
    mapping = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
    return mapping.get(int(alignment), None)


def format_line_spacing(spacing, rule):
    """Format line spacing value based on rule."""
    if spacing is None:
        return None
    # rule: 0=AT_LEAST, 1=EXACTLY, 2=MULTIPLE, 4=AT_LEAST(alt)
    if rule is not None and int(rule) in (0, 1, 4):
        # Exact or at-least: value is in EMU, convert to pt
        return emu_to_pt(spacing)
    # MULTIPLE: spacing is a float multiplier (typically 1.0-3.0 range)
    if isinstance(spacing, float) and spacing < 100:
        return round(spacing, 2)
    # Large int values are EMU (python-docx Emu subclass of int)
    if isinstance(spacing, int) and spacing > 100:
        return emu_to_pt(spacing)
    if isinstance(spacing, (int, float)):
        return round(float(spacing), 2)
    return emu_to_pt(spacing)


def slugify_font(font_name):
    """Convert font name to CSS-safe slug."""
    return re.sub(r'[^a-z0-9-]', '-', font_name.lower()).strip('-')


def pt_to_slug(pt_value):
    """Convert point value to CSS class slug (12.0 -> '12', 10.5 -> '10-5')."""
    if pt_value is None:
        return '0'
    if pt_value == int(pt_value):
        return str(int(pt_value))
    return str(pt_value).replace('.', '-')


# ── Paragraph extraction ────────────────────────────────────────────

def build_paragraph_raw(paragraph):
    """Build structured dict of paragraph formatting metadata."""
    pf = paragraph.paragraph_format

    raw = {
        'style': paragraph.style.name if paragraph.style else None,
        'alignment': alignment_to_str(paragraph.alignment),
        'indent': {
            'first_line': emu_to_pt(pf.first_line_indent),
            'left': emu_to_pt(pf.left_indent),
            'right': emu_to_pt(pf.right_indent),
        },
        'spacing': {
            'before': emu_to_pt(pf.space_before),
            'after': emu_to_pt(pf.space_after),
            'line': format_line_spacing(pf.line_spacing, pf.line_spacing_rule),
        },
        'runs': [],
    }

    for run in paragraph.runs:
        if not run.text:
            continue
        text = replace_unicode_chars(run.text)

        font_name = None
        try:
            font_name = run.font.name
        except Exception:
            pass

        font_size = None
        try:
            font_size = emu_to_pt(run.font.size)
        except Exception:
            pass

        color_hex = None
        try:
            if run.font.color and run.font.color.rgb:
                color_hex = str(run.font.color.rgb)
        except Exception:
            pass

        run_data = {
            'text': text,
            'bold': bool(run.bold) if run.bold is not None else False,
            'italic': bool(run.italic) if run.italic is not None else False,
            'underline': bool(run.underline) if run.underline is not None else False,
            'superscript': bool(run.font.superscript) if run.font.superscript is not None else False,
            'subscript': bool(run.font.subscript) if run.font.subscript is not None else False,
            'font': font_name,
            'size': font_size,
            'color': color_hex,
        }
        raw['runs'].append(run_data)

    return raw


def build_paragraph_html(paragraph, css_collector):
    """Build HTML rendering of paragraph with CSS classes."""
    pf = paragraph.paragraph_format

    # Paragraph-level classes
    p_classes = []

    align = alignment_to_str(paragraph.alignment)
    if align:
        p_classes.append(f'yy-align-{align}')
        css_collector['aligns'].add(align)

    first_indent = emu_to_pt(pf.first_line_indent)
    if first_indent is not None and first_indent != 0:
        slug = pt_to_slug(first_indent)
        p_classes.append(f'yy-indent-first-{slug}')
        css_collector['indents_first'].add(first_indent)

    left_indent = emu_to_pt(pf.left_indent)
    if left_indent is not None and left_indent != 0:
        slug = pt_to_slug(left_indent)
        p_classes.append(f'yy-indent-left-{slug}')
        css_collector['indents_left'].add(left_indent)

    right_indent = emu_to_pt(pf.right_indent)
    if right_indent is not None and right_indent != 0:
        slug = pt_to_slug(right_indent)
        p_classes.append(f'yy-indent-right-{slug}')
        css_collector['indents_right'].add(right_indent)

    space_before = emu_to_pt(pf.space_before)
    if space_before is not None and space_before != 0:
        slug = pt_to_slug(space_before)
        p_classes.append(f'yy-space-before-{slug}')
        css_collector['spaces_before'].add(space_before)

    space_after = emu_to_pt(pf.space_after)
    if space_after is not None and space_after != 0:
        slug = pt_to_slug(space_after)
        p_classes.append(f'yy-space-after-{slug}')
        css_collector['spaces_after'].add(space_after)

    line_sp = format_line_spacing(pf.line_spacing, pf.line_spacing_rule)
    if line_sp is not None:
        slug = pt_to_slug(line_sp)
        p_classes.append(f'yy-line-{slug}')
        css_collector['line_spacings'].add(line_sp)

    # Build run HTML
    run_htmls = []
    for run in paragraph.runs:
        text = replace_unicode_chars(run.text)
        if not text:
            continue

        # Escape HTML entities in text
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        # Run-level span classes
        span_classes = []

        font_name = None
        try:
            font_name = run.font.name
        except Exception:
            pass
        if font_name:
            font_slug = slugify_font(font_name)
            span_classes.append(f'yy-font-{font_slug}')
            css_collector['fonts'].add((font_slug, font_name))

        size_pt = None
        try:
            size_pt = emu_to_pt(run.font.size)
        except Exception:
            pass
        if size_pt is not None:
            span_classes.append(f'yy-size-{pt_to_slug(size_pt)}')
            css_collector['sizes'].add(size_pt)

        color_hex = None
        try:
            if run.font.color and run.font.color.rgb:
                color_hex = str(run.font.color.rgb).lower()
        except Exception:
            pass
        if color_hex:
            span_classes.append(f'yy-color-{color_hex}')
            css_collector['colors'].add(color_hex)

        # Build the span
        if span_classes:
            html_frag = f'<span class="{" ".join(span_classes)}">{text}</span>'
        else:
            html_frag = text

        # Wrap with formatting tags (innermost to outermost)
        if run.font.subscript:
            html_frag = f'<sub>{html_frag}</sub>'
        if run.font.superscript:
            html_frag = f'<sup>{html_frag}</sup>'
        if run.underline:
            html_frag = f'<u>{html_frag}</u>'
        if run.italic:
            html_frag = f'<i>{html_frag}</i>'
        if run.bold:
            html_frag = f'<b>{html_frag}</b>'

        run_htmls.append(html_frag)

    # Consolidate adjacent identical tags
    inner_html = consolidate_html(''.join(run_htmls))

    # Build <p> tag
    class_attr = f' class="{" ".join(p_classes)}"' if p_classes else ''
    return f'<p{class_attr}>{inner_html}</p>'


def extract_paragraphs_from_doc(doc_path, css_collector, content_start_idx=0, last_page=None, page_map=None):
    """
    Extract paragraphs from a Word document, filtering by footer page range.

    Only includes paragraphs starting from content_start_idx (footer page 1)
    and excludes paragraphs on the last page.

    Args:
        doc_path: Path to the .docx file
        css_collector: Dict to collect CSS class data
        content_start_idx: python-docx paragraph index where footer page 1 begins
        last_page: Maximum page number (paragraphs on this page are excluded)
        page_map: Dict of para_idx -> page_number from build_footer_page_map

    Returns:
        (paragraphs_list, chapter_boundaries)
        paragraphs_list: list of dicts with paragraph_number, text_raw, text_html, page
        chapter_boundaries: list of (para_idx, chapter_number) tuples
    """
    doc = Document(str(doc_path))
    paragraphs = []
    chapter_boundaries = []

    for para_idx, paragraph in enumerate(doc.paragraphs):
        # Skip paragraphs before content start (front matter)
        if para_idx < content_start_idx:
            continue

        # Skip paragraphs on the last page
        page = page_map.get(para_idx) if page_map else None
        if last_page is not None and page is not None and page >= last_page:
            continue

        # Detect chapter boundaries
        style_name = paragraph.style.name if paragraph.style else ''
        if style_name in ('yy_chapter_#', 'Heading 1'):
            text = paragraph.text.strip()
            if text:
                first_line = text.split('\n')[0].strip()
                if first_line.isdigit():
                    chapter_boundaries.append((para_idx, int(first_line)))

        # Build raw JSON
        raw_data = build_paragraph_raw(paragraph)

        # Build HTML
        html = build_paragraph_html(paragraph, css_collector)

        paragraphs.append({
            'paragraph_number': para_idx,
            'text_raw': json.dumps(raw_data, ensure_ascii=False),
            'text_html': html,
            'page': page,
        })

    return paragraphs, chapter_boundaries


# ── Chapter assignment ───────────────────────────────────────────────

def get_chapter_for_paragraph(para_idx, chapter_boundaries):
    """Return chapter number for a paragraph based on boundary list, or None."""
    current_chapter = None
    for boundary_idx, chapter_num in chapter_boundaries:
        if para_idx >= boundary_idx:
            current_chapter = chapter_num
        else:
            break
    return current_chapter


# ── CSS generation ───────────────────────────────────────────────────

def new_css_collector():
    return {
        'fonts': set(),
        'sizes': set(),
        'colors': set(),
        'aligns': set(),
        'indents_first': set(),
        'indents_left': set(),
        'indents_right': set(),
        'spaces_before': set(),
        'spaces_after': set(),
        'line_spacings': set(),
    }


def generate_css(css_collector, output_path):
    """Generate CSS stylesheet from collected formatting data."""
    lines = []
    lines.append('/* YY Books Global Stylesheet */')
    lines.append(f'/* Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} */')
    total_classes = sum(len(v) for v in css_collector.values())
    lines.append(f'/* Unique formatting classes: {total_classes} */')
    lines.append('')

    # Fonts
    lines.append('/* -- Fonts -- */')
    for font_slug, font_name in sorted(css_collector['fonts']):
        lines.append(f'.yy-font-{font_slug} {{ font-family: "{font_name}"; }}')
    lines.append('')

    # Sizes
    lines.append('/* -- Sizes -- */')
    for size in sorted(css_collector['sizes']):
        slug = pt_to_slug(size)
        lines.append(f'.yy-size-{slug} {{ font-size: {size}pt; }}')
    lines.append('')

    # Colors
    lines.append('/* -- Colors -- */')
    for color in sorted(css_collector['colors']):
        lines.append(f'.yy-color-{color} {{ color: #{color}; }}')
    lines.append('')

    # Alignment
    lines.append('/* -- Alignment -- */')
    for align in sorted(css_collector['aligns']):
        lines.append(f'.yy-align-{align} {{ text-align: {align}; }}')
    lines.append('')

    # Indents
    lines.append('/* -- Indents -- */')
    for val in sorted(css_collector['indents_first']):
        slug = pt_to_slug(val)
        lines.append(f'.yy-indent-first-{slug} {{ text-indent: {val}pt; }}')
    for val in sorted(css_collector['indents_left']):
        slug = pt_to_slug(val)
        lines.append(f'.yy-indent-left-{slug} {{ margin-left: {val}pt; }}')
    for val in sorted(css_collector['indents_right']):
        slug = pt_to_slug(val)
        lines.append(f'.yy-indent-right-{slug} {{ margin-right: {val}pt; }}')
    lines.append('')

    # Spacing
    lines.append('/* -- Spacing -- */')
    for val in sorted(css_collector['spaces_before']):
        slug = pt_to_slug(val)
        lines.append(f'.yy-space-before-{slug} {{ margin-top: {val}pt; }}')
    for val in sorted(css_collector['spaces_after']):
        slug = pt_to_slug(val)
        lines.append(f'.yy-space-after-{slug} {{ margin-bottom: {val}pt; }}')
    lines.append('')

    # Line spacing
    lines.append('/* -- Line Spacing -- */')
    for val in sorted(css_collector['line_spacings']):
        slug = pt_to_slug(val)
        if isinstance(val, float) and val < 5:
            lines.append(f'.yy-line-{slug} {{ line-height: {val}; }}')
        else:
            lines.append(f'.yy-line-{slug} {{ line-height: {val}pt; }}')
    lines.append('')

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines) + '\n')

    logging.info(f"CSS stylesheet written: {output_path} ({total_classes} classes)")


# ── Main ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Parse YY Word documents into yy_paragraph table')
    parser.add_argument('--directory', default=DEFAULT_DIRECTORY, help='Documents directory')
    parser.add_argument('--dry-run', action='store_true', help='Extract but do not write to DB')
    parser.add_argument('--verbose', action='store_true', help='Debug logging')
    parser.add_argument('--single', help='Process only one document (filename stem)')
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S',
    )

    # Ensure stdout handles Unicode
    if sys.stdout.encoding != 'utf-8':
        sys.stdout.reconfigure(encoding='utf-8')

    directory = Path(args.directory)
    if not directory.exists():
        logging.error(f"Directory not found: {directory}")
        sys.exit(1)

    conn = get_db_connection()
    logging.info("Connected to PostgreSQL")

    stats = {
        'docs_processed': 0,
        'docs_skipped': [],
        'total_paragraphs': 0,
        'content_paragraphs': 0,
        'front_matter_skipped': 0,
        'last_page_skipped': 0,
        'paragraphs_inserted': 0,
        'page_numbers_found': 0,
        'page_numbers_null': 0,
        'chapters_assigned': 0,
        'chapters_null': 0,
    }

    # ── Phase 1: Load FK lookups ──
    logging.info("\n=== Phase 1: Loading FK lookup tables ===")
    volume_map = load_volume_mapping(conn)
    logging.info(f"  Volume mappings: {len(volume_map)}")
    yy_chapter_map = load_yy_chapter_mapping(conn)
    logging.info(f"  YY chapter mappings: {len(yy_chapter_map)}")

    # ── Phase 2: Discover documents ──
    logging.info("\n=== Phase 2: Discovering documents ===")
    docx_files = [f for f in directory.glob("YY*.docx")
                  if not f.name.startswith("~$")]
    docx_files.sort()

    matched_files = []
    for doc_path in docx_files:
        stem = doc_path.stem
        vol_info = volume_map.get(stem)
        if vol_info is None:
            norm_stem = stem.replace('\u2019', "'").replace('\u2018', "'")
            vol_info = volume_map.get(norm_stem)
        if vol_info is not None:
            if args.single and stem != args.single:
                continue
            matched_files.append((doc_path, vol_info[0], vol_info[1]))
        else:
            logging.debug(f"  Skipping {stem} (no volume mapping)")
            stats['docs_skipped'].append(stem)

    logging.info(f"  Found {len(matched_files)} documents to process")
    if not matched_files:
        logging.error("No documents matched. Check --directory and --single.")
        sys.exit(1)

    # ── Phase 3: Extract paragraphs (COM page numbers + content filtering) ──
    logging.info("\n=== Phase 3: Extracting content paragraphs (page 1 to last-1) ===")
    css_collector = new_css_collector()
    all_data = []  # list of (doc_path, vol_key, series_key, paragraphs, chapter_boundaries)

    for i, (doc_path, vol_key, series_key) in enumerate(matched_files):
        logging.info(f"  [{i+1}/{len(matched_files)}] {doc_path.stem} (vol={vol_key}, series={series_key})")
        try:
            # Step 1: Get content_start from XML section properties (reliable)
            _, content_start, _ = build_footer_page_map(doc_path)

            # Step 2: Get COM page numbers via iteration (fast single pass)
            com_page_map = get_com_page_map_iter(str(doc_path.absolute()), timeout=600)

            doc_temp = Document(str(doc_path))
            total_paras = len(doc_temp.paragraphs)
            stats['total_paragraphs'] += total_paras

            com_content = {k: v for k, v in com_page_map.items() if k >= content_start}
            logging.info(f"    COM iter: {len(com_page_map)} total, {len(com_content)} content page numbers")

            # Step 3: Interpolate missing page numbers
            page_map = dict(com_page_map)
            sorted_known = sorted(k for k in com_page_map.keys() if k >= content_start)
            for idx in range(content_start, total_paras):
                if idx not in page_map:
                    prev_page = None
                    for k in reversed(sorted_known):
                        if k < idx:
                            prev_page = com_page_map[k]
                            break
                    if prev_page is not None:
                        page_map[idx] = prev_page
                    elif sorted_known:
                        page_map[idx] = com_page_map[sorted_known[0]]

            # Step 4: Determine last_page from COM results
            content_pages = [p for idx, p in page_map.items() if idx >= content_start]
            last_page = max(content_pages) if content_pages else 1

            # Check if last-page skip is safe
            last_page_count = sum(1 for idx, pg in page_map.items() if idx >= content_start and pg >= last_page)
            effective_last_page = last_page
            if last_page_count > 100:
                logging.warning(f"    Unreliable page breaks: {last_page_count} paragraphs on page {last_page}; not skipping last page")
                effective_last_page = last_page + 1  # effectively no last-page filter

            # Step 5: Extract only content paragraphs (page 1 through last_page - 1)
            paragraphs, chapter_boundaries = extract_paragraphs_from_doc(
                doc_path, css_collector,
                content_start_idx=content_start,
                last_page=effective_last_page,
                page_map=page_map,
            )

            front_skipped = content_start
            last_skipped = total_paras - content_start - len(paragraphs)
            stats['front_matter_skipped'] += front_skipped
            stats['last_page_skipped'] += last_skipped
            stats['content_paragraphs'] += len(paragraphs)
            stats['docs_processed'] += 1

            logging.info(f"    {total_paras} total, {len(paragraphs)} content (skipped {front_skipped} front + {last_skipped} last-page), "
                         f"{len(chapter_boundaries)} chapters, pages 1-{last_page-1} of {last_page}")

            all_data.append((doc_path, vol_key, series_key, paragraphs, chapter_boundaries))
        except Exception as e:
            logging.error(f"    Failed to extract: {e}")
            import traceback
            traceback.print_exc()
            stats['docs_skipped'].append(doc_path.stem)

    # ── Phase 4: Insert into yy_paragraph ──
    if not args.dry_run:
        logging.info("\n=== Phase 4: Inserting into yy_paragraph ===")
        cur = conn.cursor()

        cur.execute("TRUNCATE yy_paragraph CASCADE")
        cur.execute("ALTER SEQUENCE yy_paragraph_paragraph_key_seq RESTART WITH 1")
        conn.commit()
        logging.info("  Truncated yy_paragraph and reset sequence")

        rows = []
        for doc_path, vol_key, series_key, paragraphs, chapter_boundaries in all_data:
            for p in paragraphs:
                para_num = p['paragraph_number']

                # Chapter assignment
                ch_num = get_chapter_for_paragraph(para_num, chapter_boundaries)
                chapter_key = yy_chapter_map.get((vol_key, ch_num)) if ch_num else None
                if chapter_key:
                    stats['chapters_assigned'] += 1
                else:
                    stats['chapters_null'] += 1

                # Page number (already in paragraph dict from extraction)
                page = p.get('page')
                if page is not None:
                    stats['page_numbers_found'] += 1
                else:
                    stats['page_numbers_null'] += 1

                rows.append((
                    series_key,
                    vol_key,
                    chapter_key,
                    page,
                    para_num,
                    p['text_raw'],
                    p['text_html'],
                ))

        logging.info(f"  Inserting {len(rows)} rows...")
        execute_values(cur, """
            INSERT INTO yy_paragraph (
                series_key, volume_key, chapter_key,
                paragraph_page, paragraph_number,
                paragraph_text_raw, paragraph_text_html
            ) VALUES %s
        """, rows, page_size=500)
        conn.commit()
        stats['paragraphs_inserted'] = len(rows)
        logging.info(f"  Inserted {len(rows)} rows into yy_paragraph")
    else:
        logging.info("\n=== Phase 4: Dry run - no DB writes ===")
        for doc_path, vol_key, series_key, paragraphs, chapter_boundaries in all_data:
            for p in paragraphs:
                para_num = p['paragraph_number']
                ch_num = get_chapter_for_paragraph(para_num, chapter_boundaries)
                chapter_key = yy_chapter_map.get((vol_key, ch_num)) if ch_num else None
                if chapter_key:
                    stats['chapters_assigned'] += 1
                else:
                    stats['chapters_null'] += 1
                if p.get('page') is not None:
                    stats['page_numbers_found'] += 1
                else:
                    stats['page_numbers_null'] += 1

    # ── Phase 5: Generate CSS ──
    logging.info("\n=== Phase 5: Generating CSS stylesheet ===")
    css_dir = Path(__file__).parent / 'public' / 'css'
    css_dir.mkdir(parents=True, exist_ok=True)
    css_filename = f"yy.books.{datetime.now().strftime('%Y%m%d')}.css"
    css_path = css_dir / css_filename
    generate_css(css_collector, css_path)

    # ── Summary ──
    conn.close()
    logging.info("\n" + "=" * 60)
    logging.info("SUMMARY")
    logging.info("=" * 60)
    logging.info(f"Documents processed:   {stats['docs_processed']}")
    if stats['docs_skipped']:
        logging.info(f"Documents skipped:     {stats['docs_skipped']}")
    logging.info(f"Total paragraphs:      {stats['total_paragraphs']}")
    logging.info(f"Front matter skipped:  {stats['front_matter_skipped']}")
    logging.info(f"Last page skipped:     {stats['last_page_skipped']}")
    logging.info(f"Content paragraphs:    {stats['content_paragraphs']}")
    logging.info(f"Paragraphs inserted:   {stats['paragraphs_inserted']}")
    logging.info(f"Page numbers found:    {stats['page_numbers_found']}")
    logging.info(f"Page numbers null:     {stats['page_numbers_null']}")
    logging.info(f"Chapters assigned:     {stats['chapters_assigned']}")
    logging.info(f"Chapters null:         {stats['chapters_null']}")
    logging.info(f"CSS file:              {css_path}")


if __name__ == '__main__':
    main()
