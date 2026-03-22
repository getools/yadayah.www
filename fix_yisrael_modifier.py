"""
Replace Yisraʿel (AYIN) with Yisraʾel (ALEPH) in all YY documents.
The ALEPH modifier ʾ should use Yada Towrah font.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

AYIN = chr(0x02BF)   # ʿ (wrong for Yisrael)
ALEPH = chr(0x02BE)  # ʾ (correct for Yisrael)

def fix_yisrael_modifier(t_elem, run_elem):
    """
    Replace AYIN with ALEPH in Yisraʿel.
    The text element should be exactly the AYIN character if it's part of Yisraʿel pattern.
    """
    if not t_elem.text or t_elem.text != AYIN:
        return False

    # Check if this AYIN is part of "Yisraʿel" pattern
    # Look at previous and next runs
    parent = run_elem.getparent()
    runs = list(parent.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'))
    run_idx = runs.index(run_elem)

    # Check previous runs for "Yisra" (skip empty runs)
    found_yisra = False
    for i in range(run_idx - 1, max(-1, run_idx - 15), -1):
        prev_run = runs[i]
        for t in prev_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                if 'Yisra' in t.text:
                    found_yisra = True
                break  # Stop at first non-empty run
        if found_yisra or (prev_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') and
                          any(t.text for t in prev_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))):
            break

    if not found_yisra:
        return False

    # Check next runs for "el" (skip empty runs)
    found_el = False
    for i in range(run_idx + 1, min(len(runs), run_idx + 15)):
        next_run = runs[i]
        for t in next_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                if t.text.startswith('el'):
                    found_el = True
                break  # Stop at first non-empty run
        if found_el or (next_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') and
                       any(t.text for t in next_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))):
            break

    if not found_el:
        return False

    # Replace AYIN with ALEPH
    t_elem.text = ALEPH

    # Ensure Yada Towrah font is set
    rPr = run_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run_elem.insert(0, rPr)

    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    rFonts.set(qn('w:ascii'), 'Yada Towrah')
    rFonts.set(qn('w:hAnsi'), 'Yada Towrah')

    return True

def process_paragraph(para):
    """Process all text elements in a paragraph."""
    fixes = 0

    for element in list(para._element.iter()):
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            run_elem = element.getparent()
            if run_elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                if fix_yisrael_modifier(element, run_elem):
                    fixes += 1

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No Yisraʿel found ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Replacing Yisraʿel (AYIN) with Yisraʾel (ALEPH)...")
print("(Processing body, tables, TOC, headers, footers)")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total Yisraʿel (AYIN) → Yisraʾel (ALEPH): {total_fixes}")
