"""
Apply half-ring modifier replacements to YY .docx files.
Reads italic_quote_words_fixed.txt and replaces original words in:
- Body text (including Table of Contents)
- Headers (all sections)
Also sets all half-ring characters (U+02BE, U+02BF) to 'Yada Towrah' font.
"""
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt
from copy import deepcopy

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')
BACKUP_DIR = DOCS_DIR / '_halfring_backup'
FIXED_FILE = Path(r'C:\Users\Joe\Work\dev\yada\translations\italic_quote_words_fixed.txt')

# Characters treated as single-quote variants
QUOTE_CHARS = set("'\u2018\u2019`\u00B4\u02BC")
HALF_RING_ALEPH = '\u02BE'  # ʾ
HALF_RING_AYIN = '\u02BF'   # ʿ
ALL_QUOTE_LIKE = QUOTE_CHARS | {HALF_RING_ALEPH, HALF_RING_AYIN}

YT_FONT = 'Yada Towrah'

# Regex character class matching any quote variant
QUOTE_CLASS = "['\u2018\u2019`\u00B4\u02BC\u02BE\u02BF]"


def is_quote(ch):
    return ch in ALL_QUOTE_LIKE


def parse_fixed_file():
    """Return dict: original_word -> fixed_word (unique per word)."""
    lines = open(FIXED_FILE, encoding='utf-8').readlines()
    pairs = {}
    for line in lines[2:]:
        if not line.strip():
            continue
        fixed = line[:40].strip()
        original = line[41:71].strip()
        if original and fixed and original != fixed:
            if original not in pairs:
                pairs[original] = fixed
    return pairs


def build_pattern(word):
    """Build regex matching word with any quote variant at quote positions."""
    parts = []
    for ch in word:
        if is_quote(ch):
            parts.append(QUOTE_CLASS)
        else:
            parts.append(re.escape(ch))
    # Word boundaries: not preceded/followed by a letter
    return re.compile(r'(?<![a-zA-Z])' + ''.join(parts) + r'(?![a-zA-Z])')


def get_char_replacements(original, fixed):
    """Return dict: char_offset -> new_char for positions that differ."""
    return {i: f for i, (o, f) in enumerate(zip(original, fixed)) if o != f}


def process_paragraph(para, patterns):
    """Apply word replacements in a paragraph. Returns match count."""
    if not para.runs:
        return 0

    # Build full text and character-to-run mapping
    text = ''
    cmap = []  # (run_index, char_index_in_run)
    for ri, run in enumerate(para.runs):
        for ci, ch in enumerate(run.text):
            text += ch
            cmap.append((ri, ci))

    if not text or not any(is_quote(ch) for ch in text):
        return 0

    count = 0
    used = set()

    for pat, char_reps in patterns:
        for m in pat.finditer(text):
            s, e = m.start(), m.end()
            if any(i in used for i in range(s, e)):
                continue
            # Apply character replacements
            for offset, new_ch in char_reps.items():
                pos = s + offset
                if pos < len(cmap):
                    ri, ci = cmap[pos]
                    t = list(para.runs[ri].text)
                    t[ci] = new_ch
                    para.runs[ri].text = ''.join(t)
            for i in range(s, e):
                used.add(i)
            count += 1

    return count


def set_half_ring_font_in_run(run):
    """If a run contains half-ring characters mixed with other text,
    split it so the half-rings get their own run with Yada Towrah font.
    If the entire run is half-rings, just set the font.
    Returns number of half-ring chars found."""
    text = run.text
    if not text:
        return 0

    hr_count = sum(1 for ch in text if ch in (HALF_RING_ALEPH, HALF_RING_AYIN))
    if hr_count == 0:
        return 0

    # If ALL chars are half-rings, just set the font on this run
    if hr_count == len(text):
        run.font.name = YT_FONT
        return hr_count

    # Mixed: need to split the run into segments
    # We'll handle this at the paragraph level instead
    return hr_count


def apply_yt_font_to_half_rings(para):
    """Ensure all half-ring characters in the paragraph use Yada Towrah font.
    Splits runs as needed so half-rings get their own run with the correct font."""
    if not para.runs:
        return 0

    # Check if any run has half-rings
    full_text = ''.join(r.text for r in para.runs)
    if HALF_RING_ALEPH not in full_text and HALF_RING_AYIN not in full_text:
        return 0

    hr_count = 0

    # Process runs - we need to split runs that have mixed content
    # Work backwards to avoid index issues when inserting new runs
    from lxml import etree
    from docx.oxml.ns import qn

    runs_to_process = list(para.runs)
    for run in runs_to_process:
        text = run.text
        if not text:
            continue

        has_hr = any(ch in (HALF_RING_ALEPH, HALF_RING_AYIN) for ch in text)
        if not has_hr:
            continue

        # Check if entire run is half-rings
        all_hr = all(ch in (HALF_RING_ALEPH, HALF_RING_AYIN) for ch in text)
        if all_hr:
            run.font.name = YT_FONT
            hr_count += len(text)
            continue

        # Mixed content: split run into segments
        # Group consecutive chars by whether they're half-rings
        segments = []
        current_text = ''
        current_is_hr = text[0] in (HALF_RING_ALEPH, HALF_RING_AYIN)
        for ch in text:
            ch_is_hr = ch in (HALF_RING_ALEPH, HALF_RING_AYIN)
            if ch_is_hr == current_is_hr:
                current_text += ch
            else:
                segments.append((current_text, current_is_hr))
                current_text = ch
                current_is_hr = ch_is_hr
        segments.append((current_text, current_is_hr))

        if len(segments) <= 1:
            # Shouldn't happen given the checks above, but just in case
            if current_is_hr:
                run.font.name = YT_FONT
                hr_count += len(text)
            continue

        # Set first segment text on the original run
        run.text = segments[0][0]
        if segments[0][1]:
            run.font.name = YT_FONT
            hr_count += len(segments[0][0])

        # Create new runs for remaining segments, copying formatting from original
        run_element = run._element
        for seg_text, seg_is_hr in segments[1:]:
            # Clone the run element
            new_run_elem = deepcopy(run_element)
            # Set text
            t_elem = new_run_elem.find(qn('w:t'))
            if t_elem is None:
                t_elem = etree.SubElement(new_run_elem, qn('w:t'))
            t_elem.text = seg_text
            # Preserve spaces
            t_elem.set(qn('xml:space'), 'preserve')

            # Set font for half-ring segments
            if seg_is_hr:
                rPr = new_run_elem.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(new_run_elem, qn('w:rPr'))
                    new_run_elem.insert(0, rPr)
                # Set rFonts
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                rFonts.set(qn('w:ascii'), YT_FONT)
                rFonts.set(qn('w:hAnsi'), YT_FONT)
                rFonts.set(qn('w:cs'), YT_FONT)
                hr_count += len(seg_text)

            # Insert after the original run
            run_element.addnext(new_run_elem)
            run_element = new_run_elem  # next segment goes after this one

    return hr_count


def process_document(doc_path, patterns):
    """Process a single document. Returns (replacement_count, font_count)."""
    doc = Document(str(doc_path))
    rep_count = 0
    font_count = 0

    # Collect all paragraphs: body + tables + headers
    all_paras = list(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for section in doc.sections:
        try:
            if not section.header.is_linked_to_previous:
                all_paras.extend(section.header.paragraphs)
        except Exception:
            pass

    # Step 1: Apply word replacements
    for para in all_paras:
        rep_count += process_paragraph(para, patterns)

    # Step 2: Set Yada Towrah font on all half-ring characters
    # Re-collect paragraphs since runs may have changed
    all_paras2 = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras2.extend(cell.paragraphs)
    for section in doc.sections:
        try:
            if not section.header.is_linked_to_previous:
                all_paras2.extend(section.header.paragraphs)
        except Exception:
            pass

    for para in all_paras2:
        font_count += apply_yt_font_to_half_rings(para)

    doc.save(str(doc_path))
    return rep_count, font_count


def main():
    pairs = parse_fixed_file()
    print(f'Loaded {len(pairs)} unique word replacements')

    # Build patterns sorted by word length descending (longest first)
    patterns = []
    for original, fixed in sorted(pairs.items(), key=lambda x: -len(x[0])):
        pat = build_pattern(original)
        reps = get_char_replacements(original, fixed)
        if reps:
            patterns.append((pat, reps))
    print(f'Built {len(patterns)} replacement patterns')

    # Create backup directory
    BACKUP_DIR.mkdir(exist_ok=True)

    files = sorted(DOCS_DIR.glob('YY*.docx'))
    files = [f for f in files if '_backup' not in f.name and '_updated' not in f.name
             and '_halfring' not in f.name]
    print(f'Processing {len(files)} files...\n')

    grand_rep = 0
    grand_font = 0
    for i, fp in enumerate(files):
        print(f'  [{i+1}/{len(files)}] {fp.name}', end='', flush=True)

        # Backup
        bak = BACKUP_DIR / fp.name
        if not bak.exists():
            shutil.copy2(fp, bak)

        try:
            rc, fc = process_document(fp, patterns)
            print(f' -> {rc} replacements, {fc} half-ring font fixes')
            grand_rep += rc
            grand_font += fc
        except Exception as e:
            print(f' ERROR: {e}')

    print(f'\nDone! {grand_rep} word replacements, {grand_font} half-ring font fixes')
    print(f'Backups saved to: {BACKUP_DIR}')


if __name__ == '__main__':
    main()
