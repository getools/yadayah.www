"""
Find where unresolved cite references appear in the Word documents.
"""

import sys
import os
import glob
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Unresolved cite variants we're looking for
SEARCH_TERMS = [
    "Shamuwʿfel",
    "Shamuw'fel",
    "1 Shamuwʿfel",
    "2 Shamuwʿfel",
    "Danyʾfel",
    "Dany'fel",
    "Yachezqʾfel",
    "Yowʾfel",
    "Yow'fel",
]

docs_dir = r'C:\users\joe\work\dev\yada\docs'
pattern = os.path.join(docs_dir, 'YY*.docx')
all_files = glob.glob(pattern)

# Filter out temp and backup files
docx_files = []
for f in all_files:
    name = os.path.basename(f)
    if name.startswith('~$'):
        continue
    name_lower = name.lower()
    if any(x in name_lower for x in ['_backup', '_updated', '_final', '_test']):
        continue
    docx_files.append(f)

print("Searching for unresolved cite references...\n")

results = {}

for doc_path in sorted(docx_files):
    doc_name = os.path.basename(doc_path)

    try:
        doc = Document(doc_path)
    except:
        continue

    found_in_doc = []

    # Search all paragraphs
    for para in doc.paragraphs:
        text = para.text
        for term in SEARCH_TERMS:
            if term in text:
                # Find the context (surrounding text)
                idx = text.find(term)
                start = max(0, idx - 30)
                end = min(len(text), idx + len(term) + 30)
                context = text[start:end].replace('\n', ' ')

                found_in_doc.append({
                    'term': term,
                    'context': context
                })

    # Search tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    for term in SEARCH_TERMS:
                        if term in text:
                            idx = text.find(term)
                            start = max(0, idx - 30)
                            end = min(len(text), idx + len(term) + 30)
                            context = text[start:end].replace('\n', ' ')

                            found_in_doc.append({
                                'term': term,
                                'context': context
                            })

    if found_in_doc:
        results[doc_name] = found_in_doc

# Print results
for doc_name, findings in results.items():
    print(f"\n{doc_name}:")

    # Group by term
    by_term = {}
    for finding in findings:
        term = finding['term']
        if term not in by_term:
            by_term[term] = []
        by_term[term].append(finding['context'])

    for term, contexts in by_term.items():
        print(f"  '{term}' - {len(contexts)} occurrences")
        # Show first 3 contexts
        for i, context in enumerate(contexts[:3]):
            print(f"    [{i+1}] ...{context}...")

print(f"\n\nSummary:")
print(f"Documents with unresolved cites: {len(results)}")

# Count by term
term_counts = {}
for findings in results.values():
    for finding in findings:
        term = finding['term']
        term_counts[term] = term_counts.get(term, 0) + 1

print("\nOccurrences by term:")
for term, count in sorted(term_counts.items(), key=lambda x: -x[1]):
    print(f"  {term}: {count}")
