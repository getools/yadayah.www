"""Check how TOC is structured in YY docs - SDT vs regular paragraphs."""
from docx import Document
from lxml import etree
from docx.oxml.ns import qn
from pathlib import Path

DOCS_DIR = Path(r'C:\Users\Joe\Work\dev\yada\docs')

# Pick one doc to inspect
doc_path = DOCS_DIR / 'YY-s01v02-An Intro to God-Mitswah-Instructions.docx'
doc = Document(str(doc_path))

body = doc.element.body

# Check for SDT (Structured Document Tags) which often wrap TOC
sdts = body.findall(qn('w:sdt'))
print(f'SDT elements in body: {len(sdts)}')

for i, sdt in enumerate(sdts):
    # Check if it's a TOC
    sdtPr = sdt.find(qn('w:sdtPr'))
    if sdtPr is not None:
        docPartObj = sdtPr.find(qn('w:docPartObj'))
        if docPartObj is not None:
            gallery = docPartObj.find(qn('w:docPartGallery'))
            if gallery is not None:
                print(f'  SDT {i}: gallery={gallery.get(qn("w:val"))}')

    # Count paragraphs inside this SDT
    sdtContent = sdt.find(qn('w:sdtContent'))
    if sdtContent is not None:
        paras = sdtContent.findall(qn('w:p'))
        print(f'  SDT {i}: {len(paras)} paragraphs inside')
        # Show first few paragraph texts
        for j, p in enumerate(paras[:5]):
            runs = p.findall('./' + qn('w:r'))
            text = ''
            for r in runs:
                t = r.find(qn('w:t'))
                if t is not None and t.text:
                    text += t.text
            if text.strip():
                print(f'    para {j}: "{text[:80]}"')

# Also check: are TOC paragraphs in doc.paragraphs?
print(f'\nTotal doc.paragraphs: {len(doc.paragraphs)}')

# Check for field codes that indicate TOC
for i, para in enumerate(doc.paragraphs[:80]):
    text = para.text
    if 'TOC' in text or (text.strip() and i < 50):
        # Check for fldChar in XML
        fld_chars = para._element.findall('.//' + qn('w:fldChar'))
        fld_simples = para._element.findall('.//' + qn('w:fldSimple'))
        instr = para._element.findall('.//' + qn('w:instrText'))
        instr_text = ' '.join(i.text for i in instr if i.text)
        if fld_chars or fld_simples or instr:
            print(f'  para[{i}]: "{text[:60]}" [field: {instr_text[:40]}]')

# Check for any quote chars remaining in early paragraphs (TOC area)
QUOTE_CHARS = set("'\u2018\u2019`\u00B4\u02BC")
print('\nQuote chars in first 80 paragraphs:')
for i, para in enumerate(doc.paragraphs[:80]):
    text = para.text
    quotes = [ch for ch in text if ch in QUOTE_CHARS]
    if quotes:
        print(f'  para[{i}]: {len(quotes)} quotes in "{text[:70]}"')

# Check SDT paragraphs for remaining quotes
print('\nQuote chars in SDT paragraphs:')
for sdt in sdts:
    sdtContent = sdt.find(qn('w:sdtContent'))
    if sdtContent is not None:
        for p in sdtContent.findall(qn('w:p')):
            runs = p.findall('.//' + qn('w:r'))
            text = ''
            for r in runs:
                t = r.find(qn('w:t'))
                if t is not None and t.text:
                    text += t.text
            quotes = [ch for ch in text if ch in QUOTE_CHARS]
            if quotes:
                print(f'  SDT para: {len(quotes)} quotes in "{text[:70]}"')
