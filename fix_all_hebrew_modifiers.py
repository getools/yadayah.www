"""
Replace Hebrew transliterations with proper ALEPH (ʾ) and AYIN (ʿ) modifiers.
Handles all apostrophe variations: ' ' ' `
Processes body text, tables, TOC, headers, and footers.
Ensures modifiers use Yada Towrah font.
"""

import sys, os, glob, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

ALEPH = chr(0x02BE)  # ʾ (right half ring)
AYIN = chr(0x02BF)   # ʿ (left half ring)

# All possible apostrophe variations to match
APOSTROPHE_VARIANTS = [
    "'",      # Regular apostrophe (U+0027)
    "'",      # Right single quote (U+2019)
    "'",      # Left single quote (U+2018)
    "`",      # Backtick (U+0060)
    "´",      # Acute accent (U+00B4)
    "ʼ",      # Modifier letter apostrophe (U+02BC)
]

# Create regex pattern to match any apostrophe variant
APOSTROPHE_PATTERN = '[' + ''.join(re.escape(c) for c in APOSTROPHE_VARIANTS) + ']'

# Replacement mappings: (pattern_before_apostrophe, pattern_after_apostrophe, use_aleph)
# use_aleph=True means ʾ (ALEPH), use_aleph=False means ʿ (AYIN)
REPLACEMENTS = [
    # ALEPH replacements (ʾ - right half ring)
    ("Yisrae", "l", True),   # Yisrae'l → Yisraʾel
    (r"^'Arek$", "", True),  # 'Arek → ʾArek
    (r"^'arek$", "", True),  # 'arek → ʾarek
    (r"^'Akad$", "", True),  # 'Akad → ʾAkad
    (r"^'akad$", "", True),  # 'akad → ʾakad
    (r"^'Ashuwr$", "", True),  # 'Ashuwr → ʾAshuwr
    (r"^'ashuwr$", "", True),  # 'ashuwr → ʾashuwr
    (r"^'elowah$", "", True),  # 'elowah → ʾelowah
    (r"^'elowahym$", "", True),  # 'elowahym → ʾelowahym
    (r"^'Adam$", "", True),  # 'Adam → ʾAdam
    (r"^ʾadam$", "", True),  # ʾadam → ʾAdam (fix case)
    ("Miqra", "$", True),  # Miqra' → Miqraʾ (at end of word)
    (r"^'Abraham's$", "", True),  # 'Abraham's → ʾAbraham's
    (r"^'Abshalowm's$", "", True),  # 'Abshalowm's → ʾAbshalowm's
    (r"^'Abraham$", "", True),  # 'Abraham → ʾAbraham
    (r"^'Ayil$", "", True),  # 'Ayil → ʾAyil
    (r"^'Abshalowm$", "", True),  # 'Abshalowm → ʾAbshalowm
    (r"^'Amal$", "", True),  # 'Amal → ʾAmal
    (r"^'amal$", "", True),  # 'amal → ʾamal
    (r"^'atem$", "", True),  # 'atem → ʾatem
    (r"^'any$", "", True),  # 'any → ʾany
    (r"^'Edowm$", "", True),  # 'Edowm → ʾEdowm
    (r"^'Amown$", "", True),  # 'Amown → ʾAmown
    (r"^'atah$", "", True),  # 'atah → ʾatah
    (r"^'el$", "", True),  # 'el → ʾel
    (r"^'ElYah$", "", True),  # 'ElYah → ʾElYah
    (r"^'anachnuw$", "", True),  # 'anachnuw → ʾanachnuw
    (r"^'Elowym$", "", True),  # 'Elowym → ʾElowym
    (r"^'ozen$", "", True),  # 'ozen → ʾozen
    (r"^'Abyb$", "", True),  # 'Abyb → ʾAbyb
    (r"^'ishah$", "", True),  # 'ishah → ʾishah
    (r"^'chasydy$", "", True),  # 'chasydy → ʾchasydy

    # AYIN replacements (ʿ - left half ring)
    ("Shin", "ar", False),  # Shin'ar → Shinʿar
    ("shin", "ar", False),  # shin'ar → shinʿar
    (r"^'arabah$", "", False),  # 'arabah → ʿarabah
    (r"^'owd$", "", False),  # 'owd → ʿowd
    (r"^'Eden$", "", False),  # 'Eden → ʿEden
    (r"^'aleichem$", "", False),  # 'aleichem → ʿaleichem
    ("Ya", "aqob", False),  # Ya'aqob → Yaʿaqob
    ("Mow", "ab", False),  # Mow'ab → Mowʿab
    ("Mow", "edym", False),  # Mow'edym → Mowʿedym
    ("Yahowsha", "$", False),  # Yahowsha' → Yahowshaʿ (at end)
    ("Naby", "$", False),  # Naby' → Nabyʿ (at end)
    ("Ma", "achach", False),  # Ma'achach → Maʿachach
    ("Yow", "ab", False),  # Yow'ab → Yowʿab
    ("Yisra", "el", False),  # Yisra'el → Yisraʿel
    ("Ga", "al", False),  # Ga'al → Gaʿal
    ("yasha", "$", False),  # yasha' → yashaʿ (at end)
    ("Yisra", "elites", False),  # Yisra'elites → Yisraʿelites
    ("Mow", "abites", False),  # Mow'abites → Mowʿabites
    ("Yada", "$", False),  # Yada' → Yadaʿ (at end)
    ("Zarowa", "$", False),  # Zarowa' → Zarowaʿ (at end)
    ("Yow", "el", False),  # Yow'el → Yowʿel
    ("Yasha", "$", False),  # Yasha' → Yashaʿ (at end)
    ("Yahowshuwa", "$", False),  # Yahowshuwa' → Yahowshuwaʿ (at end)
    ("mari", "yth", False),  # mari'yth → mariʿyth
    ("mal", "akym", False),  # mal'akym → malʿakym
    ("yesha", "y", False),  # yesha'y → yeshaʿy
    ("arba", "ym", False),  # arba'ym → arbaʿym
    ("sha", "b", False),  # sha'b → shaʿb
    ("Howsha", "$", False),  # Howsha' → Howshaʿ (at end)
    ("mits", "ar", False),  # mits'ar → mitsʿar
    ("Sha", "uwl", False),  # Sha'uwl → Shaʿuwl
]

def compile_patterns():
    """Compile regex patterns for efficient matching."""
    compiled = []
    for before, after, use_aleph in REPLACEMENTS:
        # Handle end-of-word markers
        after_pattern = after.replace('$', '')

        # Build the full pattern
        if before.startswith('^') and before.endswith('$'):
            # Full word match (e.g., ^'el$)
            word = before[1:-1]  # Remove ^ and $
            pattern = f"^{word[0]}{APOSTROPHE_PATTERN}{word[1:]}$"
        elif before.startswith('^'):
            # Start of word (e.g., ^'Adam)
            word = before[1:]
            pattern = f"^{word[0]}{APOSTROPHE_PATTERN}{word[1:]}"
        elif after == '$':
            # End of word with apostrophe (e.g., Naby')
            pattern = f"{before}{APOSTROPHE_PATTERN}$"
        else:
            # Middle of word
            pattern = f"{before}{APOSTROPHE_PATTERN}{after_pattern}"

        compiled.append((re.compile(pattern), before, after, use_aleph))

    return compiled

COMPILED_PATTERNS = compile_patterns()

def replace_in_text(text):
    """
    Replace apostrophes with modifiers in text.
    Returns list of (original_text, new_text, modifier_char, modifier_index) tuples.
    """
    if not text:
        return []

    replacements = []

    for pattern, before, after, use_aleph in COMPILED_PATTERNS:
        modifier = ALEPH if use_aleph else AYIN

        # Find all matches
        for match in pattern.finditer(text):
            matched_text = match.group(0)

            # Find the apostrophe position in the matched text
            apostrophe_pos = None
            for char in APOSTROPHE_VARIANTS:
                if char in matched_text:
                    apostrophe_pos = matched_text.index(char)
                    break

            if apostrophe_pos is not None:
                # Build replacement: before + modifier + after
                if before.startswith('^') and before.endswith('$'):
                    word = before[1:-1]
                    new_text = word[0] + modifier + word[1:]
                elif before.startswith('^'):
                    word = before[1:]
                    new_text = word[0] + modifier + word[1:] + matched_text[apostrophe_pos+1:]
                elif after == '$':
                    new_text = matched_text[:apostrophe_pos] + modifier
                else:
                    after_text = after.replace('$', '')
                    new_text = matched_text[:apostrophe_pos] + modifier + after_text

                if matched_text != new_text:
                    # Calculate absolute position in original text
                    abs_pos = match.start() + apostrophe_pos
                    replacements.append((matched_text, new_text, modifier, abs_pos))

    return replacements

def split_and_apply_modifier(run, text, modifier_char, modifier_pos):
    """
    Split run at modifier position and apply Yada Towrah font to modifier.
    Returns True if replacement was made.
    """
    if modifier_pos < 0 or modifier_pos >= len(text):
        return False

    # Get parent paragraph
    para = run._element.getparent()
    if para is None:
        return False

    run_idx = list(para).index(run._element)

    # Get original formatting
    rPr = run._element.find(qn('w:rPr'))

    def clone_rPr(exclude_font=False):
        """Clone run properties, optionally excluding font."""
        new_rPr = OxmlElement('w:rPr')
        if rPr is not None:
            for child in rPr:
                if exclude_font and child.tag == qn('w:rFonts'):
                    continue
                import copy
                new_rPr.append(copy.deepcopy(child))
        return new_rPr

    def create_run_with_text(content, is_modifier=False):
        """Create new run with text and formatting."""
        new_run = OxmlElement('w:r')
        new_rPr = clone_rPr(exclude_font=is_modifier)

        if is_modifier:
            # Add Yada Towrah font for modifier
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Yada Towrah')
            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
            new_rPr.append(rFonts)

        new_run.append(new_rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        new_run.append(t)

        return new_run

    # Split text into three parts: before, modifier, after
    before_text = text[:modifier_pos]
    after_text = text[modifier_pos+1:]  # Skip the original apostrophe

    # Remove original run
    para.remove(run._element)

    # Insert new runs
    current_idx = run_idx

    if before_text:
        before_run = create_run_with_text(before_text, is_modifier=False)
        para.insert(current_idx, before_run)
        current_idx += 1

    # Insert modifier with Yada Towrah font
    modifier_run = create_run_with_text(modifier_char, is_modifier=True)
    para.insert(current_idx, modifier_run)
    current_idx += 1

    if after_text:
        after_run = create_run_with_text(after_text, is_modifier=False)
        para.insert(current_idx, after_run)

    return True

def process_paragraph(para):
    """Process all runs in a paragraph."""
    fixes = 0

    # Process runs (iterate over copy since we modify the list)
    for run in list(para.runs):
        if not run.text:
            continue

        replacements = replace_in_text(run.text)

        for matched_text, new_text, modifier, pos in replacements:
            # Find which character is the modifier
            modifier_idx = new_text.index(modifier)

            # Split and apply
            if split_and_apply_modifier(run, run.text, modifier, pos):
                fixes += 1
                break  # Run has been modified, move to next

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    # Process tables (including TOC)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Replacing Hebrew transliterations with proper modifiers...")
print("(Processing body, tables, TOC, headers, footers)")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total replacements: {total_fixes}")
