"""
Replace ALL apostrophe variants in Hebrew transliterations with proper ALEPH (ʾ) and AYIN (ʿ) modifiers.
FINAL VERSION: Handles cross-run splits where apostrophe and name are in different runs.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

ALEPH = chr(0x02BE)  # ʾ
AYIN = chr(0x02BF)   # ʿ

# ALL possible apostrophe variations
APOSTROPHE_VARIANTS = [
    chr(0x0027),  # Regular apostrophe '
    chr(0x2019),  # Right single quote ' (curly apostrophe)
    chr(0x2018),  # Left single quote ' (left curly quote)
    chr(0x0060),  # Backtick `
    chr(0x00B4),  # Acute accent ´
    chr(0x02BC),  # Modifier letter apostrophe ʼ
]

# Base patterns - names that should have modifiers
ALEPH_NAMES = [
    "Abram", "Adam", "Abraham", "el", "Arek", "arek", "Akad", "akad",
    "Ashuwr", "ashuwr", "elowah", "elowahym", "Abraham's", "Abshalowm's",
    "Ayil", "Abshalowm", "Amal", "amal", "atem", "any", "Edowm", "Amown",
    "Amnown", "atah", "ElYah", "anachnuw", "Elowym", "ozen", "Abyb",
    "ishah", "chasydy"
]

AYIN_NAMES = [
    "arabah", "owd", "Eden", "aleichem"
]

# Patterns with apostrophe in middle
MID_ALEPH = [
    ("Yisrae'l", "Yisraʾel"),
    ("Miqra'", "Miqraʾ"),
]

MID_AYIN = [
    ("Yisra'el", "Yisraʿel"),
    ("Ya'aqob", "Yaʿaqob"),
    ("Mow'ab", "Mowʿab"),
    ("Mow'edym", "Mowʿedym"),
    ("Yahowsha'", "Yahowshaʿ"),
    ("Yow'ab", "Yowʿab"),
    ("Ga'al", "Gaʿal"),
    ("yasha'", "yashaʿ"),
    ("Yisra'elites", "Yisraʿelites"),
    ("Mow'abites", "Mowʿabites"),
    ("Yada'", "Yadaʿ"),
    ("Zarowa'", "Zarowaʿ"),
    ("Yow'el", "Yowʿel"),
    ("Yasha'", "Yashaʿ"),
    ("Howsha'", "Howshaʿ"),
    ("Sha'uwl", "Shaʿuwl"),
    ("Shin'ar", "Shinʿar"),
    ("shin'ar", "shinʿar"),
    ("Naby'", "Nabyʿ"),
    ("Ma'achach", "Maʿachach"),
    ("Yahowshuwa'", "Yahowshuwaʿ"),
    ("mari'yth", "mariʿyth"),
    ("mal'akym", "malʿakym"),
    ("yesha'y", "yeshaʿy"),
    ("arba'ym", "arbaʿym"),
    ("sha'b", "shaʿb"),
    ("mits'ar", "mitsʿar"),
]

def merge_two_runs(para, run1_idx):
    """Merge run at run1_idx with the next run. Returns new merged text."""
    runs = para.runs
    if run1_idx >= len(runs) - 1:
        return None

    run1 = runs[run1_idx]
    run2 = runs[run1_idx + 1]

    # Combine text
    merged_text = run1.text + run2.text

    # Update first run with merged text
    run1.text = merged_text

    # Remove second run from XML
    p_elem = para._element
    r2_elem = run2._element
    p_elem.remove(r2_elem)

    return merged_text

def split_run_with_font(run_elem, text, modifier_char):
    """Split run to isolate modifier with Yada Towrah font."""
    if modifier_char not in text:
        return False

    para = run_elem.getparent()
    if para is None:
        return False

    run_idx = list(para).index(run_elem)
    rPr = run_elem.find(qn('w:rPr'))

    def clone_rPr(exclude_font=False):
        new_rPr = OxmlElement('w:rPr')
        if rPr is not None:
            for child in rPr:
                if exclude_font and child.tag == qn('w:rFonts'):
                    continue
                import copy
                new_rPr.append(copy.deepcopy(child))
        return new_rPr

    def create_run(content, is_modifier=False):
        new_run = OxmlElement('w:r')
        new_rPr = clone_rPr(exclude_font=is_modifier)

        if is_modifier:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Yada Towrah')
            rFonts.set(qn('w:hAnsi'), 'Yada Towrah')
            new_rPr.append(rFonts)

        new_run.append(new_rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        new_run.append(t)

        return new_run

    para.remove(run_elem)

    current_idx = run_idx
    parts = text.split(modifier_char)

    for i, part in enumerate(parts):
        if i > 0:
            modifier_run = create_run(modifier_char, is_modifier=True)
            para.insert(current_idx, modifier_run)
            current_idx += 1

        if part:
            part_run = create_run(part, is_modifier=False)
            para.insert(current_idx, part_run)
            current_idx += 1

    return True

def process_paragraph(para):
    """Process paragraph - handles both same-run and cross-run cases."""
    fixes = 0

    # PHASE 1: Handle cross-run cases (apostrophe at end of run N, name at start of run N+1)
    run_idx = 0
    while run_idx < len(para.runs) - 1:
        run = para.runs[run_idx]
        next_run = para.runs[run_idx + 1]

        if not run.text or not next_run.text:
            run_idx += 1
            continue

        # Check if run ends with apostrophe and next run starts with a name
        ends_with_apos = False
        apos_char = None
        for apos in APOSTROPHE_VARIANTS:
            if run.text.endswith(apos):
                ends_with_apos = True
                apos_char = apos
                break

        if ends_with_apos:
            # Check if next run starts with ALEPH name
            for name in ALEPH_NAMES:
                if next_run.text.startswith(name):
                    # Merge runs
                    merged_text = merge_two_runs(para, run_idx)
                    if merged_text:
                        # Now replace in the merged run
                        pattern = f"{apos_char}{name}"
                        replacement = f"{ALEPH}{name}"
                        if pattern in merged_text:
                            new_text = merged_text.replace(pattern, replacement, 1)
                            para.runs[run_idx].text = new_text

                            # Split on ALEPH
                            if split_run_with_font(para.runs[run_idx]._element, new_text, ALEPH):
                                fixes += 1
                                # Don't increment run_idx, reprocess from same position
                                continue

            # Check if next run starts with AYIN name
            for name in AYIN_NAMES:
                if next_run.text.startswith(name):
                    merged_text = merge_two_runs(para, run_idx)
                    if merged_text:
                        pattern = f"{apos_char}{name}"
                        replacement = f"{AYIN}{name}"
                        if pattern in merged_text:
                            new_text = merged_text.replace(pattern, replacement, 1)
                            para.runs[run_idx].text = new_text

                            if split_run_with_font(para.runs[run_idx]._element, new_text, AYIN):
                                fixes += 1
                                continue

        run_idx += 1

    # PHASE 2: Handle same-run cases
    run_idx = 0
    while run_idx < len(para.runs):
        run = para.runs[run_idx]

        if not run.text:
            run_idx += 1
            continue

        original_text = run.text
        new_text = original_text

        # Apply all MID_ALEPH replacements
        for old, new in MID_ALEPH:
            for apos in APOSTROPHE_VARIANTS:
                pattern = old.replace("'", apos)
                if pattern in new_text and pattern != new:
                    new_text = new_text.replace(pattern, new)

        # Apply all MID_AYIN replacements
        for old, new in MID_AYIN:
            for apos in APOSTROPHE_VARIANTS:
                pattern = old.replace("'", apos)
                if pattern in new_text and pattern != new:
                    new_text = new_text.replace(pattern, new)

        # Apply leading apostrophe + name replacements
        for name in ALEPH_NAMES:
            for apos in APOSTROPHE_VARIANTS:
                pattern = f"{apos}{name}"
                replacement = f"{ALEPH}{name}"
                if pattern in new_text and pattern != replacement:
                    new_text = new_text.replace(pattern, replacement)

        for name in AYIN_NAMES:
            for apos in APOSTROPHE_VARIANTS:
                pattern = f"{apos}{name}"
                replacement = f"{AYIN}{name}"
                if pattern in new_text and pattern != replacement:
                    new_text = new_text.replace(pattern, replacement)

        if new_text != original_text:
            run.text = new_text

            has_aleph = ALEPH in new_text
            has_ayin = AYIN in new_text

            if has_aleph and has_ayin:
                # Split ALEPH first
                if split_run_with_font(run._element, new_text, ALEPH):
                    fixes += new_text.count(ALEPH)
                    continue
            elif has_aleph:
                if split_run_with_font(run._element, new_text, ALEPH):
                    fixes += new_text.count(ALEPH)
                    continue
            elif has_ayin:
                if split_run_with_font(run._element, new_text, AYIN):
                    fixes += new_text.count(AYIN)
                    continue

        run_idx += 1

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    for para in doc.paragraphs:
        fixes += process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += process_paragraph(para)

    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += process_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No changes needed ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Final pass: Fixing cross-run apostrophe splits...")
print()

all_files = glob.glob(os.path.join(docs_dir, 'YY*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in sorted(docx_files):
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total cross-run fixes: {total_fixes}")
