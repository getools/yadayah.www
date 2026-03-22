"""
Fix "Strong's" - replace modifier characters (ʿ or ʾ) with regular apostrophe (').
The regular apostrophe should use the same font as surrounding text (Times New Roman).
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import time

AYIN = chr(0x02BF)  # ʿ
ALEPH = chr(0x02BE)  # ʾ
APOSTROPHE = "'"

def fix_strongs_in_paragraph(para):
    """Fix Strong's in a paragraph by replacing modifiers with regular apostrophe."""
    fixes = 0

    # Check if paragraph contains "Strong" and a modifier
    if 'Strong' not in para.text:
        return 0

    if AYIN not in para.text and ALEPH not in para.text:
        return 0

    # Process runs to find and fix "Strongʿ" or "Strongʾ"
    for run in para.runs:
        if not run.text:
            continue

        original = run.text

        # Replace Strongʿ with Strong'
        text = original.replace(f'Strong{AYIN}', f'Strong{APOSTROPHE}')
        text = text.replace(f'strong{AYIN}', f'strong{APOSTROPHE}')

        # Replace Strongʾ with Strong'
        text = text.replace(f'Strong{ALEPH}', f'Strong{APOSTROPHE}')
        text = text.replace(f'strong{ALEPH}', f'strong{APOSTROPHE}')

        if text != original:
            run.text = text
            fixes += 1

    return fixes

def process_document(docx_path):
    """Process a document to fix Strong's."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += fix_strongs_in_paragraph(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += fix_strongs_in_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                fixes += fix_strongs_in_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += fix_strongs_in_paragraph(para)

        if section.footer:
            for para in section.footer.paragraphs:
                fixes += fix_strongs_in_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fixes += fix_strongs_in_paragraph(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No Strong's with modifiers ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing Strong's - replacing modifier characters with regular apostrophe...")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total Strong's instances fixed: {total_fixes}")
