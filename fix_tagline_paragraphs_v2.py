"""
Fix first paragraph after taglines to be justified and indented.

Only process paragraphs after the first chapter heading (not edition info or TOC).
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time
import re

def find_content_start(doc):
    """
    Find where main content starts by looking for first chapter heading.
    Chapter headings use style 'yy_chapter_#' and are typically just a number.
    """
    for i, para in enumerate(doc.paragraphs):
        # Look for chapter number style
        if para.style and 'chapter' in para.style.name.lower():
            return i

        # Or look for numeric chapter headings followed by chapter title
        text = para.text.strip()
        if text.isdigit() and int(text) == 1:
            # Check if next paragraph exists and looks like a chapter title
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if '~' in next_text:  # Chapter titles have ~ separator
                    return i

    # Fallback: look for first "RESOURCES" or numbered chapter
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == 'RESOURCES' or (text.isdigit() and int(text) <= 20):
            return i

    return 0  # Default to beginning if not found

def is_tagline_paragraph(para):
    """
    Detect if paragraph is a tagline:
    - Has italic style (paragraph style has italic font)
    - Ends with ... or … (ellipsis)
    - Left-aligned or None
    """
    if not para.text.strip():
        return False

    # Check if ends with ellipsis
    text = para.text.strip()
    if not text.endswith('...') and not text.endswith('…'):
        return False

    # Check if left-aligned
    if para.alignment not in (WD_ALIGN_PARAGRAPH.LEFT, None):
        return False

    # Check if paragraph style has italic
    if para.style and hasattr(para.style, 'font') and para.style.font.italic:
        return True

    # Check if runs are italic
    if not para.runs:
        return False

    italic_chars = 0
    total_chars = 0

    for run in para.runs:
        if run.text.strip():
            total_chars += len(run.text.strip())
            if run.italic:
                italic_chars += len(run.text.strip())

    if total_chars > 0 and italic_chars / total_chars > 0.8:
        return True

    return False

def fix_paragraph_formatting(doc, content_start_idx, stats):
    """Fix first paragraph after taglines - make justified and indented."""
    paragraphs = doc.paragraphs

    for i in range(content_start_idx, len(paragraphs) - 1):
        para = paragraphs[i]

        # Skip if in TOC or Resources sections
        text = para.text.strip()
        if text == 'RESOURCES':
            stats['in_resources'] = True
        if stats.get('in_resources'):
            continue

        # If current is tagline, find next non-empty paragraph
        if is_tagline_paragraph(para):
            # Find next non-empty paragraph
            next_para = None
            next_idx = i + 1
            while next_idx < len(paragraphs):
                candidate = paragraphs[next_idx]
                if candidate.text.strip():
                    next_para = candidate
                    break
                next_idx += 1

            if next_para is None:
                continue

            # Check if next paragraph needs fixing
            if (next_para.alignment == WD_ALIGN_PARAGRAPH.LEFT or
                next_para.alignment is None):

                # Fix alignment - set to justified
                next_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Fix indentation - set first line indent to 0.5 inches (36 points)
                if next_para.paragraph_format.first_line_indent is None or \
                   next_para.paragraph_format.first_line_indent.pt < 10:
                    next_para.paragraph_format.first_line_indent = Pt(36)
                    stats['para_fixes'] += 1

                    if stats['verbose']:
                        print(f"    Fixed para after: {para.text[:50]}...")

def process_document(docx_path, verbose=False):
    """Process a document."""
    print(f"\nProcessing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return {'para_fixes': 0}

    stats = {
        'para_fixes': 0,
        'verbose': verbose,
        'in_resources': False
    }

    # Find where content starts (after edition info and TOC)
    content_start_idx = find_content_start(doc)

    if verbose:
        print(f"  Content starts at paragraph index: {content_start_idx}")

    # Fix paragraph formatting
    fix_paragraph_formatting(doc, content_start_idx, stats)

    elapsed = time.time() - start_time
    fixes = stats['para_fixes']

    if fixes > 0:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_fixed{ext}"

        try:
            doc.save(output_path)
            print(f"  ✓ Fixed {fixes} paragraphs ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR: {e}")
            return {'para_fixes': 0}
    else:
        print(f"  - No changes ({elapsed:.1f}s)")

    return {'para_fixes': fixes}

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing paragraphs after taglines (skipping edition info and TOC)...")
print()

docx_files = [
    os.path.join(docs_dir, f)
    for f in os.listdir(docs_dir)
    if f.startswith('YY-') and f.endswith('.docx')
    and '_fixed' not in f and not f.startswith('~$')
]

print(f"Found {len(docx_files)} documents\n" + "=" * 80)

total_updated = 0
total_fixes = 0

for path in docx_files:
    result = process_document(path, verbose=False)
    if result['para_fixes'] > 0:
        total_updated += 1
        total_fixes += result['para_fixes']

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Files updated: {total_updated}")
print(f"Paragraph fixes: {total_fixes}")
print(f"\nUpdated files: {docs_dir}/*_fixed.docx")
