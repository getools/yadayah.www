"""
Fix Strong's by replacing Strongʿs or Strongʾs in paragraph text.
Since runs may not be consecutive, we need to work with combined text.
"""

import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import time

AYIN = chr(0x02BF)  # ʿ
ALEPH = chr(0x02BE)  # ʾ
APOSTROPHE = "'"

def fix_strongs_in_xml(para):
    """
    Fix Strong's by replacing modifier in XML text elements.
    We look for Strong + modifier pattern and replace the modifier.
    """
    # First check if paragraph contains the pattern
    if 'Strong' not in para.text:
        return 0

    if f'Strong{AYIN}' not in para.text and f'Strong{ALEPH}' not in para.text:
        return 0

    fixes = 0

    # Iterate through all text elements in the paragraph
    for element in para._element.iter():
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            if element.text:
                original = element.text

                # Check if this element is the modifier after "Strong"
                if element.text in [AYIN, ALEPH]:
                    # Check previous sibling runs to see if any end with "Strong"
                    run_elem = element.getparent()
                    parent = run_elem.getparent()

                    # Get all runs in order
                    runs = list(parent.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r',
                                               namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))

                    run_idx = runs.index(run_elem)

                    # Look at previous runs to find "Strong"
                    found_strong = False
                    for i in range(run_idx - 1, max(-1, run_idx - 10), -1):
                        prev_run = runs[i]
                        for t_elem in prev_run.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                            if t_elem.text and 'Strong' in t_elem.text:
                                found_strong = True
                                break
                        if found_strong:
                            break

                    if found_strong:
                        # Replace modifier with apostrophe
                        element.text = APOSTROPHE
                        fixes += 1

    return fixes

def process_document(docx_path):
    """Process a document."""
    print(f"Processing: {os.path.basename(docx_path)}")
    start_time = time.time()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  ERROR: {e}")
        return 0

    fixes = 0

    # Process body paragraphs
    for para in doc.paragraphs:
        fixes += fix_strongs_in_xml(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fixes += fix_strongs_in_xml(para)

    elapsed = time.time() - start_time

    if fixes > 0:
        try:
            doc.save(docx_path)
            print(f"  ✓ Fixed {fixes} instances ({elapsed:.1f}s)")
        except Exception as e:
            print(f"  ERROR saving: {e}")
            return 0
    else:
        print(f"  - No Strong's with modifiers ({elapsed:.1f}s)")

    return fixes

# Main
docs_dir = r'C:\users\joe\work\dev\yada\docs'
print("Fixing Strong's - replacing modifiers (ʿ/ʾ) with regular apostrophe (')...")
print()

# Get all YY files
all_files = glob.glob(os.path.join(docs_dir, 'YY-s*.docx'))
docx_files = []
for f in all_files:
    name = os.path.basename(f).lower()
    if any(x in name for x in ['_updated', '_final', '_fmt', '_formatted', '_fixed', '_test', '_reverted', '_tagline', '_debug', '_backup']):
        continue
    docx_files.append(f)

print(f"Found {len(docx_files)} documents\n" + "=" * 80 + "\n")

total_fixes = 0
for path in docx_files:
    total_fixes += process_document(path)

print("\n" + "=" * 80)
print("COMPLETE")
print("=" * 80)
print(f"Files processed: {len(docx_files)}")
print(f"Total Strong's modifiers replaced: {total_fixes}")
